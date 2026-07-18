"""Generate editable contract documents from structured contract records."""

from __future__ import annotations

import hashlib
import mimetypes
from copy import deepcopy
from dataclasses import dataclass
from decimal import Decimal
from io import BytesIO
from typing import Protocol

from django.core.files.base import ContentFile
from django.utils import timezone
from docx import Document as WordDocument
from docx.document import Document as WordDocumentType

from operations.models import (
    CenterLegalProfile,
    ContractLegalSnapshot,
    ContractSignedFile,
    Document,
    DonationContract,
    OrganizationServiceContract,
    ServiceContract,
)

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MAX_TEMPLATE_BYTES = 5 * 1024 * 1024
PLACEHOLDER_BLANK = "_______________"


class ContractDocumentError(RuntimeError):
    """Raised when a stored Word template cannot be used safely."""


class _ContractWithNumber(Protocol):
    pk: int | None
    number: str


@dataclass(frozen=True)
class PlaceholderGroup:
    title: str
    placeholders: tuple[str, ...]


@dataclass(frozen=True)
class GeneratedContractFile:
    payload: BytesIO
    filename: str
    document: Document | None = None


PLACEHOLDER_GROUPS = (
    PlaceholderGroup(
        "Центр",
        (
            "center.full_name",
            "center.short_name",
            "center.director_full_name",
            "center.director_short_name",
            "center.director_position",
            "center.authority_basis",
            "center.license_number",
            "center.license_date",
            "center.license_authority",
            "center.ogrn",
            "center.inn",
            "center.kpp",
            "center.legal_address",
            "center.location_address",
            "center.phone",
            "center.email",
            "center.site",
            "center.bank_name",
            "center.bank_bik",
            "center.bank_account",
            "center.bank_corr_account",
        ),
    ),
    PlaceholderGroup(
        "Договор",
        (
            "contract.number",
            "contract.signed_on",
            "contract.valid_from",
            "contract.valid_until",
            "contract.validity",
            "contract.status",
            "contract.type",
            "contract.template",
            "contract.city",
            "contract.amount",
            "contract.amount_words",
            "contract.monthly_amount",
            "contract.payment_due_days",
        ),
    ),
    PlaceholderGroup(
        "Получатель",
        (
            "child.full_name",
            "child.birth_date",
            "child.phone",
            "child.email",
            "child.address",
        ),
    ),
    PlaceholderGroup(
        "Представитель",
        (
            "representative.full_name",
            "representative.relationship",
            "representative.phone",
            "representative.phone_alt",
            "representative.email",
            "representative.passport_series",
            "representative.passport_number",
            "representative.passport_issued_by",
            "representative.passport_issued_on",
            "representative.registration_address",
            "representative.signs_contract",
            "representative.receives_schedule",
            "representative.is_payer",
        ),
    ),
    PlaceholderGroup(
        "Контрагент",
        (
            "counterparty.name",
            "counterparty.type",
            "counterparty.inn",
            "counterparty.kpp",
            "counterparty.ogrn",
            "counterparty.legal_address",
            "counterparty.postal_address",
            "counterparty.bank_details",
            "counterparty.contact_person",
            "counterparty.phone",
            "counterparty.email",
            "counterparty.signer_full_name",
            "counterparty.signer_position",
            "counterparty.authority_basis",
            "counterparty.bank_name",
            "counterparty.bank_bik",
            "counterparty.bank_account",
            "counterparty.bank_corr_account",
        ),
    ),
    PlaceholderGroup(
        "Финансирование",
        (
            "funding_source.name",
            "funding_source.type",
            "funding_source.starts_on",
            "funding_source.ends_on",
            "funding_source.transfer_policy",
            "funding_source.project_name",
        ),
    ),
    PlaceholderGroup(
        "Пожертвование",
        (
            "donation.amount_limit",
            "donation.amount",
            "donation.monthly_amount",
            "donation.periodicity",
        ),
    ),
    PlaceholderGroup(
        "Спецификация услуг",
        (
            "service_spec.rows",
            "service_spec.service_name",
            "service_spec.quantity",
            "service_spec.unit",
            "service_spec.hours",
            "service_spec.price",
            "service_spec.amount",
            "service_spec.period",
        ),
    ),
    PlaceholderGroup(
        "Сертификат",
        (
            "certificate.type",
            "certificate.number",
            "certificate.total_amount",
            "certificate.remaining_amount",
            "certificate.valid_from",
            "certificate.valid_until",
            "certificate.payer_name",
        ),
    ),
)


def placeholder_reference_groups() -> list[PlaceholderGroup]:
    return list(PLACEHOLDER_GROUPS)


def _empty_placeholder_values() -> dict[str, str]:
    return {
        placeholder: PLACEHOLDER_BLANK
        for group in PLACEHOLDER_GROUPS
        for placeholder in group.placeholders
    }


def _safe_filename(prefix: str, contract: _ContractWithNumber) -> str:
    source = contract.number or str(contract.pk or "")
    suffix = "".join(ch if ch.isalnum() or ch in "-_" else "_" for ch in source).strip("_")
    return f"{prefix}_{suffix or 'draft'}.docx"


def _date_label(value) -> str:
    if not value:
        return PLACEHOLDER_BLANK
    return value.strftime("%d.%m.%Y")


def _money_label(value: Decimal | None) -> str:
    if value is None:
        return "без лимита"
    return f"{value:,.2f}".replace(",", " ").replace(".", ",") + " ₽"


def _quantity_label(value: Decimal | None) -> str:
    if value is None:
        return PLACEHOLDER_BLANK
    text = format(value.normalize(), "f")
    if "." in text:
        text = text.rstrip("0").rstrip(".")
    return text or "0"


def _text(value: object, fallback: str = PLACEHOLDER_BLANK) -> str:
    if value is None or value == "":
        return fallback
    return str(value)


def _bool_label(value: bool | None) -> str:
    if value is None:
        return PLACEHOLDER_BLANK
    return "да" if value else "нет"


def _validity_label(valid_from, valid_until) -> str:
    if valid_from and valid_until:
        return f"{_date_label(valid_from)} - {_date_label(valid_until)}"
    if valid_from:
        return f"с {_date_label(valid_from)}"
    if valid_until:
        return f"до {_date_label(valid_until)}"
    return PLACEHOLDER_BLANK


def _template_label(template) -> str:
    if template is None:
        return "Базовый системный шаблон"
    version = f" v{template.version}" if template.version else ""
    return f"{template.title}{version}"


def _snapshot_date(value) -> str:
    if not value:
        return ""
    return value.isoformat()


def _snapshot_datetime(value) -> str:
    if not value:
        return ""
    return timezone.localtime(value).isoformat()


def _snapshot_decimal(value: Decimal | None) -> str:
    if value is None:
        return ""
    return str(value)


def _template_snapshot(template) -> dict[str, object]:
    if template is None:
        return {
            "id": None,
            "title": "Базовый системный шаблон",
            "version": "",
            "template_type": "",
            "template_type_display": "",
            "file_name": "",
            "updated_at": "",
        }
    return {
        "id": template.pk,
        "title": template.title,
        "version": template.version,
        "template_type": template.template_type,
        "template_type_display": template.get_template_type_display(),
        "file_name": template.file.name if template.file else "",
        "updated_at": _snapshot_datetime(template.updated_at),
    }


def _center_snapshot() -> dict[str, object]:
    profile = CenterLegalProfile.get_active()
    if profile is None:
        return {}
    return {
        "id": profile.pk,
        "full_name": profile.full_name,
        "short_name": profile.short_name,
        "director_full_name": profile.director_full_name,
        "director_short_name": profile.director_short_name,
        "director_position": profile.director_position,
        "authority_basis": profile.authority_basis,
        "license_number": profile.license_number,
        "license_date": _snapshot_date(profile.license_date),
        "license_authority": profile.license_authority,
        "ogrn": profile.ogrn,
        "inn": profile.inn,
        "kpp": profile.kpp,
        "legal_address": profile.legal_address,
        "location_address": profile.location_address,
        "phone": profile.phone,
        "email": profile.email,
        "site": profile.site,
        "bank_name": profile.bank_name,
        "bank_bik": profile.bank_bik,
        "bank_account": profile.bank_account,
        "bank_corr_account": profile.bank_corr_account,
        "updated_at": _snapshot_datetime(profile.updated_at),
    }


def _recipient_snapshot(contract: ServiceContract) -> dict[str, object]:
    child = contract.child
    address = child.registration_address or child.residential_address
    return {
        "id": child.pk,
        "full_name": child.full_name,
        "last_name": child.last_name,
        "first_name": child.first_name,
        "middle_name": child.middle_name,
        "birth_date": _snapshot_date(child.birth_date),
        "phone": child.phone,
        "email": child.email,
        "address": address,
        "registration_address": child.registration_address,
        "residential_address": child.residential_address,
        "status": child.status,
        "status_display": child.get_status_display(),
        "updated_at": _snapshot_datetime(child.updated_at),
    }


def _representative_snapshot(contract: ServiceContract) -> dict[str, object]:
    link = contract.representative_link
    representative = link.representative
    return {
        "link_id": link.pk,
        "representative_id": representative.pk,
        "full_name": representative.full_name,
        "last_name": representative.last_name,
        "first_name": representative.first_name,
        "middle_name": representative.middle_name,
        "relationship_type": link.relationship_type,
        "relationship_type_display": link.get_relationship_type_display(),
        "phone": representative.phone,
        "phone_alt": representative.phone_alt,
        "email": representative.email,
        "passport_series": representative.passport_series,
        "passport_number": representative.passport_number,
        "passport_issued_by": representative.passport_issued_by,
        "passport_issued_on": _snapshot_date(representative.passport_issued_on),
        "registration_address": representative.registration_address,
        "is_primary": link.is_primary,
        "signs_contract": link.signs_contract,
        "receives_schedule": link.receives_schedule,
        "is_payer": link.is_payer,
        "representative_updated_at": _snapshot_datetime(representative.updated_at),
        "link_updated_at": _snapshot_datetime(link.updated_at),
    }


def _counterparty_snapshot(
    contract: DonationContract | OrganizationServiceContract,
) -> dict[str, object]:
    counterparty = contract.counterparty
    return {
        "id": counterparty.pk,
        "name": counterparty.name,
        "counterparty_type": counterparty.counterparty_type,
        "counterparty_type_display": counterparty.get_counterparty_type_display(),
        "inn": counterparty.inn,
        "kpp": counterparty.kpp,
        "ogrn": counterparty.ogrn,
        "legal_address": counterparty.legal_address,
        "postal_address": counterparty.postal_address,
        "bank_details": counterparty.bank_details,
        "contact_person": counterparty.contact_person,
        "phone": counterparty.phone,
        "email": counterparty.email,
        "updated_at": _snapshot_datetime(counterparty.updated_at),
    }


def _funding_source_snapshot(contract) -> dict[str, object]:
    funding_source = getattr(contract, "funding_source", None)
    if funding_source is None:
        return {}
    return {
        "id": funding_source.pk,
        "name": funding_source.name,
        "source_type": funding_source.source_type,
        "source_type_display": funding_source.get_source_type_display(),
        "starts_on": _snapshot_date(funding_source.starts_on),
        "ends_on": _snapshot_date(funding_source.ends_on),
        "transfer_policy": funding_source.transfer_policy,
        "transfer_policy_display": funding_source.get_transfer_policy_display(),
        "updated_at": _snapshot_datetime(funding_source.updated_at),
    }


def _line_period_label(line) -> str:
    return _validity_label(line.starts_on, line.ends_on)


def _line_hours_label(line) -> str:
    if line.unit == line.Unit.HOUR:
        return _quantity_label(line.quantity)
    duration_minutes = getattr(line.service, "default_duration_minutes", None)
    if line.unit == line.Unit.SESSION and duration_minutes:
        hours = (line.quantity * Decimal(duration_minutes)) / Decimal("60")
        return _quantity_label(hours.quantize(Decimal("0.01")))
    return PLACEHOLDER_BLANK


def _service_lines(contract: ServiceContract | OrganizationServiceContract):
    return list(contract.service_lines.select_related("service").order_by("sort_order", "pk"))


def _service_line_snapshot(line) -> dict[str, object]:
    return {
        "id": line.pk,
        "service_id": line.service_id,
        "service_code": line.service.code,
        "service_name": line.service_name or line.service.name,
        "directory_service_name": line.service.name,
        "quantity": _snapshot_decimal(line.quantity),
        "unit": line.unit,
        "unit_display": line.get_unit_display(),
        "hours": _line_hours_label(line),
        "unit_price": _snapshot_decimal(line.unit_price),
        "amount": _snapshot_decimal(line.amount),
        "starts_on": _snapshot_date(line.starts_on),
        "ends_on": _snapshot_date(line.ends_on),
        "period": _line_period_label(line),
        "sort_order": line.sort_order,
        "notes": line.notes,
        "updated_at": _snapshot_datetime(line.updated_at),
    }


def _service_lines_snapshot(
    contract: ServiceContract | OrganizationServiceContract,
) -> list[dict[str, object]]:
    return [_service_line_snapshot(line) for line in _service_lines(contract)]


def _certificate_snapshot(contract: ServiceContract) -> dict[str, object]:
    certificate = getattr(contract, "certificate", None)
    if certificate is None:
        return {}
    return {
        "id": certificate.pk,
        "child_id": certificate.child_id,
        "certificate_type": certificate.certificate_type,
        "certificate_type_display": certificate.get_certificate_type_display(),
        "number": certificate.number,
        "total_amount": _snapshot_decimal(certificate.total_amount),
        "remaining_amount": _snapshot_decimal(certificate.remaining_amount),
        "valid_from": _snapshot_date(certificate.valid_from),
        "valid_until": _snapshot_date(certificate.valid_until),
        "note": certificate.note,
        "updated_at": _snapshot_datetime(certificate.updated_at),
    }


def _service_contract_snapshot(contract: ServiceContract, document: Document) -> dict[str, object]:
    service_lines = _service_lines_snapshot(contract)
    total_amount = sum((Decimal(line["amount"] or "0") for line in service_lines), Decimal("0"))
    return {
        "id": contract.pk,
        "document_id": document.pk,
        "number": contract.number,
        "contract_type": contract.contract_type,
        "contract_type_display": contract.get_contract_type_display(),
        "funding_source_id": contract.funding_source_id,
        "certificate_id": contract.certificate_id,
        "certificate": _certificate_snapshot(contract),
        "signed_on": _snapshot_date(contract.signed_on),
        "valid_from": _snapshot_date(contract.valid_from),
        "valid_until": _snapshot_date(contract.valid_until),
        "amount": _snapshot_decimal(total_amount),
        "service_lines": service_lines,
        "status": contract.status,
        "status_display": contract.get_status_display(),
        "updated_at": _snapshot_datetime(contract.updated_at),
    }


def _donation_contract_snapshot(contract: DonationContract, document: Document) -> dict[str, object]:
    return {
        "id": contract.pk,
        "document_id": document.pk,
        "number": contract.number,
        "contract_type": contract.contract_type,
        "contract_type_display": contract.get_contract_type_display(),
        "signed_on": _snapshot_date(contract.signed_on),
        "valid_from": _snapshot_date(contract.valid_from),
        "valid_until": _snapshot_date(contract.valid_until),
        "amount_limit": _snapshot_decimal(contract.amount_limit),
        "status": contract.status,
        "status_display": contract.get_status_display(),
        "updated_at": _snapshot_datetime(contract.updated_at),
    }


def _organization_service_contract_snapshot(
    contract: OrganizationServiceContract,
    document: Document,
) -> dict[str, object]:
    service_lines = _service_lines_snapshot(contract)
    total_amount = sum((Decimal(line["amount"] or "0") for line in service_lines), Decimal("0"))
    return {
        "id": contract.pk,
        "document_id": document.pk,
        "number": contract.number,
        "contract_type": contract.contract_type,
        "contract_type_display": contract.get_contract_type_display(),
        "funding_source_id": contract.funding_source_id,
        "signed_on": _snapshot_date(contract.signed_on),
        "valid_from": _snapshot_date(contract.valid_from),
        "valid_until": _snapshot_date(contract.valid_until),
        "amount": _snapshot_decimal(total_amount),
        "service_lines": service_lines,
        "status": contract.status,
        "status_display": contract.get_status_display(),
        "updated_at": _snapshot_datetime(contract.updated_at),
    }


def _snapshot_actor(actor):
    if actor is not None and getattr(actor, "is_authenticated", False):
        return actor
    return None


def _document_snapshot(document: Document) -> ContractLegalSnapshot | None:
    return getattr(document, "contract_legal_snapshot", None)


def _ensure_service_snapshot_owner(contract: ServiceContract, document: Document) -> None:
    snapshot = _document_snapshot(document)
    if snapshot is None:
        return
    if (
        snapshot.contract_kind != ContractLegalSnapshot.ContractKind.SERVICE
        or snapshot.service_contract_id != contract.pk
    ):
        raise ContractDocumentError(
            "Связанный документ уже имеет юридический snapshot другого договора."
        )


def _ensure_donation_snapshot_owner(contract: DonationContract, document: Document) -> None:
    snapshot = _document_snapshot(document)
    if snapshot is None:
        return
    if (
        snapshot.contract_kind != ContractLegalSnapshot.ContractKind.DONATION
        or snapshot.donation_contract_id != contract.pk
    ):
        raise ContractDocumentError(
            "Связанный документ уже имеет юридический snapshot другого договора."
        )


def _ensure_organization_snapshot_owner(
    contract: OrganizationServiceContract,
    document: Document,
) -> None:
    snapshot = _document_snapshot(document)
    if snapshot is None:
        return
    if (
        snapshot.contract_kind != ContractLegalSnapshot.ContractKind.ORGANIZATION_SERVICE
        or snapshot.organization_contract_id != contract.pk
    ):
        raise ContractDocumentError(
            "Связанный документ уже имеет юридический snapshot другого договора."
        )


def _save_service_legal_snapshot(
    contract: ServiceContract,
    document: Document,
    *,
    actor=None,
) -> ContractLegalSnapshot:
    snapshot = _document_snapshot(document)
    if snapshot is None:
        snapshot = ContractLegalSnapshot(document=document)
    else:
        _ensure_service_snapshot_owner(contract, document)
    snapshot.contract_kind = ContractLegalSnapshot.ContractKind.SERVICE
    snapshot.service_contract = contract
    snapshot.donation_contract = None
    snapshot.generated_by = _snapshot_actor(actor)
    snapshot.contract_snapshot = _service_contract_snapshot(contract, document)
    snapshot.center_snapshot = _center_snapshot()
    snapshot.recipient_snapshot = _recipient_snapshot(contract)
    snapshot.representative_snapshot = _representative_snapshot(contract)
    snapshot.counterparty_snapshot = {}
    snapshot.funding_source_snapshot = _funding_source_snapshot(contract)
    snapshot.template_snapshot = _template_snapshot(contract.template)
    snapshot.note = "Сформирован автоматически при генерации Word-файла договора с получателем."
    snapshot.full_clean()
    snapshot.save()
    return snapshot


def _save_donation_legal_snapshot(
    contract: DonationContract,
    document: Document,
    *,
    actor=None,
) -> ContractLegalSnapshot:
    snapshot = _document_snapshot(document)
    if snapshot is None:
        snapshot = ContractLegalSnapshot(document=document)
    else:
        _ensure_donation_snapshot_owner(contract, document)
    snapshot.contract_kind = ContractLegalSnapshot.ContractKind.DONATION
    snapshot.service_contract = None
    snapshot.donation_contract = contract
    snapshot.generated_by = _snapshot_actor(actor)
    snapshot.contract_snapshot = _donation_contract_snapshot(contract, document)
    snapshot.center_snapshot = _center_snapshot()
    snapshot.recipient_snapshot = {}
    snapshot.representative_snapshot = {}
    snapshot.counterparty_snapshot = _counterparty_snapshot(contract)
    snapshot.funding_source_snapshot = _funding_source_snapshot(contract)
    snapshot.template_snapshot = _template_snapshot(contract.template)
    snapshot.note = "Сформирован автоматически при генерации Word-файла договора пожертвования."
    snapshot.full_clean()
    snapshot.save()
    return snapshot


def _save_organization_legal_snapshot(
    contract: OrganizationServiceContract,
    document: Document,
    *,
    actor=None,
) -> ContractLegalSnapshot:
    snapshot = _document_snapshot(document)
    if snapshot is None:
        snapshot = ContractLegalSnapshot(document=document)
    else:
        _ensure_organization_snapshot_owner(contract, document)
    snapshot.contract_kind = ContractLegalSnapshot.ContractKind.ORGANIZATION_SERVICE
    snapshot.service_contract = None
    snapshot.donation_contract = None
    snapshot.organization_contract = contract
    snapshot.generated_by = _snapshot_actor(actor)
    snapshot.contract_snapshot = _organization_service_contract_snapshot(contract, document)
    snapshot.center_snapshot = _center_snapshot()
    snapshot.recipient_snapshot = {}
    snapshot.representative_snapshot = {}
    snapshot.counterparty_snapshot = _counterparty_snapshot(contract)
    snapshot.funding_source_snapshot = _funding_source_snapshot(contract)
    snapshot.template_snapshot = _template_snapshot(contract.template)
    snapshot.note = "Сформирован автоматически при генерации Word-файла B2B-договора услуг."
    snapshot.full_clean()
    snapshot.save()
    return snapshot


def _file_basename(name: str) -> str:
    return name.rsplit("/", 1)[-1].rsplit("\\", 1)[-1]


def _read_document_file(document: Document) -> bytes:
    if not document.file:
        raise ContractDocumentError("У связанного документа нет файла для архива.")
    try:
        document.file.open("rb")
        try:
            return document.file.read()
        finally:
            document.file.close()
    except OSError as exc:
        raise ContractDocumentError("Не удалось прочитать файл договора для архива.") from exc


def _signed_file_common_fields(snapshot: ContractLegalSnapshot) -> dict[str, object]:
    return {
        "contract_snapshot": deepcopy(snapshot.contract_snapshot),
        "center_snapshot": deepcopy(snapshot.center_snapshot),
        "recipient_snapshot": deepcopy(snapshot.recipient_snapshot),
        "representative_snapshot": deepcopy(snapshot.representative_snapshot),
        "counterparty_snapshot": deepcopy(snapshot.counterparty_snapshot),
        "funding_source_snapshot": deepcopy(snapshot.funding_source_snapshot),
        "template_snapshot": deepcopy(snapshot.template_snapshot),
    }


def _archive_contract_signed_file(
    *,
    contract_kind: str,
    document: Document,
    snapshot: ContractLegalSnapshot,
    actor=None,
    service_contract: ServiceContract | None = None,
    donation_contract: DonationContract | None = None,
    organization_contract: OrganizationServiceContract | None = None,
) -> ContractSignedFile:
    payload = _read_document_file(document)
    if not payload:
        raise ContractDocumentError("Нельзя архивировать пустой файл договора.")
    original_filename = _file_basename(document.file.name)
    signed_file = ContractSignedFile(
        contract_kind=contract_kind,
        service_contract=service_contract,
        donation_contract=donation_contract,
        organization_contract=organization_contract,
        source_document=document,
        original_filename=original_filename,
        content_type=mimetypes.guess_type(original_filename)[0] or "",
        file_size=len(payload),
        file_sha256=hashlib.sha256(payload).hexdigest(),
        signed_on=(
            service_contract.signed_on
            if service_contract is not None and service_contract.signed_on
            else donation_contract.signed_on
            if donation_contract is not None and donation_contract.signed_on
            else organization_contract.signed_on
            if organization_contract is not None and organization_contract.signed_on
            else timezone.localdate()
        ),
        uploaded_by=_snapshot_actor(actor),
        note="Архивная копия подписанного файла договора из текущего связанного документа.",
        **_signed_file_common_fields(snapshot),
    )
    signed_file.file.save(original_filename, ContentFile(payload), save=False)
    signed_file.save()
    return signed_file


def archive_service_contract_signed_file(
    contract: ServiceContract,
    *,
    actor=None,
) -> ContractSignedFile:
    document = contract.document
    if document is None:
        raise ContractDocumentError("Сначала сформируйте Word-файл договора с получателем.")
    _ensure_service_snapshot_owner(contract, document)
    snapshot = _document_snapshot(document)
    if snapshot is None:
        raise ContractDocumentError(
            "Сначала сформируйте Word-файл, чтобы зафиксировать юридический snapshot."
        )
    return _archive_contract_signed_file(
        contract_kind=ContractSignedFile.ContractKind.SERVICE,
        document=document,
        snapshot=snapshot,
        actor=actor,
        service_contract=contract,
    )


def archive_donation_contract_signed_file(
    contract: DonationContract,
    *,
    actor=None,
) -> ContractSignedFile:
    document = contract.document
    if document is None:
        raise ContractDocumentError("Сначала сформируйте Word-файл договора пожертвования.")
    _ensure_donation_snapshot_owner(contract, document)
    snapshot = _document_snapshot(document)
    if snapshot is None:
        raise ContractDocumentError(
            "Сначала сформируйте Word-файл, чтобы зафиксировать юридический snapshot."
        )
    return _archive_contract_signed_file(
        contract_kind=ContractSignedFile.ContractKind.DONATION,
        document=document,
        snapshot=snapshot,
        actor=actor,
        donation_contract=contract,
    )


def archive_organization_service_contract_signed_file(
    contract: OrganizationServiceContract,
    *,
    actor=None,
) -> ContractSignedFile:
    document = contract.document
    if document is None:
        raise ContractDocumentError("Сначала сформируйте Word-файл B2B-договора услуг.")
    _ensure_organization_snapshot_owner(contract, document)
    snapshot = _document_snapshot(document)
    if snapshot is None:
        raise ContractDocumentError(
            "Сначала сформируйте Word-файл, чтобы зафиксировать юридический snapshot."
        )
    return _archive_contract_signed_file(
        contract_kind=ContractSignedFile.ContractKind.ORGANIZATION_SERVICE,
        document=document,
        snapshot=snapshot,
        actor=actor,
        organization_contract=contract,
    )


def _center_placeholder_values() -> dict[str, str]:
    profile = CenterLegalProfile.get_active()
    if profile is None:
        return {}
    return {
        "center.full_name": _text(profile.full_name),
        "center.short_name": _text(profile.short_name, profile.full_name),
        "center.director_full_name": _text(profile.director_full_name),
        "center.director_short_name": _text(profile.director_short_name),
        "center.director_position": _text(profile.director_position),
        "center.authority_basis": _text(profile.authority_basis),
        "center.license_number": _text(profile.license_number),
        "center.license_date": _date_label(profile.license_date),
        "center.license_authority": _text(profile.license_authority),
        "center.ogrn": _text(profile.ogrn),
        "center.inn": _text(profile.inn),
        "center.kpp": _text(profile.kpp),
        "center.legal_address": _text(profile.legal_address),
        "center.location_address": _text(profile.location_address),
        "center.phone": _text(profile.phone),
        "center.email": _text(profile.email),
        "center.site": _text(profile.site),
        "center.bank_name": _text(profile.bank_name),
        "center.bank_bik": _text(profile.bank_bik),
        "center.bank_account": _text(profile.bank_account),
        "center.bank_corr_account": _text(profile.bank_corr_account),
    }


def _funding_source_placeholder_values(contract) -> dict[str, str]:
    funding_source = getattr(contract, "funding_source", None)
    if funding_source is None:
        return {}
    return {
        "funding_source.name": funding_source.name,
        "funding_source.type": funding_source.get_source_type_display(),
        "funding_source.starts_on": _date_label(funding_source.starts_on),
        "funding_source.ends_on": _date_label(funding_source.ends_on),
        "funding_source.transfer_policy": funding_source.get_transfer_policy_display(),
        "funding_source.project_name": _text(getattr(funding_source, "project_name", "")),
    }


def _service_line_text(line) -> str:
    return (
        f"{line.service_name or line.service.name}: "
        f"{_quantity_label(line.quantity)} {line.get_unit_display()} x "
        f"{_money_label(line.unit_price)} = {_money_label(line.amount)}"
    )


def _service_spec_placeholder_values(
    contract: ServiceContract | OrganizationServiceContract,
) -> dict[str, str]:
    lines = _service_lines(contract)
    total_amount = sum((line.amount for line in lines), Decimal("0"))
    values = {
        "contract.amount": _money_label(total_amount) if lines else PLACEHOLDER_BLANK,
        "service_spec.rows": "\n".join(_service_line_text(line) for line in lines)
        if lines
        else PLACEHOLDER_BLANK,
    }
    if not lines:
        return values
    first = lines[0]
    values.update(
        {
            "service_spec.service_name": _text(first.service_name or first.service.name),
            "service_spec.quantity": _quantity_label(first.quantity),
            "service_spec.unit": first.get_unit_display(),
            "service_spec.hours": _line_hours_label(first),
            "service_spec.price": _money_label(first.unit_price),
            "service_spec.amount": _money_label(first.amount),
            "service_spec.period": _line_period_label(first),
        }
    )
    return values


def _certificate_placeholder_values(contract: ServiceContract) -> dict[str, str]:
    certificate = getattr(contract, "certificate", None)
    if certificate is None:
        return {}
    return {
        "certificate.type": certificate.get_certificate_type_display(),
        "certificate.number": _text(certificate.number),
        "certificate.total_amount": _money_label(certificate.total_amount),
        "certificate.remaining_amount": _money_label(certificate.remaining_amount),
        "certificate.valid_from": _date_label(certificate.valid_from),
        "certificate.valid_until": _date_label(certificate.valid_until),
        "certificate.payer_name": PLACEHOLDER_BLANK,
    }


def service_contract_placeholders(contract: ServiceContract) -> dict[str, str]:
    signer_link = contract.representative_link
    signer = signer_link.representative
    child_address = contract.child.registration_address or contract.child.residential_address
    values = _empty_placeholder_values()
    values.update(_center_placeholder_values())
    values.update(_funding_source_placeholder_values(contract))
    values.update(_service_spec_placeholder_values(contract))
    values.update(_certificate_placeholder_values(contract))
    values.update(
        {
            "contract.number": _text(contract.number, "б/н"),
            "contract.signed_on": _date_label(contract.signed_on),
            "contract.valid_from": _date_label(contract.valid_from),
            "contract.valid_until": _date_label(contract.valid_until),
            "contract.validity": _validity_label(contract.valid_from, contract.valid_until),
            "contract.status": contract.get_status_display(),
            "contract.type": contract.get_contract_type_display(),
            "contract.template": _template_label(contract.template),
            "child.full_name": contract.child.full_name,
            "child.birth_date": _date_label(contract.child.birth_date),
            "child.phone": _text(contract.child.phone),
            "child.email": _text(contract.child.email),
            "child.address": _text(child_address),
            "representative.full_name": signer.full_name,
            "representative.relationship": signer_link.get_relationship_type_display(),
            "representative.phone": _text(signer.phone),
            "representative.phone_alt": _text(signer.phone_alt),
            "representative.email": _text(signer.email),
            "representative.passport_series": _text(signer.passport_series),
            "representative.passport_number": _text(signer.passport_number),
            "representative.passport_issued_by": _text(signer.passport_issued_by),
            "representative.passport_issued_on": _date_label(signer.passport_issued_on),
            "representative.registration_address": _text(signer.registration_address),
            "representative.signs_contract": _bool_label(signer_link.signs_contract),
            "representative.receives_schedule": _bool_label(signer_link.receives_schedule),
            "representative.is_payer": _bool_label(signer_link.is_payer),
        }
    )
    return values


def donation_contract_placeholders(contract: DonationContract) -> dict[str, str]:
    counterparty = contract.counterparty
    values = _empty_placeholder_values()
    values.update(_center_placeholder_values())
    values.update(_funding_source_placeholder_values(contract))
    values.update(
        {
            "contract.number": _text(contract.number, "б/н"),
            "contract.signed_on": _date_label(contract.signed_on),
            "contract.valid_from": _date_label(contract.valid_from),
            "contract.valid_until": _date_label(contract.valid_until),
            "contract.validity": _validity_label(contract.valid_from, contract.valid_until),
            "contract.status": contract.get_status_display(),
            "contract.type": contract.get_contract_type_display(),
            "contract.template": _template_label(contract.template),
            "counterparty.name": counterparty.name,
            "counterparty.type": counterparty.get_counterparty_type_display(),
            "counterparty.inn": _text(counterparty.inn, ""),
            "counterparty.kpp": _text(counterparty.kpp, ""),
            "counterparty.ogrn": _text(counterparty.ogrn, ""),
            "counterparty.legal_address": _text(counterparty.legal_address, ""),
            "counterparty.postal_address": _text(counterparty.postal_address, ""),
            "counterparty.bank_details": _text(counterparty.bank_details, ""),
            "counterparty.contact_person": _text(counterparty.contact_person),
            "counterparty.phone": _text(counterparty.phone),
            "counterparty.email": _text(counterparty.email),
            "donation.amount_limit": _money_label(contract.amount_limit),
            "donation.amount": _money_label(contract.amount_limit),
        }
    )
    return values


def organization_service_contract_placeholders(
    contract: OrganizationServiceContract,
) -> dict[str, str]:
    counterparty = contract.counterparty
    values = _empty_placeholder_values()
    values.update(_center_placeholder_values())
    values.update(_funding_source_placeholder_values(contract))
    values.update(_service_spec_placeholder_values(contract))
    values.update(
        {
            "contract.number": _text(contract.number, "б/н"),
            "contract.signed_on": _date_label(contract.signed_on),
            "contract.valid_from": _date_label(contract.valid_from),
            "contract.valid_until": _date_label(contract.valid_until),
            "contract.validity": _validity_label(contract.valid_from, contract.valid_until),
            "contract.status": contract.get_status_display(),
            "contract.type": contract.get_contract_type_display(),
            "contract.template": _template_label(contract.template),
            "counterparty.name": counterparty.name,
            "counterparty.type": counterparty.get_counterparty_type_display(),
            "counterparty.inn": _text(counterparty.inn, ""),
            "counterparty.kpp": _text(counterparty.kpp, ""),
            "counterparty.ogrn": _text(counterparty.ogrn, ""),
            "counterparty.legal_address": _text(counterparty.legal_address, ""),
            "counterparty.postal_address": _text(counterparty.postal_address, ""),
            "counterparty.bank_details": _text(counterparty.bank_details, ""),
            "counterparty.contact_person": _text(counterparty.contact_person),
            "counterparty.phone": _text(counterparty.phone),
            "counterparty.email": _text(counterparty.email),
        }
    )
    return values


def _replace_placeholders(text: str, values: dict[str, str]) -> str:
    result = text
    for key, value in values.items():
        result = result.replace("{{ " + key + " }}", value)
        result = result.replace("{{" + key + "}}", value)
    return result


def _replace_in_paragraph(paragraph, values: dict[str, str]) -> None:
    current = paragraph.text
    updated = _replace_placeholders(current, values)
    if updated == current:
        return
    if not paragraph.runs:
        paragraph.add_run(updated)
        return
    paragraph.runs[0].text = updated
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_in_table(table, values: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, values)
            for nested_table in cell.tables:
                _replace_in_table(nested_table, values)


def _replace_in_document(document: WordDocumentType, values: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, values)
    for table in document.tables:
        _replace_in_table(table, values)
    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                _replace_in_paragraph(paragraph, values)
            for table in part.tables:
                _replace_in_table(table, values)


def _load_template_document(contract) -> WordDocumentType | None:
    template = contract.template
    if template is None or not template.file:
        return None
    if template.file.size > MAX_TEMPLATE_BYTES:
        raise ContractDocumentError("Файл шаблона больше 5 МБ. Загрузите более компактный .docx.")
    try:
        template.file.open("rb")
        try:
            return WordDocument(template.file)
        finally:
            template.file.close()
    except Exception as exc:
        raise ContractDocumentError("Не удалось прочитать .docx шаблон договора.") from exc


def _service_fallback_document(contract: ServiceContract) -> WordDocumentType:
    values = service_contract_placeholders(contract)
    document = WordDocument()
    document.add_heading("Договор на оказание реабилитационных услуг", level=1)
    document.add_paragraph(f"№ {values['contract.number']} от {values['contract.signed_on']}")
    document.add_paragraph(f"Шаблон: {values['contract.template']}")

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in [
        ("Получатель", values["child.full_name"]),
        ("Дата рождения", values["child.birth_date"]),
        ("Подписант", values["representative.full_name"]),
        ("Тип договора", values["contract.type"]),
        ("Статус реестра", values["contract.status"]),
        ("Срок действия", f"{values['contract.valid_from']} - {values['contract.valid_until']}"),
    ]:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value

    document.add_paragraph(
        "Стоимость услуг определяется действующим прейскурантом центра. "
        "Расписание, переносы и отмены ведутся в операционной системе центра."
    )
    document.add_paragraph("Подписи сторон:")
    document.add_paragraph("Представитель: ______________________ /_______________/")
    document.add_paragraph("Центр: ______________________________ /_______________/")
    return document


def _donation_fallback_document(contract: DonationContract) -> WordDocumentType:
    values = donation_contract_placeholders(contract)
    document = WordDocument()
    document.add_heading("Договор пожертвования", level=1)
    document.add_paragraph(f"№ {values['contract.number']} от {values['contract.signed_on']}")
    document.add_paragraph(f"Шаблон: {values['contract.template']}")

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in [
        ("Жертвователь/спонсор", values["counterparty.name"]),
        ("ИНН", values["counterparty.inn"]),
        ("Источник финансирования", values["funding_source.name"]),
        ("Тип договора", values["contract.type"]),
        ("Статус реестра", values["contract.status"]),
        ("Лимит суммы", values["donation.amount_limit"]),
        ("Срок действия", f"{values['contract.valid_from']} - {values['contract.valid_until']}"),
    ]:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value

    document.add_paragraph(
        "Этот документ не создает автоматические платежи, балансы получателей, "
        "ledger-записи или начисления специалистам."
    )
    document.add_paragraph("Подписи сторон:")
    document.add_paragraph("Жертвователь/спонсор: ______________ /_______________/")
    document.add_paragraph("Центр: _____________________________ /_______________/")
    return document


def _organization_service_fallback_document(contract: OrganizationServiceContract) -> WordDocumentType:
    values = organization_service_contract_placeholders(contract)
    document = WordDocument()
    document.add_heading("B2B-договор на оказание реабилитационных услуг", level=1)
    document.add_paragraph(f"№ {values['contract.number']} от {values['contract.signed_on']}")
    document.add_paragraph(f"Шаблон: {values['contract.template']}")

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in [
        ("Организация", values["counterparty.name"]),
        ("ИНН", values["counterparty.inn"]),
        ("Источник финансирования", values["funding_source.name"]),
        ("Тип договора", values["contract.type"]),
        ("Статус реестра", values["contract.status"]),
        ("Срок действия", values["contract.validity"]),
        ("Сумма договора", values["contract.amount"]),
        ("Спецификация услуг", values["service_spec.rows"]),
    ]:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value

    document.add_paragraph(
        "Этот документ фиксирует юридическое основание оказания услуг организации. "
        "Он не создает автоматические платежи, списания по балансам, занятия или акты."
    )
    document.add_paragraph("Подписи сторон:")
    document.add_paragraph("Организация: ______________________ /_______________/")
    document.add_paragraph("Центр: _____________________________ /_______________/")
    return document


def _render_document(document: WordDocumentType) -> BytesIO:
    payload = BytesIO()
    document.save(payload)
    payload.seek(0)
    return payload


def render_service_contract_docx(contract: ServiceContract) -> GeneratedContractFile:
    document = _load_template_document(contract)
    if document is None:
        document = _service_fallback_document(contract)
    else:
        _replace_in_document(document, service_contract_placeholders(contract))
    return GeneratedContractFile(
        payload=_render_document(document),
        filename=_safe_filename("service_contract", contract),
    )


def render_donation_contract_docx(contract: DonationContract) -> GeneratedContractFile:
    document = _load_template_document(contract)
    if document is None:
        document = _donation_fallback_document(contract)
    else:
        _replace_in_document(document, donation_contract_placeholders(contract))
    return GeneratedContractFile(
        payload=_render_document(document),
        filename=_safe_filename("donation_contract", contract),
    )


def render_organization_service_contract_docx(
    contract: OrganizationServiceContract,
) -> GeneratedContractFile:
    document = _load_template_document(contract)
    if document is None:
        document = _organization_service_fallback_document(contract)
    else:
        _replace_in_document(document, organization_service_contract_placeholders(contract))
    return GeneratedContractFile(
        payload=_render_document(document),
        filename=_safe_filename("organization_service_contract", contract),
    )


def save_service_contract_docx(contract: ServiceContract, *, actor=None) -> GeneratedContractFile:
    generated = render_service_contract_docx(contract)
    document = contract.document
    if document is not None:
        if document.target_type != Document.TargetType.RECIPIENT:
            raise ContractDocumentError(
                "Связанный документ не является документом получателя. Исправьте карточку договора."
            )
        if document.child_id != contract.child_id:
            raise ContractDocumentError(
                "Связанный документ относится к другому получателю. Исправьте карточку договора."
            )
        if document.category != Document.Category.CONTRACT:
            raise ContractDocumentError(
                "Связанный документ должен быть категорией договора. Исправьте карточку договора."
            )
        _ensure_service_snapshot_owner(contract, document)
    else:
        document = Document(
            target_type=Document.TargetType.RECIPIENT,
            child=contract.child,
            category=Document.Category.CONTRACT,
        )

    document.title = _service_document_title(contract)
    document.issued_on = contract.signed_on or timezone.localdate()
    document.expires_on = contract.valid_until
    if actor is not None and getattr(actor, "is_authenticated", False):
        document.uploaded_by = actor
    document.note = "Сформирован автоматически из карточки договора и Word-шаблона."
    document.file.save(generated.filename, ContentFile(generated.payload.getvalue()), save=False)
    document.full_clean()
    document.save()

    if contract.document_id != document.pk:
        contract.document = document
        contract.save(update_fields=["document", "updated_at"])

    _save_service_legal_snapshot(contract, document, actor=actor)

    generated.payload.seek(0)
    return GeneratedContractFile(
        payload=generated.payload,
        filename=generated.filename,
        document=document,
    )


def save_donation_contract_docx(contract: DonationContract, *, actor=None) -> GeneratedContractFile:
    generated = render_donation_contract_docx(contract)
    document = contract.document
    if document is not None:
        if document.category != Document.Category.CONTRACT:
            raise ContractDocumentError(
                "Связанный документ должен быть категорией договора. Исправьте карточку договора."
            )
        if document.target_type == Document.TargetType.RECIPIENT:
            raise ContractDocumentError(
                "Договор пожертвования нельзя сохранять в документ получателя."
            )
        if (
            document.target_type == Document.TargetType.COUNTERPARTY
            and document.counterparty_id != contract.counterparty_id
        ):
            raise ContractDocumentError(
                "Связанный документ относится к другому контрагенту. Исправьте карточку договора."
            )
        _ensure_donation_snapshot_owner(contract, document)
    else:
        document = Document(
            target_type=Document.TargetType.COUNTERPARTY,
            counterparty=contract.counterparty,
            category=Document.Category.CONTRACT,
        )

    document.title = _donation_document_title(contract)
    document.issued_on = contract.signed_on or timezone.localdate()
    document.expires_on = contract.valid_until
    if actor is not None and getattr(actor, "is_authenticated", False):
        document.uploaded_by = actor
    document.note = "Сформирован автоматически из карточки договора пожертвования и Word-шаблона."
    document.file.save(generated.filename, ContentFile(generated.payload.getvalue()), save=False)
    document.full_clean()
    document.save()

    if contract.document_id != document.pk:
        contract.document = document
        contract.save(update_fields=["document", "updated_at"])

    _save_donation_legal_snapshot(contract, document, actor=actor)

    generated.payload.seek(0)
    return GeneratedContractFile(
        payload=generated.payload,
        filename=generated.filename,
        document=document,
    )


def save_organization_service_contract_docx(
    contract: OrganizationServiceContract,
    *,
    actor=None,
) -> GeneratedContractFile:
    generated = render_organization_service_contract_docx(contract)
    document = contract.document
    if document is not None:
        if document.category != Document.Category.CONTRACT:
            raise ContractDocumentError(
                "Связанный документ должен быть категорией договора. Исправьте карточку B2B-договора."
            )
        if document.target_type == Document.TargetType.RECIPIENT:
            raise ContractDocumentError(
                "B2B-договор нельзя сохранять в документ получателя."
            )
        if (
            document.target_type == Document.TargetType.COUNTERPARTY
            and document.counterparty_id != contract.counterparty_id
        ):
            raise ContractDocumentError(
                "Связанный документ относится к другой организации. Исправьте карточку B2B-договора."
            )
        _ensure_organization_snapshot_owner(contract, document)
    else:
        document = Document(
            target_type=Document.TargetType.COUNTERPARTY,
            counterparty=contract.counterparty,
            category=Document.Category.CONTRACT,
        )

    document.title = _organization_service_document_title(contract)
    document.issued_on = contract.signed_on or timezone.localdate()
    document.expires_on = contract.valid_until
    if actor is not None and getattr(actor, "is_authenticated", False):
        document.uploaded_by = actor
    document.note = "Сформирован автоматически из карточки B2B-договора услуг и Word-шаблона."
    document.file.save(generated.filename, ContentFile(generated.payload.getvalue()), save=False)
    document.full_clean()
    document.save()

    if contract.document_id != document.pk:
        contract.document = document
        contract.save(update_fields=["document", "updated_at"])

    _save_organization_legal_snapshot(contract, document, actor=actor)

    generated.payload.seek(0)
    return GeneratedContractFile(
        payload=generated.payload,
        filename=generated.filename,
        document=document,
    )


def _service_document_title(contract: ServiceContract) -> str:
    number = contract.number or "б/н"
    title = f"Договор {number} — {contract.child.full_name}"
    return title[:200]


def _donation_document_title(contract: DonationContract) -> str:
    number = contract.number or "б/н"
    title = f"Договор пожертвования {number} — {contract.counterparty.name}"
    return title[:200]


def _organization_service_document_title(contract: OrganizationServiceContract) -> str:
    number = contract.number or "б/н"
    title = f"B2B-договор услуг {number} — {contract.counterparty.name}"
    return title[:200]

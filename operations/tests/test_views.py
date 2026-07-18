"""Тесты новых views: tomorrow, timesheet, grant_report, mass_reschedule, recommendations, documents, consents, payments."""

from __future__ import annotations

from datetime import datetime, time, timedelta
from decimal import Decimal
from io import BytesIO

from django.contrib.auth.models import User
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from django.utils import timezone

from operations.forms import AppointmentConfirmationSendForm
from operations.models import (
    Appointment,
    AppointmentConfirmation,
    AppointmentParticipant,
    AppointmentRescheduleChain,
    AppointmentReschedulePlan,
    AppointmentRescheduleStep,
    AppointmentRoomOverride,
    AppointmentStaffAssignment,
    BalanceAccount,
    CenterExpense,
    CenterExpenseCategory,
    CenterLegalProfile,
    Certificate,
    Child,
    Consent,
    ContractLegalSnapshot,
    ContractTemplate,
    Counterparty,
    Discount,
    Document,
    DonationContract,
    EquipmentAsset,
    ExpenseFundingSplit,
    FinancialIntegrityCheckRun,
    FinancialIntegrityFinding,
    FinancialIntegrityFindingEvent,
    FundingServiceQuota,
    FundingSource,
    FundingStaffAllocation,
    GrantRecipientAllocation,
    LedgerEntry,
    ParentGuardian,
    Payment,
    PayrollAccrual,
    PayrollSheet,
    ProgramBlock,
    RecipientRepresentative,
    Recommendation,
    Room,
    Service,
    ServiceContract,
    StaffCompensationRule,
    StaffMember,
    TreatmentProgram,
)
from operations.services import (
    billing as billing_svc,
    financial_integrity_checks as financial_integrity_checks_svc,
    rescheduling_plans as plan_svc,
)


def _local_dt(day, clock):
    return timezone.make_aware(datetime.combine(day, clock), timezone.get_current_timezone())


class NewViewsTestBase(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.admin = User.objects.create_user(
            "admin", password="x", is_staff=True, is_superuser=True
        )
        cls.parent = ParentGuardian.objects.create(
            last_name="Иванова",
            first_name="Мария",
            phone="+7 900 000-00-01",
            email="rep@example.local",
        )
        cls.child = Child.objects.create(
            last_name="Иванов", first_name="Ваня", primary_parent=cls.parent
        )
        cls.staff_user = User.objects.create_user("specialist1", password="x")
        cls.staff = StaffMember.objects.create(
            user=cls.staff_user, full_name="Иванова Н. Г.", specializations="Логопед"
        )
        cls.service = Service.objects.create(
            name="Логопед",
            code="LOG",
            category=Service.Category.SPEECH,
            default_duration_minutes=30,
            default_price=Decimal("1500"),
        )
        cls.service2 = Service.objects.create(
            name="Дефектолог",
            code="DEF",
            category=Service.Category.DEFECTOLOGY,
            default_duration_minutes=45,
            default_price=Decimal("1800"),
        )
        cls.room = Room.objects.create(name="Кабинет 1")
        cls.funding = FundingSource.objects.create(
            name="Личные средства",
            source_type=FundingSource.SourceType.PERSONAL,
        )
        cls.funding_grant = FundingSource.objects.create(
            name="Грант",
            source_type=FundingSource.SourceType.GRANT,
        )
        cls.expense_category = CenterExpenseCategory.objects.create(
            name="Хозяйственные расходы",
            expense_type=CenterExpenseCategory.ExpenseType.HOUSEHOLD,
        )
        cls.equipment_expense_category = CenterExpenseCategory.objects.create(
            name="Оборудование",
            expense_type=CenterExpenseCategory.ExpenseType.EQUIPMENT,
        )
        cls.counterparty = Counterparty.objects.create(
            name="Поставщик материалов",
            counterparty_type=Counterparty.CounterpartyType.VENDOR,
        )
        cls.account = BalanceAccount.objects.create(
            child=cls.child,
            funding_source=cls.funding,
            unit=BalanceAccount.Unit.SESSIONS,
            service_scope=BalanceAccount.ServiceScope.SPECIFIC_SERVICE,
            service=cls.service,
            initial_amount=Decimal("10"),
        )
        cls.account_grant = BalanceAccount.objects.create(
            child=cls.child,
            funding_source=cls.funding_grant,
            unit=BalanceAccount.Unit.SESSIONS,
            service_scope=BalanceAccount.ServiceScope.SPECIFIC_SERVICE,
            service=cls.service2,
            initial_amount=Decimal("20"),
        )

    def setUp(self):
        self.client.force_login(self.admin)


class TomorrowViewTests(NewViewsTestBase):
    def test_tomorrow_renders_with_default_date(self):
        response = self.client.get(reverse("tomorrow"))
        self.assertEqual(response.status_code, 200)
        self.assertIn("overview", response.context)
        self.assertIn("tomorrow_summary_items", response.context)
        self.assertIn("tomorrow_control_items", response.context)
        self.assertIn("tomorrow_next_action", response.context)
        self.assertContains(response, "Следующий шаг")
        self.assertContains(response, "Контроль дня")
        self.assertContains(response, 'id="tomorrow-appointments"')

    def test_tomorrow_accepts_date_param(self):
        target = timezone.localdate() + timedelta(days=3)
        response = self.client.get(reverse("tomorrow"), {"date": target.isoformat()})
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context["overview"].date, target)

    def test_tomorrow_rejects_specialist(self):
        self.client.logout()
        self.client.force_login(self.staff_user)
        response = self.client.get(reverse("tomorrow"))
        self.assertEqual(response.status_code, 302)

    def test_tomorrow_shows_group_participants_and_staff(self):
        day = timezone.localdate() + timedelta(days=1)
        start = _local_dt(day, time(10, 0))
        second_parent = ParentGuardian.objects.create(
            last_name="Петрова",
            first_name="Анна",
            phone="+7 900 000-00-02",
        )
        second_child = Child.objects.create(
            last_name="Петров",
            first_name="Илья",
            primary_parent=second_parent,
        )
        second_staff = StaffMember.objects.create(full_name="Петрова И. А.")
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            session_type=Appointment.SessionType.GROUP,
            title="Группа общения",
        )
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=second_child,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )
        AppointmentStaffAssignment.objects.create(
            appointment=appointment,
            staff_member=second_staff,
            role=AppointmentStaffAssignment.Role.ASSISTANT,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )

        response = self.client.get(reverse("tomorrow"), {"date": day.isoformat()})

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Группа общения")
        self.assertContains(response, "Петров Илья")
        self.assertContains(response, "Петрова И. А.")
        self.assertContains(response, "tomorrow-table")
        self.assertContains(response, 'data-label="Время"')
        self.assertContains(response, 'data-label="Кабинет"')

    def test_tomorrow_prefers_single_participant_over_legacy_child(self):
        day = timezone.localdate() + timedelta(days=1)
        start = _local_dt(day, time(11, 0))
        participant_child = Child.objects.create(last_name="Завтра", first_name="Участник")
        assigned_staff = StaffMember.objects.create(full_name="Завтра Назначенный")
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
        )
        appointment.participants.filter(child=self.child).delete()
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=participant_child,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )
        appointment.staff_assignments.filter(staff_member=self.staff).delete()
        AppointmentStaffAssignment.objects.create(
            appointment=appointment,
            staff_member=assigned_staff,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )

        response = self.client.get(reverse("tomorrow"), {"date": day.isoformat()})

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, participant_child.full_name)
        self.assertContains(response, assigned_staff.full_name)


    def test_tomorrow_shows_confirmation_targets_and_low_balances(self):
        day = timezone.localdate() + timedelta(days=1)
        start = _local_dt(day, time(12, 0))
        participant_child = Child.objects.create(last_name="Tomorrow", first_name="Target")
        assistant = StaffMember.objects.create(
            full_name="Tomorrow Assistant",
            email="tomorrow-assistant@example.local",
        )
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            session_type=Appointment.SessionType.GROUP,
            title="Tomorrow confirmation group",
        )
        participant = AppointmentParticipant.objects.create(
            appointment=appointment,
            child=participant_child,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )
        staff_assignment = AppointmentStaffAssignment.objects.create(
            appointment=appointment,
            staff_member=assistant,
            role=AppointmentStaffAssignment.Role.ASSISTANT,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )
        AppointmentConfirmation.objects.create(
            appointment=appointment,
            target_type=AppointmentConfirmation.TargetType.REPRESENTATIVE,
            participant=participant,
            email="tomorrow-target@example.local",
            subject="Confirm",
            message="Confirm participant.",
        )
        AppointmentConfirmation.objects.create(
            appointment=appointment,
            target_type=AppointmentConfirmation.TargetType.SPECIALIST,
            staff_assignment=staff_assignment,
            email=assistant.email,
            subject="Confirm",
            message="Confirm staff.",
        )
        low_balance_child = Child.objects.create(last_name="Low", first_name="Balance")
        BalanceAccount.objects.create(
            child=low_balance_child,
            funding_source=self.funding,
            unit=BalanceAccount.Unit.SESSIONS,
            service_scope=BalanceAccount.ServiceScope.ANY,
            initial_amount=Decimal("6"),
        )

        response = self.client.get(reverse("tomorrow"), {"date": day.isoformat()})

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, participant_child.full_name)
        self.assertContains(response, assistant.full_name)
        self.assertContains(response, "tomorrow-target@example.local")
        self.assertContains(response, "tomorrow-assistant@example.local")
        self.assertContains(response, low_balance_child.full_name)


class WorkQueueViewTests(NewViewsTestBase):
    def _financial_integrity_fixture(self):
        day = timezone.localdate() - timedelta(days=1)
        start = _local_dt(day, time(10, 0))
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            status=Appointment.Status.COMPLETED,
            attendance_status=Appointment.AttendanceStatus.ATTENDED,
            billing_decision=Appointment.BillingDecision.DO_NOT_CHARGE,
        )
        LedgerEntry.objects.create(
            appointment=appointment,
            account=self.account,
            entry_type=LedgerEntry.EntryType.DEBIT,
            amount=Decimal("-1"),
            reason="Stale work queue ledger",
            created_by=self.admin,
        )
        financial_integrity_checks_svc.run_financial_integrity_check(
            appointments=[appointment],
            requested_by=self.admin,
        )
        return appointment

    def test_work_queue_shows_next_action_and_section_links(self):
        day = timezone.localdate() - timedelta(days=1)
        start = _local_dt(day, time(10, 0))
        Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            status=Appointment.Status.COMPLETED,
            attendance_status=Appointment.AttendanceStatus.ATTENDED,
        )

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("queue_summary_items", response.context)
        self.assertIn("queue_next_action", response.context)
        self.assertEqual(response.context["queue_next_action"]["title"], "Решения по списанию")
        self.assertContains(response, "Следующее действие")
        self.assertContains(response, "#queue-billing")
        self.assertContains(response, 'id="queue-billing"')

    def test_dashboard_surfaces_financial_integrity_issues(self):
        self._financial_integrity_fixture()

        response = self.client.get(reverse("dashboard"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context["financial_integrity_count"], 1)
        self.assertEqual(response.context["financial_integrity_warning_count"], 1)
        self.assertEqual(response.context["priority_total"], 1)
        self.assertContains(response, "Проверить финансы")
        self.assertContains(response, "финансовых расхождений")
        self.assertContains(response, f'{reverse("work_queue")}#queue-financial-integrity')
        self.assertContains(response, reverse("financial_integrity_report"))

    def test_dashboard_uses_persisted_financial_findings_without_live_audit(self):
        appointment = self._financial_integrity_fixture()
        LedgerEntry.objects.filter(appointment=appointment).delete()

        response = self.client.get(reverse("dashboard"))

        self.assertEqual(response.status_code, 200)
        self.assertFalse(LedgerEntry.objects.filter(appointment=appointment).exists())
        self.assertEqual(response.context["financial_integrity_count"], 1)
        self.assertEqual(response.context["financial_integrity_warning_count"], 1)

    def test_resolved_financial_findings_are_hidden_from_queue_and_dashboard(self):
        self._financial_integrity_fixture()
        FinancialIntegrityFinding.objects.update(
            status=FinancialIntegrityFinding.Status.RESOLVED,
            resolved_at=timezone.now(),
        )

        dashboard_response = self.client.get(reverse("dashboard"))
        queue_response = self.client.get(reverse("work_queue"))

        self.assertEqual(dashboard_response.status_code, 200)
        self.assertEqual(queue_response.status_code, 200)
        self.assertEqual(dashboard_response.context["financial_integrity_count"], 0)
        self.assertEqual(queue_response.context["financial_integrity_issue_count"], 0)
        self.assertNotContains(queue_response, "stale_debit_ledger_without_charge_fact")

    def test_work_queue_surfaces_financial_integrity_issues(self):
        appointment = self._financial_integrity_fixture()

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context["financial_integrity_issue_count"], 1)
        financial_summary = next(
            item
            for item in response.context["queue_summary_items"]
            if item["href"] == "#queue-financial-integrity"
        )
        self.assertEqual(financial_summary["value"], 1)
        self.assertEqual(financial_summary["tone"], "warning")
        self.assertContains(response, "#queue-financial-integrity")
        self.assertContains(response, 'id="queue-financial-integrity"')
        self.assertContains(response, "Финансовый контроль")
        self.assertContains(response, "stale_debit_ledger_without_charge_fact")
        self.assertContains(response, "Debit ledger-проводка")
        self.assertContains(response, reverse("appointment_detail", args=[appointment.pk]))
        finding = FinancialIntegrityFinding.objects.get()
        self.assertContains(response, reverse("financial_integrity_finding_detail", args=[finding.pk]))
        self.assertContains(response, reverse("financial_integrity_report"))
        self.assertNotContains(response, "Исправить автоматически")

    def test_work_queue_shows_financial_integrity_run_hint_without_runs(self):
        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(FinancialIntegrityCheckRun.objects.count(), 0)
        self.assertIsNone(response.context["financial_integrity_latest_run"])
        self.assertEqual(response.context["financial_integrity_latest_run_tone"], "info")
        self.assertContains(response, "Сохраненная проверка еще не запускалась")
        self.assertContains(response, "run_financial_integrity_check --run-type scheduled")

    def test_work_queue_shows_latest_financial_integrity_run_summary(self):
        self._financial_integrity_fixture()
        run = FinancialIntegrityCheckRun.objects.get()

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(FinancialIntegrityCheckRun.objects.count(), 1)
        self.assertEqual(response.context["financial_integrity_latest_run"], run)
        self.assertEqual(response.context["financial_integrity_latest_run_tone"], "warning")
        self.assertContains(response, f"#{run.pk}")
        self.assertContains(response, "расхождений: 1")

    def test_work_queue_shows_failed_financial_integrity_run_summary(self):
        run = FinancialIntegrityCheckRun.objects.create(
            run_type=FinancialIntegrityCheckRun.RunType.SCHEDULED,
            status=FinancialIntegrityCheckRun.Status.FAILED,
            started_at=timezone.now(),
            finished_at=timezone.now(),
            error_message="QA runner failure",
        )

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context["financial_integrity_latest_run"], run)
        self.assertEqual(response.context["financial_integrity_latest_run_tone"], "danger")
        self.assertContains(response, "QA runner failure")

    def test_financial_integrity_report_is_read_only_and_counts_period(self):
        self._financial_integrity_fixture()
        finding = FinancialIntegrityFinding.objects.get()
        run_count = FinancialIntegrityCheckRun.objects.count()
        event_count = FinancialIntegrityFindingEvent.objects.count()

        response = self.client.get(reverse("financial_integrity_report"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(FinancialIntegrityCheckRun.objects.count(), run_count)
        self.assertEqual(FinancialIntegrityFindingEvent.objects.count(), event_count)
        report = response.context["report"]
        self.assertEqual(report["active_summary"]["total"], 1)
        self.assertEqual(report["event_counts"]["created"], 1)
        self.assertEqual(report["run_counts"]["total"], 1)
        self.assertEqual(report["run_counts"]["checked"], 1)
        self.assertEqual(report["run_counts"]["issues"], 1)
        self.assertEqual(report["period"]["selected_period"], "30")
        self.assertContains(response, "financial-report-summary-grid")
        self.assertContains(response, reverse("financial_integrity_finding_detail", args=[finding.pk]))
        self.assertContains(response, f'{reverse("work_queue")}#queue-financial-integrity')

    def test_financial_integrity_report_custom_period_keeps_current_active_snapshot(self):
        self._financial_integrity_fixture()
        date_from = timezone.localdate() - timedelta(days=90)
        date_to = timezone.localdate() - timedelta(days=60)

        response = self.client.get(
            reverse("financial_integrity_report"),
            {
                "date_from": date_from.isoformat(),
                "date_to": date_to.isoformat(),
            },
        )

        self.assertEqual(response.status_code, 200)
        report = response.context["report"]
        self.assertEqual(report["period"]["selected_period"], "custom")
        self.assertEqual(report["event_counts"]["created"], 0)
        self.assertEqual(report["run_counts"]["total"], 0)
        self.assertEqual(report["active_summary"]["total"], 1)

    def test_financial_integrity_report_requires_admin(self):
        self._financial_integrity_fixture()
        user = User.objects.create_user("financial-report-non-admin", password="x")
        self.client.force_login(user)

        response = self.client.get(reverse("financial_integrity_report"))

        self.assertEqual(response.status_code, 302)

    def test_financial_integrity_finding_detail_shows_sources_and_actions(self):
        self._financial_integrity_fixture()
        finding = FinancialIntegrityFinding.objects.get()
        event_count = FinancialIntegrityFindingEvent.objects.count()

        response = self.client.get(reverse("financial_integrity_finding_detail", args=[finding.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(FinancialIntegrityFindingEvent.objects.count(), event_count)
        self.assertContains(response, finding.issue_key)
        self.assertContains(response, finding.code)
        self.assertContains(response, finding.message)
        self.assertContains(response, reverse("appointment_detail", args=[finding.appointment_id]))
        self.assertContains(response, reverse("balance_account_edit", args=[finding.account_id]))
        self.assertContains(response, reverse("funding_source_edit", args=[finding.funding_source_id]))
        self.assertContains(response, 'name="action" value="acknowledge"')
        self.assertContains(response, 'name="action" value="ignore"')
        self.assertContains(response, reverse("financial_integrity_finding_recheck", args=[finding.pk]))
        self.assertContains(response, "financial-integrity-history")
        self.assertContains(response, "financial-integrity-event")
        self.assertContains(response, "financial-integrity-payload")
        self.assertEqual(len(response.context["event_rows"]), 1)
        self.assertEqual(
            response.context["event_rows"][0]["event"].event_type,
            FinancialIntegrityFindingEvent.EventType.CREATED,
        )

    def test_financial_integrity_finding_detail_handles_missing_source_objects(self):
        self._financial_integrity_fixture()
        finding = FinancialIntegrityFinding.objects.get()
        appointment_url = reverse("appointment_detail", args=[finding.appointment_id])
        finding.appointment = None
        finding.appointment_participant = None
        finding.ledger_entry = None
        finding.account = None
        finding.funding_source = None
        finding.participant_name = "Detached Participant"
        finding.account_label = "Detached Account"
        finding.funding_source_name = "Detached Funding"
        finding.payload = {"source": "detached"}
        finding.save(
            update_fields=[
                "appointment",
                "appointment_participant",
                "ledger_entry",
                "account",
                "funding_source",
                "participant_name",
                "account_label",
                "funding_source_name",
                "payload",
                "updated_at",
            ]
        )

        response = self.client.get(reverse("financial_integrity_finding_detail", args=[finding.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, appointment_url)
        self.assertNotContains(response, reverse("financial_integrity_finding_recheck", args=[finding.pk]))
        self.assertContains(response, "Detached Participant")
        self.assertContains(response, "Detached Account")
        self.assertContains(response, "Detached Funding")

    def test_financial_integrity_finding_recheck_scopes_to_appointment(self):
        self._financial_integrity_fixture()
        finding = FinancialIntegrityFinding.objects.get()
        detail_url = reverse("financial_integrity_finding_detail", args=[finding.pk])
        run_count = FinancialIntegrityCheckRun.objects.count()

        response = self.client.post(
            reverse("financial_integrity_finding_recheck", args=[finding.pk]),
            {"next": detail_url},
        )

        self.assertEqual(response.status_code, 302)
        self.assertEqual(response["Location"], detail_url)
        self.assertEqual(FinancialIntegrityCheckRun.objects.count(), run_count + 1)
        run = FinancialIntegrityCheckRun.objects.latest("pk")
        self.assertEqual(run.status, FinancialIntegrityCheckRun.Status.COMPLETED)
        self.assertEqual(run.run_type, FinancialIntegrityCheckRun.RunType.MANUAL)
        self.assertEqual(run.requested_by, self.admin)
        self.assertEqual(run.candidate_count, 1)
        self.assertEqual(run.issue_count, 1)
        event = FinancialIntegrityFindingEvent.objects.latest("pk")
        self.assertEqual(event.event_type, FinancialIntegrityFindingEvent.EventType.SCOPED_RECHECK)
        self.assertEqual(event.finding, finding)
        self.assertEqual(event.run, run)
        self.assertEqual(event.actor, self.admin)
        self.assertEqual(event.status_from, FinancialIntegrityFinding.Status.OPEN)
        self.assertEqual(event.status_to, FinancialIntegrityFinding.Status.OPEN)
        finding.refresh_from_db()
        self.assertEqual(finding.last_seen_run, run)
        self.assertEqual(finding.status, FinancialIntegrityFinding.Status.OPEN)

    def test_financial_integrity_finding_recheck_without_appointment_does_not_run(self):
        self._financial_integrity_fixture()
        finding = FinancialIntegrityFinding.objects.get()
        finding.appointment = None
        finding.save(update_fields=["appointment", "updated_at"])
        detail_url = reverse("financial_integrity_finding_detail", args=[finding.pk])
        run_count = FinancialIntegrityCheckRun.objects.count()

        response = self.client.post(
            reverse("financial_integrity_finding_recheck", args=[finding.pk]),
            {"next": detail_url},
        )

        self.assertEqual(response.status_code, 302)
        self.assertEqual(response["Location"], detail_url)
        self.assertEqual(FinancialIntegrityCheckRun.objects.count(), run_count)

    def test_work_queue_financial_integrity_can_acknowledge_finding(self):
        self._financial_integrity_fixture()
        finding = FinancialIntegrityFinding.objects.get()

        response = self.client.post(
            reverse("financial_integrity_finding_triage", args=[finding.pk]),
            {
                "action": "acknowledge",
                "next": f"{reverse('work_queue')}#queue-financial-integrity",
            },
        )

        self.assertEqual(response.status_code, 302)
        self.assertEqual(response["Location"], f"{reverse('work_queue')}#queue-financial-integrity")
        finding.refresh_from_db()
        self.assertEqual(finding.status, FinancialIntegrityFinding.Status.ACKNOWLEDGED)
        self.assertEqual(finding.triaged_by, self.admin)
        self.assertIsNotNone(finding.triaged_at)

    def test_work_queue_financial_integrity_can_return_acknowledged_finding_to_open(self):
        self._financial_integrity_fixture()
        finding = FinancialIntegrityFinding.objects.get()
        finding.status = FinancialIntegrityFinding.Status.ACKNOWLEDGED
        finding.triaged_by = self.admin
        finding.triaged_at = timezone.now()
        finding.save(update_fields=["status", "triaged_by", "triaged_at", "updated_at"])

        response = self.client.post(
            reverse("financial_integrity_finding_triage", args=[finding.pk]),
            {
                "action": "return_to_open",
                "next": f"{reverse('work_queue')}#queue-financial-integrity",
            },
        )

        self.assertEqual(response.status_code, 302)
        finding.refresh_from_db()
        self.assertEqual(finding.status, FinancialIntegrityFinding.Status.OPEN)

    def test_work_queue_financial_integrity_ignore_requires_note(self):
        self._financial_integrity_fixture()
        finding = FinancialIntegrityFinding.objects.get()

        response = self.client.post(
            reverse("financial_integrity_finding_triage", args=[finding.pk]),
            {
                "action": "ignore",
                "note": "   ",
                "next": f"{reverse('work_queue')}#queue-financial-integrity",
            },
        )

        self.assertEqual(response.status_code, 302)
        finding.refresh_from_db()
        self.assertEqual(finding.status, FinancialIntegrityFinding.Status.OPEN)
        self.assertEqual(finding.triage_note, "")
        queue_response = self.client.get(reverse("work_queue"))
        self.assertEqual(queue_response.context["financial_integrity_issue_count"], 1)
        self.assertContains(queue_response, "stale_debit_ledger_without_charge_fact")

    def test_work_queue_financial_integrity_ignore_hides_finding_from_queue(self):
        self._financial_integrity_fixture()
        finding = FinancialIntegrityFinding.objects.get()

        response = self.client.post(
            reverse("financial_integrity_finding_triage", args=[finding.pk]),
            {
                "action": "ignore",
                "note": "Проверено вручную, корректировка не нужна.",
                "next": f"{reverse('work_queue')}#queue-financial-integrity",
            },
        )

        self.assertEqual(response.status_code, 302)
        finding.refresh_from_db()
        self.assertEqual(finding.status, FinancialIntegrityFinding.Status.IGNORED)
        self.assertEqual(finding.triage_note, "Проверено вручную, корректировка не нужна.")
        queue_response = self.client.get(reverse("work_queue"))
        self.assertEqual(queue_response.context["financial_integrity_issue_count"], 0)
        self.assertNotContains(queue_response, "stale_debit_ledger_without_charge_fact")
        detail_response = self.client.get(reverse("financial_integrity_finding_detail", args=[finding.pk]))
        self.assertContains(detail_response, "financial-integrity-event-note")
        self.assertEqual(detail_response.context["event_rows"][0]["event"].note, finding.triage_note)

    def test_work_queue_financial_integrity_triage_rejects_external_next(self):
        self._financial_integrity_fixture()
        finding = FinancialIntegrityFinding.objects.get()

        response = self.client.post(
            reverse("financial_integrity_finding_triage", args=[finding.pk]),
            {
                "action": "acknowledge",
                "next": "https://evil.example/phish",
            },
        )

        self.assertEqual(response.status_code, 302)
        self.assertEqual(response["Location"], f"{reverse('work_queue')}#queue-financial-integrity")
        finding.refresh_from_db()
        self.assertEqual(finding.status, FinancialIntegrityFinding.Status.ACKNOWLEDGED)

    def test_work_queue_financial_integrity_triage_requires_admin(self):
        self._financial_integrity_fixture()
        finding = FinancialIntegrityFinding.objects.get()
        user = User.objects.create_user("not-admin", password="x")
        self.client.force_login(user)

        response = self.client.post(
            reverse("financial_integrity_finding_triage", args=[finding.pk]),
            {
                "action": "acknowledge",
                "next": f"{reverse('work_queue')}#queue-financial-integrity",
            },
        )

        self.assertEqual(response.status_code, 302)
        finding.refresh_from_db()
        self.assertEqual(finding.status, FinancialIntegrityFinding.Status.OPEN)

    def test_group_billing_task_links_to_participant_decisions(self):
        day = timezone.localdate() - timedelta(days=1)
        start = _local_dt(day, time(10, 0))
        second_parent = ParentGuardian.objects.create(
            last_name="Сидорова",
            first_name="Ольга",
            phone="+7 900 000-00-03",
        )
        second_child = Child.objects.create(
            last_name="Сидоров",
            first_name="Павел",
            primary_parent=second_parent,
        )
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            status=Appointment.Status.COMPLETED,
            attendance_status=Appointment.AttendanceStatus.ATTENDED,
            session_type=Appointment.SessionType.GROUP,
            title="Группа чтения",
        )
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=second_child,
            attendance_status=Appointment.AttendanceStatus.ATTENDED,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Группа чтения")
        self.assertContains(response, "Сидоров Павел")
        self.assertContains(response, "Решить по участникам")
        self.assertNotContains(response, 'value="charge"')

    def test_group_billing_task_uses_undecided_participant_over_legacy_summary(self):
        day = timezone.localdate() - timedelta(days=1)
        start = _local_dt(day, time(10, 0))
        second_child = Child.objects.create(last_name="Очередь", first_name="Без решения")
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            status=Appointment.Status.COMPLETED,
            attendance_status=Appointment.AttendanceStatus.ATTENDED,
            billing_decision=Appointment.BillingDecision.CHARGE,
            billing_account=self.account,
            session_type=Appointment.SessionType.GROUP,
            title="Группа с нерешенным участником",
        )
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=second_child,
            billing_decision=Appointment.BillingDecision.UNDECIDED,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertIn(appointment, list(response.context["needs_billing"]))
        self.assertContains(response, "Группа с нерешенным участником")
        self.assertContains(response, second_child.full_name)

    def test_billing_task_prefers_single_participant_and_assignment(self):
        day = timezone.localdate() - timedelta(days=1)
        start = _local_dt(day, time(11, 0))
        participant_child = Child.objects.create(last_name="Очередь", first_name="Получатель")
        assigned_staff = StaffMember.objects.create(full_name="Очередь Специалист")
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            status=Appointment.Status.COMPLETED,
            attendance_status=Appointment.AttendanceStatus.ATTENDED,
        )
        appointment.participants.filter(child=self.child).delete()
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=participant_child,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )
        appointment.staff_assignments.filter(staff_member=self.staff).delete()
        AppointmentStaffAssignment.objects.create(
            appointment=appointment,
            staff_member=assigned_staff,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, participant_child.full_name)
        self.assertContains(response, assigned_staff.full_name)

    def test_billing_task_uses_single_participant_attendance_label(self):
        day = timezone.localdate() - timedelta(days=1)
        start = _local_dt(day, time(11, 0))
        participant_child = Child.objects.create(last_name="Очередь", first_name="Посещение")
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            status=Appointment.Status.COMPLETED,
            attendance_status=Appointment.AttendanceStatus.ATTENDED,
        )
        appointment.participants.filter(child=self.child).delete()
        participant = AppointmentParticipant.objects.create(
            appointment=appointment,
            child=participant_child,
            attendance_status=Appointment.AttendanceStatus.MISSED,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, participant_child.full_name)
        self.assertContains(response, participant.get_attendance_status_display())

    def test_billing_task_quick_charge_uses_single_participant_account(self):
        day = timezone.localdate() - timedelta(days=1)
        start = _local_dt(day, time(11, 0))
        participant_child = Child.objects.create(last_name="Очередь", first_name="Счет")
        participant_account = BalanceAccount.objects.create(
            child=participant_child,
            funding_source=self.funding,
            unit=BalanceAccount.Unit.SESSIONS,
            service_scope=BalanceAccount.ServiceScope.ANY,
            initial_amount=Decimal("5"),
        )
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            status=Appointment.Status.COMPLETED,
            attendance_status=Appointment.AttendanceStatus.ATTENDED,
        )
        appointment.participants.filter(child=self.child).delete()
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=participant_child,
            billing_account=participant_account,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, participant_child.full_name)
        self.assertContains(response, f'value="{participant_account.pk}"')

    def test_group_with_unmarked_participant_stays_in_attendance_queue(self):
        day = timezone.localdate() - timedelta(days=1)
        start = _local_dt(day, time(10, 0))
        second_parent = ParentGuardian.objects.create(
            last_name="Кузнецова",
            first_name="Ирина",
            phone="+7 900 000-00-04",
        )
        second_child = Child.objects.create(
            last_name="Кузнецов",
            first_name="Дима",
            primary_parent=second_parent,
        )
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            status=Appointment.Status.COMPLETED,
            attendance_status=Appointment.AttendanceStatus.ATTENDED,
            session_type=Appointment.SessionType.GROUP,
            title="Группа моторики",
        )
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=second_child,
            attendance_status=Appointment.AttendanceStatus.UNKNOWN,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertIn(appointment, list(response.context["needs_attendance"]))
        self.assertContains(response, "Группа моторики")
        self.assertContains(response, "Кузнецов Дима")

    def test_group_confirmation_task_shows_participant_names(self):
        day = timezone.localdate() + timedelta(days=1)
        start = _local_dt(day, time(10, 0))
        second_parent = ParentGuardian.objects.create(
            last_name="Орлова",
            first_name="Нина",
            phone="+7 900 000-00-05",
            email="orlova@example.local",
        )
        second_child = Child.objects.create(
            last_name="Орлов",
            first_name="Лев",
            primary_parent=second_parent,
        )
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            session_type=Appointment.SessionType.GROUP,
            title="Группа согласования",
        )
        participant = AppointmentParticipant.objects.create(
            appointment=appointment,
            child=second_child,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )
        AppointmentConfirmation.objects.create(
            appointment=appointment,
            target_type=AppointmentConfirmation.TargetType.REPRESENTATIVE,
            participant=participant,
            representative=second_parent,
            email=second_parent.email,
            subject="Подтвердите занятие",
            message="Проверьте дату.",
        )

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Группа согласования")
        self.assertContains(response, "Орлов Лев")

    def test_confirmation_task_prefers_single_participant_over_legacy_child(self):
        day = timezone.localdate() + timedelta(days=1)
        start = _local_dt(day, time(11, 0))
        legacy_child = self.child
        participant_child = Child.objects.create(last_name="Snapshot", first_name="Only")
        appointment = Appointment.objects.create(
            child=legacy_child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
        )
        appointment.participants.filter(child=legacy_child).delete()
        participant = AppointmentParticipant.objects.create(
            appointment=appointment,
            child=participant_child,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )
        AppointmentConfirmation.objects.create(
            appointment=appointment,
            target_type=AppointmentConfirmation.TargetType.RECIPIENT,
            participant=participant,
            email="snapshot-only@example.local",
            subject="Подтвердите занятие",
            message="Проверьте дату.",
        )

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, participant_child.full_name)

    def test_confirmation_task_shows_target_staff_assignment(self):
        day = timezone.localdate() + timedelta(days=1)
        start = _local_dt(day, time(12, 0))
        assistant = StaffMember.objects.create(
            full_name="Assistant Queue",
            email="assistant-queue@example.local",
        )
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            session_type=Appointment.SessionType.GROUP,
            title="Queue staff group",
        )
        second_child = Child.objects.create(last_name="Queue", first_name="Participant")
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=second_child,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )
        assignment = AppointmentStaffAssignment.objects.create(
            appointment=appointment,
            staff_member=assistant,
            role=AppointmentStaffAssignment.Role.ASSISTANT,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )
        AppointmentConfirmation.objects.create(
            appointment=appointment,
            target_type=AppointmentConfirmation.TargetType.SPECIALIST,
            staff_assignment=assignment,
            email=assistant.email,
            subject="Confirm",
            message="Please confirm.",
        )

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Queue staff group")
        self.assertContains(response, "Assistant Queue")

    def test_reschedule_step_confirmation_links_to_step_detail(self):
        day = timezone.localdate() + timedelta(days=5)
        start = _local_dt(day, time(10, 0))
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            billing_account=self.account,
        )
        plan = AppointmentReschedulePlan.objects.create(
            status=AppointmentReschedulePlan.Status.READY,
            plan_type=AppointmentReschedulePlan.PlanType.MANUAL,
            reason="Confirmation queue step fixture",
            created_by=self.admin,
        )
        step = AppointmentRescheduleStep.objects.create(
            plan=plan,
            position=1,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.VALID,
            confirmation_status=AppointmentRescheduleStep.ConfirmationStatus.WAITING,
            source_appointment=appointment,
            proposed_starts_at=start + timedelta(hours=1),
            proposed_ends_at=start + timedelta(hours=1, minutes=30),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
        )
        AppointmentConfirmation.objects.create(
            appointment=appointment,
            reschedule_step=step,
            target_type=AppointmentConfirmation.TargetType.REPRESENTATIVE,
            email="step-confirmation@example.local",
            subject="Согласуйте перенос",
            message="Проверьте новое время.",
        )

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Согласование шага переноса")
        self.assertContains(response, "Открыть шаг переноса")
        self.assertContains(
            response,
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#step-{step.pk}',
        )
        self.assertNotContains(response, f'{reverse("appointment_detail", args=[appointment.pk])}">Открыть занятие')

    def test_reschedule_step_confirmation_ignores_terminal_plan_in_queue(self):
        day = timezone.localdate() + timedelta(days=5)
        start = _local_dt(day, time(10, 0))
        active_appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            billing_account=self.account,
        )
        terminal_appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start + timedelta(hours=1),
            ends_at=start + timedelta(hours=1, minutes=30),
            billing_account=self.account,
        )
        active_plan = AppointmentReschedulePlan.objects.create(
            status=AppointmentReschedulePlan.Status.READY,
            plan_type=AppointmentReschedulePlan.PlanType.MANUAL,
            reason="Active confirmation queue step fixture",
            created_by=self.admin,
        )
        terminal_plan = AppointmentReschedulePlan.objects.create(
            status=AppointmentReschedulePlan.Status.CANCELLED,
            plan_type=AppointmentReschedulePlan.PlanType.MANUAL,
            reason="Terminal confirmation queue step fixture",
            created_by=self.admin,
        )
        active_step = AppointmentRescheduleStep.objects.create(
            plan=active_plan,
            position=1,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.VALID,
            confirmation_status=AppointmentRescheduleStep.ConfirmationStatus.WAITING,
            source_appointment=active_appointment,
            proposed_starts_at=start + timedelta(hours=2),
            proposed_ends_at=start + timedelta(hours=2, minutes=30),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
        )
        terminal_step = AppointmentRescheduleStep.objects.create(
            plan=terminal_plan,
            position=1,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.VALID,
            confirmation_status=AppointmentRescheduleStep.ConfirmationStatus.WAITING,
            source_appointment=terminal_appointment,
            proposed_starts_at=start + timedelta(hours=3),
            proposed_ends_at=start + timedelta(hours=3, minutes=30),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
        )
        active_confirmation = AppointmentConfirmation.objects.create(
            appointment=active_appointment,
            reschedule_step=active_step,
            target_type=AppointmentConfirmation.TargetType.REPRESENTATIVE,
            email="active-step-confirmation@example.local",
            subject="Согласуйте перенос",
            message="Проверьте новое время.",
        )
        terminal_confirmation = AppointmentConfirmation.objects.create(
            appointment=terminal_appointment,
            reschedule_step=terminal_step,
            target_type=AppointmentConfirmation.TargetType.REPRESENTATIVE,
            email="terminal-step-confirmation@example.local",
            subject="Согласуйте перенос",
            message="Проверьте новое время.",
        )

        response = self.client.get(reverse("work_queue"))
        dashboard_response = self.client.get(reverse("dashboard"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(dashboard_response.status_code, 200)
        self.assertIn(active_confirmation, list(response.context["confirmation_tasks"]))
        self.assertNotIn(terminal_confirmation, list(response.context["confirmation_tasks"]))
        self.assertEqual(dashboard_response.context["confirmation_tasks"], 1)
        self.assertContains(response, "active-step-confirmation@example.local")
        self.assertNotContains(response, "terminal-step-confirmation@example.local")

    def test_low_balance_task_links_to_recipient_payment_and_account(self):
        low_balance_child = Child.objects.create(last_name="Низкий", first_name="Баланс")
        account = BalanceAccount.objects.create(
            child=low_balance_child,
            funding_source=self.funding,
            unit=BalanceAccount.Unit.SESSIONS,
            service_scope=BalanceAccount.ServiceScope.SPECIFIC_SERVICE,
            service=self.service,
            initial_amount=Decimal("1"),
        )

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, low_balance_child.full_name)
        self.assertContains(response, reverse("recipient_detail", args=[low_balance_child.pk]))
        self.assertContains(response, reverse("payment_create_for_account", args=[account.pk]))
        self.assertContains(response, reverse("balance_account_edit", args=[account.pk]))
        self.assertContains(response, account.warning_label)
        self.assertContains(response, "queue-balance-table")
        self.assertContains(response, 'data-label="Остаток"')
        self.assertContains(response, 'data-label="Действия"')


class StaffTimesheetViewTests(NewViewsTestBase):
    def test_timesheet_renders(self):
        response = self.client.get(reverse("staff_timesheet", args=[self.staff.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertIn("form", response.context)
        self.assertIn("sheet", response.context)
        self.assertIn("timesheet_summary_items", response.context)
        self.assertIn("timesheet_attention_items", response.context)
        self.assertContains(response, "Сводка табеля")
        self.assertContains(response, "Контроль перед выплатой")

    def test_timesheet_shows_payroll_control_summary(self):
        day = timezone.localdate()
        start = _local_dt(day, time(10, 0))
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            billing_account=self.account,
            status=Appointment.Status.COMPLETED,
            billing_decision=Appointment.BillingDecision.CHARGE,
        )

        response = self.client.get(
            reverse("staff_timesheet", args=[self.staff.pk]),
            {"date_from": day.isoformat(), "date_to": day.isoformat()},
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Сводка табеля")
        self.assertContains(response, "К начислению")
        self.assertContains(response, "Есть занятия без ставки")
        self.assertContains(response, "Начисления еще не сохранены")
        self.assertContains(response, reverse("appointment_detail", args=[appointment.pk]))
        self.assertContains(response, "timesheet-mobile-table")
        self.assertContains(response, 'data-label="Основание"')
        self.assertContains(response, 'data-label="К начислению"')

    def test_timesheet_csv_export(self):
        day = timezone.localdate()
        start = _local_dt(day, time(10, 0))
        Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            billing_account=self.account,
            status=Appointment.Status.COMPLETED,
            billing_decision=Appointment.BillingDecision.CHARGE,
        )
        StaffCompensationRule.objects.create(
            staff_member=self.staff,
            service=self.service,
            amount=Decimal("500"),
        )
        response = self.client.get(
            reverse("staff_timesheet", args=[self.staff.pk]),
            {
                "date_from": (day - timedelta(days=1)).isoformat(),
                "date_to": (day + timedelta(days=1)).isoformat(),
                "csv": "1",
            },
        )
        self.assertEqual(response.status_code, 200)
        self.assertIn("text/csv", response["Content-Type"])
        body = response.content.decode("utf-8")
        self.assertIn(
            "Дата;Всего;Проведено;Отменено;Не явился;Часов;К начислению;Сумма начисления", body
        )
        self.assertIn("Итого;", body)
        self.assertIn(";1;500", body)

    def test_timesheet_grant_allocation_rate_is_not_warning(self):
        day = timezone.localdate()
        start = _local_dt(day, time(10, 0))
        FundingStaffAllocation.objects.create(
            funding_source=self.funding,
            service=self.service,
            staff_member=self.staff,
            allocated_sessions=10,
            session_pay_amount=Decimal("520"),
            starts_on=day,
            ends_on=day,
        )
        Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            billing_account=self.account,
            status=Appointment.Status.COMPLETED,
            billing_decision=Appointment.BillingDecision.CHARGE,
        )

        response = self.client.get(
            reverse("staff_timesheet", args=[self.staff.pk]),
            {"date_from": day.isoformat(), "date_to": day.isoformat()},
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "грантовая квота")
        self.assertNotContains(response, 'class="warning-row"')

    def test_timesheet_post_generates_payroll_accruals(self):
        day = timezone.localdate()
        start = _local_dt(day, time(10, 0))
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            billing_account=self.account,
            status=Appointment.Status.COMPLETED,
        )
        billing_svc.apply_decision(
            appointment,
            decision=Appointment.BillingDecision.CHARGE,
            account=self.account,
            actor=self.admin,
        )
        StaffCompensationRule.objects.create(
            staff_member=self.staff,
            service=self.service,
            amount=Decimal("500"),
        )

        response = self.client.post(
            reverse("staff_timesheet", args=[self.staff.pk]),
            {
                "date_from": day.isoformat(),
                "date_to": day.isoformat(),
                "action": "generate_accruals",
            },
        )

        self.assertRedirects(
            response,
            f"{reverse('staff_timesheet', args=[self.staff.pk])}?date_from={day.isoformat()}&date_to={day.isoformat()}",
        )
        accrual = PayrollAccrual.objects.get(staff_member=self.staff)
        self.assertEqual(accrual.amount, Decimal("500"))

    def test_timesheet_can_create_and_approve_payroll_sheet(self):
        day = timezone.localdate()
        start = _local_dt(day, time(10, 0))
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            billing_account=self.account,
            status=Appointment.Status.COMPLETED,
        )
        billing_svc.apply_decision(
            appointment,
            decision=Appointment.BillingDecision.CHARGE,
            account=self.account,
            actor=self.admin,
        )
        StaffCompensationRule.objects.create(
            staff_member=self.staff,
            service=self.service,
            amount=Decimal("500"),
        )

        create_response = self.client.post(
            reverse("staff_timesheet", args=[self.staff.pk]),
            {
                "date_from": day.isoformat(),
                "date_to": day.isoformat(),
                "action": "create_payroll_sheet",
            },
        )

        payroll_sheet = PayrollSheet.objects.get(staff_member=self.staff)
        self.assertRedirects(
            create_response, reverse("payroll_sheet_detail", args=[payroll_sheet.pk])
        )
        detail_url = reverse("payroll_sheet_detail", args=[payroll_sheet.pk])
        draft_response = self.client.get(detail_url)
        self.assertIn("payroll_sheet_summary_items", draft_response.context)
        self.assertIn("payroll_sheet_next_action", draft_response.context)
        self.assertEqual(
            draft_response.context["payroll_sheet_next_action"]["title"],
            "Проверить и утвердить",
        )
        self.assertContains(draft_response, "Контроль расчетного листа")
        self.assertContains(draft_response, 'id="payroll-sheet-lines"')
        self.assertContains(draft_response, 'data-label="Сумма"')
        self.assertContains(draft_response, reverse("staff_timesheet", args=[self.staff.pk]))

        approve_response = self.client.post(detail_url, {"action": "approve"})
        payroll_sheet.refresh_from_db()
        self.assertRedirects(approve_response, detail_url)
        self.assertEqual(payroll_sheet.status, PayrollSheet.Status.APPROVED)
        self.assertEqual(
            PayrollAccrual.objects.get(staff_member=self.staff).status,
            PayrollAccrual.Status.APPROVED,
        )
        detail_response = self.client.get(detail_url)
        self.assertEqual(
            detail_response.context["payroll_sheet_next_action"]["title"],
            "Лист утвержден",
        )
        self.assertContains(detail_response, self.funding.name)
        self.assertContains(detail_response, "500,00 / За занятие")

    def test_timesheet_invalid_range_shows_error(self):
        response = self.client.get(
            reverse("staff_timesheet", args=[self.staff.pk]),
            {"date_from": "2099-01-01", "date_to": "2099-01-05"},
        )
        self.assertEqual(response.status_code, 200)


class GrantReportViewTests(NewViewsTestBase):
    def test_grant_report_renders(self):
        response = self.client.get(reverse("grant_report"))
        self.assertEqual(response.status_code, 200)
        self.assertIn("form", response.context)
        self.assertIn("grant_summary_items", response.context)
        self.assertIn("grant_attention_items", response.context)
        self.assertContains(response, "Сводка гранта")
        self.assertContains(response, "Контроль гранта")
        self.assertContains(response, reverse("funding_service_quota_create"))
        self.assertContains(response, reverse("funding_staff_allocation_create"))
        self.assertContains(response, reverse("grant_recipient_allocation_create"))

    def test_grant_quota_forms_show_operator_control(self):
        quota = FundingServiceQuota.objects.create(
            funding_source=self.funding_grant,
            service=self.service2,
            planned_sessions=30,
            starts_on="2026-01-01",
            ends_on="2026-12-31",
        )
        staff_allocation = FundingStaffAllocation.objects.create(
            service_quota=quota,
            funding_source=self.funding_grant,
            service=self.service2,
            staff_member=self.staff,
            allocated_sessions=8,
            session_pay_amount=Decimal("500.00"),
            starts_on="2026-01-01",
            ends_on="2026-12-31",
        )
        recipient_allocation = GrantRecipientAllocation.objects.create(
            funding_source=self.funding_grant,
            child=self.child,
            service=self.service2,
            allocated_sessions=12,
            balance_account=self.account_grant,
            valid_from="2026-01-01",
            valid_until="2026-12-31",
        )

        cases = [
            (
                f"{reverse('funding_service_quota_create')}?funding={self.funding_grant.pk}",
                "Контроль квоты услуги",
                "Следующий шаг",
            ),
            (
                reverse("funding_service_quota_edit", args=[quota.pk]),
                "Контроль квоты услуги",
                "Связанные распределения",
            ),
            (
                f"{reverse('funding_staff_allocation_create')}?quota={quota.pk}",
                "Контроль распределения специалисту",
                "Ставка специалисту",
            ),
            (
                reverse("funding_staff_allocation_edit", args=[staff_allocation.pk]),
                "Контроль распределения специалисту",
                "Связь с квотой",
            ),
            (
                f"{reverse('grant_recipient_allocation_create')}?funding={self.funding_grant.pk}",
                "Контроль выделения получателю",
                "Счет в занятиях",
            ),
            (
                reverse("grant_recipient_allocation_edit", args=[recipient_allocation.pk]),
                "Контроль выделения получателю",
                "Связанный счет",
            ),
        ]

        for url, title, control_item in cases:
            with self.subTest(url=url):
                response = self.client.get(url)
                self.assertEqual(response.status_code, 200)
                self.assertIn("object_form_control_items", response.context)
                self.assertContains(response, 'class="object-form-layout"')
                self.assertContains(response, title)
                self.assertContains(response, control_item)

    def test_grant_recipient_allocation_create_creates_balance_account(self):
        response = self.client.post(
            reverse("grant_recipient_allocation_create"),
            {
                "funding_source": self.funding_grant.pk,
                "child": self.child.pk,
                "service": self.service2.pk,
                "allocated_sessions": 14,
                "balance_account": "",
                "valid_from": "2026-01-01",
                "valid_until": "2026-12-31",
                "note": "Грант ребенку",
            },
        )

        allocation = GrantRecipientAllocation.objects.get(child=self.child, service=self.service2)
        self.assertEqual(allocation.allocated_sessions, 14)
        self.assertEqual(allocation.balance_account.child, self.child)
        self.assertEqual(allocation.balance_account.funding_source, self.funding_grant)
        self.assertEqual(allocation.balance_account.service, self.service2)
        self.assertEqual(allocation.balance_account.initial_amount, Decimal("14"))
        self.assertRedirects(
            response,
            (
                f"{reverse('grant_report')}?funding={self.funding_grant.pk}"
                "&date_from=2026-01-01&date_to=2026-12-31"
            ),
        )

    def test_grant_recipient_allocation_rejects_foreign_balance_account(self):
        response = self.client.post(
            reverse("grant_recipient_allocation_create"),
            {
                "funding_source": self.funding_grant.pk,
                "child": self.child.pk,
                "service": self.service2.pk,
                "allocated_sessions": 14,
                "balance_account": self.account.pk,
                "valid_from": "2026-01-01",
                "valid_until": "2026-12-31",
                "note": "",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn("balance_account", response.context["form"].errors)
        self.assertFalse(GrantRecipientAllocation.objects.exists())

    def test_grant_recipient_allocation_delete_keeps_balance_account(self):
        account = BalanceAccount.objects.create(
            child=self.child,
            funding_source=self.funding_grant,
            unit=BalanceAccount.Unit.SESSIONS,
            service_scope=BalanceAccount.ServiceScope.SPECIFIC_SERVICE,
            service=self.service2,
            initial_amount=Decimal("8"),
        )
        allocation = GrantRecipientAllocation.objects.create(
            funding_source=self.funding_grant,
            child=self.child,
            service=self.service2,
            allocated_sessions=8,
            balance_account=account,
            valid_from="2026-01-01",
            valid_until="2026-12-31",
        )

        response = self.client.post(
            reverse("grant_recipient_allocation_delete", args=[allocation.pk])
        )

        self.assertFalse(GrantRecipientAllocation.objects.filter(pk=allocation.pk).exists())
        self.assertTrue(BalanceAccount.objects.filter(pk=account.pk).exists())
        self.assertRedirects(
            response,
            (
                f"{reverse('grant_report')}?funding={self.funding_grant.pk}"
                "&date_from=2026-01-01&date_to=2026-12-31"
            ),
        )

    def test_funding_service_quota_create_redirects_to_allocation(self):
        response = self.client.post(
            reverse("funding_service_quota_create"),
            {
                "funding_source": self.funding_grant.pk,
                "service": self.service2.pk,
                "planned_sessions": 24,
                "starts_on": "2026-01-01",
                "ends_on": "2026-12-31",
                "note": "План на год",
            },
        )

        quota = FundingServiceQuota.objects.get(
            funding_source=self.funding_grant, service=self.service2
        )
        self.assertEqual(quota.planned_sessions, 24)
        self.assertRedirects(
            response,
            f"{reverse('funding_staff_allocation_create')}?quota={quota.pk}",
        )

    def test_funding_service_quota_edit_updates_plan(self):
        quota = FundingServiceQuota.objects.create(
            funding_source=self.funding_grant,
            service=self.service2,
            planned_sessions=20,
            starts_on="2026-01-01",
            ends_on="2026-12-31",
        )

        response = self.client.post(
            reverse("funding_service_quota_edit", args=[quota.pk]),
            {
                "funding_source": self.funding_grant.pk,
                "service": self.service2.pk,
                "planned_sessions": 30,
                "starts_on": "2026-01-01",
                "ends_on": "2026-12-31",
                "note": "Уточненный план",
            },
        )

        quota.refresh_from_db()
        self.assertEqual(quota.planned_sessions, 30)
        self.assertRedirects(
            response,
            (
                f"{reverse('grant_report')}?funding={self.funding_grant.pk}"
                "&date_from=2026-01-01&date_to=2026-12-31"
            ),
        )

    def test_funding_service_quota_delete_blocks_existing_allocations(self):
        quota = FundingServiceQuota.objects.create(
            funding_source=self.funding_grant,
            service=self.service2,
            planned_sessions=20,
            starts_on="2026-01-01",
            ends_on="2026-12-31",
        )
        FundingStaffAllocation.objects.create(
            service_quota=quota,
            funding_source=self.funding_grant,
            service=self.service2,
            staff_member=self.staff,
            allocated_sessions=5,
        )

        response = self.client.post(reverse("funding_service_quota_delete", args=[quota.pk]))

        self.assertTrue(FundingServiceQuota.objects.filter(pk=quota.pk).exists())
        self.assertRedirects(
            response,
            (
                f"{reverse('grant_report')}?funding={self.funding_grant.pk}"
                "&date_from=2026-01-01&date_to=2026-12-31"
            ),
        )

    def test_funding_staff_allocation_create_direct(self):
        response = self.client.post(
            reverse("funding_staff_allocation_create"),
            {
                "funding_source": self.funding_grant.pk,
                "service": self.service2.pk,
                "staff_member": self.staff.pk,
                "allocated_sessions": 12,
                "session_pay_amount": "450.00",
                "starts_on": "2026-01-01",
                "ends_on": "2026-12-31",
                "note": "Распределение руководителем",
            },
        )

        allocation = FundingStaffAllocation.objects.get(
            staff_member=self.staff, service=self.service2
        )
        self.assertIsNone(allocation.service_quota_id)
        self.assertEqual(allocation.funding_source, self.funding_grant)
        self.assertEqual(allocation.allocated_sessions, 12)
        self.assertEqual(allocation.session_pay_amount, Decimal("450.00"))
        self.assertRedirects(
            response,
            (
                f"{reverse('grant_report')}?funding={self.funding_grant.pk}"
                "&date_from=2026-01-01&date_to=2026-12-31"
            ),
        )

    def test_funding_staff_allocation_create_from_service_quota(self):
        quota = FundingServiceQuota.objects.create(
            funding_source=self.funding_grant,
            service=self.service2,
            planned_sessions=30,
            starts_on="2026-01-01",
            ends_on="2026-12-31",
        )

        response = self.client.post(
            reverse("funding_staff_allocation_create"),
            {
                "service_quota": quota.pk,
                "staff_member": self.staff.pk,
                "allocated_sessions": 8,
                "session_pay_amount": "",
                "starts_on": "2026-02-01",
                "ends_on": "2026-05-31",
                "note": "",
            },
        )

        allocation = FundingStaffAllocation.objects.get(
            service_quota=quota, staff_member=self.staff
        )
        self.assertEqual(allocation.funding_source, self.funding_grant)
        self.assertEqual(allocation.service, self.service2)
        self.assertEqual(allocation.allocated_sessions, 8)
        self.assertRedirects(
            response,
            (
                f"{reverse('grant_report')}?funding={self.funding_grant.pk}"
                "&date_from=2026-02-01&date_to=2026-05-31"
            ),
        )

    def test_funding_staff_allocation_rejects_quota_overallocation(self):
        quota = FundingServiceQuota.objects.create(
            funding_source=self.funding_grant,
            service=self.service2,
            planned_sessions=10,
        )
        FundingStaffAllocation.objects.create(
            service_quota=quota,
            funding_source=self.funding_grant,
            service=self.service2,
            staff_member=self.staff,
            allocated_sessions=8,
        )
        another_staff = StaffMember.objects.create(full_name="Петрова О. С.")

        response = self.client.post(
            reverse("funding_staff_allocation_create"),
            {
                "service_quota": quota.pk,
                "staff_member": another_staff.pk,
                "allocated_sessions": 5,
                "session_pay_amount": "",
                "starts_on": "",
                "ends_on": "",
                "note": "",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn("allocated_sessions", response.context["form"].errors)
        self.assertFalse(
            FundingStaffAllocation.objects.filter(
                service_quota=quota, staff_member=another_staff
            ).exists()
        )

    def test_funding_staff_allocation_edit_excludes_self_from_quota_limit(self):
        quota = FundingServiceQuota.objects.create(
            funding_source=self.funding_grant,
            service=self.service2,
            planned_sessions=10,
            starts_on="2026-01-01",
            ends_on="2026-12-31",
        )
        allocation = FundingStaffAllocation.objects.create(
            service_quota=quota,
            funding_source=self.funding_grant,
            service=self.service2,
            staff_member=self.staff,
            allocated_sessions=8,
            starts_on="2026-01-01",
            ends_on="2026-12-31",
        )

        response = self.client.post(
            reverse("funding_staff_allocation_edit", args=[allocation.pk]),
            {
                "service_quota": quota.pk,
                "funding_source": "",
                "service": "",
                "staff_member": self.staff.pk,
                "allocated_sessions": 10,
                "session_pay_amount": "500.00",
                "starts_on": "2026-01-01",
                "ends_on": "2026-12-31",
                "note": "На весь план",
            },
        )

        allocation.refresh_from_db()
        self.assertEqual(allocation.allocated_sessions, 10)
        self.assertEqual(allocation.session_pay_amount, Decimal("500.00"))
        self.assertRedirects(
            response,
            (
                f"{reverse('grant_report')}?funding={self.funding_grant.pk}"
                "&date_from=2026-01-01&date_to=2026-12-31"
            ),
        )

    def test_funding_staff_allocation_delete_removes_row(self):
        allocation = FundingStaffAllocation.objects.create(
            funding_source=self.funding_grant,
            service=self.service2,
            staff_member=self.staff,
            allocated_sessions=5,
            starts_on="2026-01-01",
            ends_on="2026-12-31",
        )

        response = self.client.post(
            reverse("funding_staff_allocation_delete", args=[allocation.pk])
        )

        self.assertFalse(FundingStaffAllocation.objects.filter(pk=allocation.pk).exists())
        self.assertRedirects(
            response,
            (
                f"{reverse('grant_report')}?funding={self.funding_grant.pk}"
                "&date_from=2026-01-01&date_to=2026-12-31"
            ),
        )

    def test_grant_report_filtered(self):
        response = self.client.get(
            reverse("grant_report"),
            {
                "funding": self.funding_grant.pk,
                "date_from": "2020-01-01",
                "date_to": "2099-01-01",
            },
        )
        self.assertEqual(response.status_code, 200)
        self.assertIn("report", response.context)
        self.assertIsNotNone(response.context["report"])
        self.assertContains(response, "Текущий остаток")
        self.assertContains(response, "Квоты не заданы")

    def test_grant_report_warns_about_unallocated_quota(self):
        FundingServiceQuota.objects.create(
            funding_source=self.funding_grant,
            service=self.service2,
            planned_sessions=20,
        )

        response = self.client.get(
            reverse("grant_report"),
            {
                "funding": self.funding_grant.pk,
                "date_from": "2020-01-01",
                "date_to": "2099-01-01",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Нераспределенная квота")
        self.assertContains(response, "20 занятий")

    def test_grant_report_links_existing_quota_to_allocation(self):
        quota = FundingServiceQuota.objects.create(
            funding_source=self.funding_grant,
            service=self.service2,
            planned_sessions=20,
        )

        response = self.client.get(
            reverse("grant_report"),
            {
                "funding": self.funding_grant.pk,
                "date_from": "2020-01-01",
                "date_to": "2099-01-01",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            f"{reverse('funding_staff_allocation_create')}?quota={quota.pk}",
        )

    def test_grant_report_tables_have_mobile_labels(self):
        quota = FundingServiceQuota.objects.create(
            funding_source=self.funding_grant,
            service=self.service2,
            planned_sessions=20,
        )
        FundingStaffAllocation.objects.create(
            service_quota=quota,
            funding_source=self.funding_grant,
            service=self.service2,
            staff_member=self.staff,
            allocated_sessions=5,
            session_pay_amount=Decimal("400"),
        )
        GrantRecipientAllocation.objects.create(
            funding_source=self.funding_grant,
            child=self.child,
            service=self.service2,
            allocated_sessions=7,
            balance_account=self.account_grant,
        )
        Certificate.objects.create(
            child=self.child,
            certificate_type=Certificate.CertificateType.SPONSOR,
            number="CERT-1",
            total_amount=Decimal("10000"),
            remaining_amount=Decimal("7500"),
        )
        Discount.objects.create(
            child=self.child,
            service=self.service2,
            percentage=Decimal("10"),
            note="Грантовая скидка",
        )

        response = self.client.get(
            reverse("grant_report"),
            {
                "funding": self.funding_grant.pk,
                "date_from": "2020-01-01",
                "date_to": "2099-01-01",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "grant-report-table")
        self.assertContains(response, "grant-balance-table")
        self.assertContains(response, "grant-quota-table")
        self.assertContains(response, "grant-recipient-allocation-table")
        self.assertContains(response, "grant-certificates-table")
        self.assertContains(response, "grant-discounts-table")
        self.assertContains(response, 'data-label="Текущий"')
        self.assertContains(response, 'data-label="Стоимость специалисту"')
        self.assertContains(response, 'data-label="Остаток счета"')
        self.assertContains(response, 'data-label="Полная сумма"')
        self.assertContains(response, 'data-label="Комментарий"')

    def test_grant_report_csv_export(self):
        quota = FundingServiceQuota.objects.create(
            funding_source=self.funding_grant,
            service=self.service2,
            planned_sessions=20,
        )
        FundingStaffAllocation.objects.create(
            service_quota=quota,
            funding_source=self.funding_grant,
            service=self.service2,
            staff_member=self.staff,
            allocated_sessions=5,
            session_pay_amount=Decimal("400"),
        )
        GrantRecipientAllocation.objects.create(
            funding_source=self.funding_grant,
            child=self.child,
            service=self.service2,
            allocated_sessions=7,
            balance_account=self.account_grant,
        )
        response = self.client.get(
            reverse("grant_report"),
            {
                "funding": self.funding_grant.pk,
                "date_from": "2020-01-01",
                "date_to": "2099-01-01",
                "csv": "1",
            },
        )
        self.assertEqual(response.status_code, 200)
        self.assertIn("text/csv", response["Content-Type"])
        body = response.content.decode("utf-8")
        self.assertIn("Счёт;Начальный", body)
        self.assertIn("Квоты по услугам", body)
        self.assertIn("Дефектолог;20;5;0;20", body)
        self.assertIn("Выделения получателям", body)
        self.assertIn("Иванов Ваня;Дефектолог;7", body)

    def test_grant_report_with_funding_pk(self):
        response = self.client.get(
            reverse("grant_report_funding", args=[self.funding_grant.pk]),
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context["funding_id"], self.funding_grant.pk)


class StaffMassRescheduleViewTests(NewViewsTestBase):
    def test_get_renders_form(self):
        today = timezone.localdate() + timedelta(days=2)
        start = _local_dt(today, time(10, 0))
        Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            billing_account=self.account,
            status=Appointment.Status.CONFIRMED,
        )

        response = self.client.get(
            reverse("staff_mass_reschedule", args=[self.staff.pk]),
            {"date_from": today.isoformat(), "date_to": today.isoformat()},
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn("staff", response.context)
        self.assertIn("mass_reschedule_summary_items", response.context)
        self.assertIn("mass_reschedule_next_action", response.context)
        self.assertEqual(
            response.context["mass_reschedule_next_action"]["title"],
            "Сохранить план отсутствия",
        )
        self.assertContains(response, "Контроль массового переноса")
        self.assertContains(response, 'id="mass-reschedule-form"')
        self.assertContains(response, 'hx-target="#mass-reschedule-result"')
        self.assertContains(response, "Сохранить план отсутствия")
        self.assertContains(response, "Отменить и разослать уведомления")

    def test_post_requires_reason(self):
        today = timezone.localdate()
        response = self.client.post(
            reverse("staff_mass_reschedule", args=[self.staff.pk]),
            {
                "date_from": today.isoformat(),
                "date_to": (today + timedelta(days=7)).isoformat(),
                "reason": "",
            },
        )
        self.assertEqual(response.status_code, 200)
        msgs = [str(m) for m in response.context["messages"]]
        self.assertTrue(any("причин" in m.lower() for m in msgs))

    def test_post_cancels_appointments(self):
        today = timezone.localdate() + timedelta(days=5)
        start = _local_dt(today, time(10, 0))
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            billing_account=self.account,
            status=Appointment.Status.CONFIRMED,
        )
        response = self.client.post(
            reverse("staff_mass_reschedule", args=[self.staff.pk]),
            {
                "date_from": today.isoformat(),
                "date_to": (today + timedelta(days=7)).isoformat(),
                "reason": "Болезнь",
            },
        )
        self.assertEqual(response.status_code, 302)
        appt.refresh_from_db()
        self.assertEqual(appt.status, Appointment.Status.CANCELLED)

    def test_post_create_plan_does_not_cancel_appointments(self):
        today = timezone.localdate() + timedelta(days=5)
        start = _local_dt(today, time(10, 0))
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            billing_account=self.account,
            status=Appointment.Status.CONFIRMED,
        )

        response = self.client.post(
            reverse("staff_mass_reschedule", args=[self.staff.pk]),
            {
                "action": "create_plan",
                "date_from": today.isoformat(),
                "date_to": today.isoformat(),
                "reason": "Больничный",
            },
        )

        plan = AppointmentReschedulePlan.objects.get(
            plan_type=AppointmentReschedulePlan.PlanType.STAFF_ABSENCE,
            staff_member=self.staff,
        )
        self.assertRedirects(response, reverse("appointment_reschedule_plan_detail", args=[plan.pk]))
        self.assertEqual(plan.steps.count(), 1)
        appt.refresh_from_db()
        self.assertEqual(appt.status, Appointment.Status.CONFIRMED)

    def test_htmx_create_plan_returns_redirect_header(self):
        today = timezone.localdate() + timedelta(days=5)
        start = _local_dt(today, time(10, 0))
        Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            billing_account=self.account,
            status=Appointment.Status.CONFIRMED,
        )

        response = self.client.post(
            reverse("staff_mass_reschedule", args=[self.staff.pk]),
            {
                "action": "create_plan",
                "date_from": today.isoformat(),
                "date_to": today.isoformat(),
                "reason": "Больничный",
            },
            HTTP_HX_REQUEST="true",
        )

        plan = AppointmentReschedulePlan.objects.get(
            plan_type=AppointmentReschedulePlan.PlanType.STAFF_ABSENCE,
            staff_member=self.staff,
        )
        self.assertEqual(response.status_code, 204)
        self.assertEqual(
            response.headers["HX-Redirect"],
            reverse("appointment_reschedule_plan_detail", args=[plan.pk]),
        )


class RecommendationViewTests(NewViewsTestBase):
    def test_list_renders(self):
        Recommendation.objects.create(
            child=self.child,
            staff_member=self.staff,
            title="Домашнее задание",
            body="Повторять упражнения.",
        )

        response = self.client.get(reverse("recommendation_list"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("recommendations", response.context)
        self.assertIn("recommendation_summary_items", response.context)
        self.assertIn("recommendation_control_items", response.context)
        self.assertIn("recommendation_next_action", response.context)
        self.assertContains(response, "Контроль рекомендаций")
        self.assertContains(response, "Следующий шаг")
        self.assertContains(response, 'id="recommendation-list"')
        self.assertContains(response, "recommendation-table")
        self.assertContains(response, 'data-label="Рекомендация"')
        self.assertContains(response, 'data-label="Действия"')

    def test_list_filtered_by_child(self):
        other_child = Child.objects.create(last_name="Другой", first_name="Получатель")
        Recommendation.objects.create(
            child=self.child,
            staff_member=self.staff,
            title="Рекомендация нужного получателя",
            body="Текст",
        )
        Recommendation.objects.create(
            child=other_child,
            staff_member=self.staff,
            title="Рекомендация другого получателя",
            body="Текст",
        )

        response = self.client.get(reverse("recommendation_list"), {"child_id": self.child.pk})

        self.assertEqual(response.status_code, 200)
        self.assertEqual([item.child_id for item in response.context["recommendations"]], [self.child.pk])
        self.assertContains(response, "Рекомендация нужного получателя")
        self.assertNotContains(response, "Рекомендация другого получателя")
        self.assertEqual(
            response.context["recommendation_create_url"],
            reverse("recommendation_create_for_child", args=[self.child.pk]),
        )

    def test_create_get(self):
        response = self.client.get(reverse("recommendation_create"))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Контроль записи")
        self.assertContains(response, 'id="recommendation-form"')

    def test_create_prefills_child(self):
        response = self.client.get(reverse("recommendation_create_for_child", args=[self.child.pk]))
        self.assertEqual(response.status_code, 200)

    def test_create_post(self):
        response = self.client.post(
            reverse("recommendation_create"),
            {
                "child": self.child.pk,
                "staff_member": self.staff.pk,
                "category": Recommendation.Category.HOME_TASK,
                "title": "Занятия 2 раза в неделю",
                "body": "Продолжать логопедические занятия дважды в неделю.",
            },
        )
        if response.status_code != 302:
            form = response.context.get("form")
            print("Form errors:", form.errors if form else "no form")
        self.assertEqual(response.status_code, 302)
        self.assertTrue(Recommendation.objects.filter(child=self.child).exists())

    def test_acknowledge(self):
        rec = Recommendation.objects.create(
            child=self.child,
            staff_member=self.staff,
            title="Тест",
            body="Тест",
        )
        response = self.client.post(reverse("recommendation_acknowledge", args=[rec.pk]))
        self.assertEqual(response.status_code, 302)
        rec.refresh_from_db()
        self.assertTrue(rec.is_acknowledged)
        self.assertIsNotNone(rec.acknowledged_at)


class DocumentViewTests(NewViewsTestBase):
    def test_list_renders(self):
        Document.objects.create(
            child=self.child,
            category=Document.Category.CONTRACT,
            title="Договор для проверки",
            file="documents/check.txt",
        )

        response = self.client.get(reverse("document_list"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("document_summary_items", response.context)
        self.assertIn("document_control_items", response.context)
        self.assertIn("document_next_action", response.context)
        self.assertContains(response, "Контроль документов")
        self.assertContains(response, "Следующий шаг")
        self.assertContains(response, 'id="document-list"')
        self.assertContains(response, "document-table")
        self.assertContains(response, 'data-label="Название"')
        self.assertContains(response, 'data-label="Действует до"')

    def test_list_filtered_by_child(self):
        other_child = Child.objects.create(last_name="Документы", first_name="Другой")
        Document.objects.create(
            child=self.child,
            category=Document.Category.CONTRACT,
            title="Договор нужного получателя",
            file="documents/contract.txt",
        )
        Document.objects.create(
            child=other_child,
            category=Document.Category.OTHER,
            title="Документ другого получателя",
            file="documents/other.txt",
        )

        response = self.client.get(reverse("document_list"), {"child_id": self.child.pk})

        self.assertEqual(response.status_code, 200)
        self.assertEqual([item.child_id for item in response.context["documents"]], [self.child.pk])
        self.assertContains(response, "Договор нужного получателя")
        self.assertNotContains(response, "Документ другого получателя")
        self.assertEqual(
            response.context["document_create_url"],
            reverse("document_create_for_child", args=[self.child.pk]),
        )

    def test_create_get(self):
        response = self.client.get(reverse("document_create"))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Контроль записи")
        self.assertContains(response, 'id="document-form"')

    def test_create_prefills_child(self):
        response = self.client.get(reverse("document_create_for_child", args=[self.child.pk]))
        self.assertEqual(response.status_code, 200)

    def test_create_post(self):
        from django.core.files.uploadedfile import SimpleUploadedFile

        response = self.client.post(
            reverse("document_create"),
            {
                "target_type": Document.TargetType.RECIPIENT,
                "child": self.child.pk,
                "counterparty": "",
                "category": Document.Category.OTHER,
                "title": "Справка",
                "issued_on": "2024-01-01",
                "expires_on": "2025-01-01",
                "file": SimpleUploadedFile("test.txt", b"hello", content_type="text/plain"),
                "note": "",
            },
        )
        self.assertEqual(response.status_code, 302)
        self.assertTrue(Document.objects.filter(child=self.child, title="Справка").exists())

    def test_create_counterparty_document_without_child(self):
        from django.core.files.uploadedfile import SimpleUploadedFile

        counterparty = Counterparty.objects.create(name="Фонд документов")

        response = self.client.post(
            reverse("document_create"),
            {
                "target_type": Document.TargetType.COUNTERPARTY,
                "child": "",
                "counterparty": counterparty.pk,
                "category": Document.Category.CONTRACT,
                "title": "Договор фонда",
                "issued_on": "2024-01-01",
                "expires_on": "",
                "file": SimpleUploadedFile("fund.txt", b"hello", content_type="text/plain"),
                "note": "",
            },
        )

        self.assertRedirects(response, reverse("document_list"))
        document = Document.objects.get(title="Договор фонда")
        self.assertIsNone(document.child_id)
        self.assertEqual(document.counterparty, counterparty)
        self.assertEqual(document.target_type, Document.TargetType.COUNTERPARTY)

    def test_create_recipient_document_requires_child(self):
        from django.core.files.uploadedfile import SimpleUploadedFile

        response = self.client.post(
            reverse("document_create"),
            {
                "target_type": Document.TargetType.RECIPIENT,
                "child": "",
                "counterparty": "",
                "category": Document.Category.OTHER,
                "title": "Без получателя",
                "issued_on": "",
                "expires_on": "",
                "file": SimpleUploadedFile("missing.txt", b"hello", content_type="text/plain"),
                "note": "",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertFalse(Document.objects.filter(title="Без получателя").exists())
        self.assertContains(response, "Для документа получателя выберите получателя")


class ConsentViewTests(NewViewsTestBase):
    def test_list_renders(self):
        Consent.objects.create(
            child=self.child,
            consent_type=Consent.ConsentType.PERSONAL_DATA,
            signed_on=timezone.localdate(),
        )

        response = self.client.get(reverse("consent_list"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("consent_summary_items", response.context)
        self.assertIn("consent_control_items", response.context)
        self.assertIn("consent_next_action", response.context)
        self.assertContains(response, "Контроль согласий")
        self.assertContains(response, "Следующий шаг")
        self.assertContains(response, 'id="consent-list"')
        self.assertContains(response, "consent-table")
        self.assertContains(response, 'data-label="Подписано"')
        self.assertContains(response, 'data-label="Документ"')

    def test_list_filtered_by_child(self):
        other_child = Child.objects.create(last_name="Согласия", first_name="Другой")
        Consent.objects.create(
            child=self.child,
            consent_type=Consent.ConsentType.PERSONAL_DATA,
            signed_on=timezone.localdate(),
        )
        Consent.objects.create(
            child=other_child,
            consent_type=Consent.ConsentType.PHOTO_VIDEO,
            signed_on=timezone.localdate(),
        )

        response = self.client.get(reverse("consent_list"), {"child_id": self.child.pk})

        self.assertEqual(response.status_code, 200)
        self.assertEqual([item.child_id for item in response.context["consents"]], [self.child.pk])
        self.assertContains(response, self.child.full_name)
        self.assertNotContains(response, other_child.full_name)
        self.assertEqual(
            response.context["consent_create_url"],
            reverse("consent_create_for_child", args=[self.child.pk]),
        )

    def test_create_get(self):
        response = self.client.get(reverse("consent_create"))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Контроль записи")
        self.assertContains(response, 'id="consent-form"')

    def test_create_prefills_child(self):
        response = self.client.get(reverse("consent_create_for_child", args=[self.child.pk]))
        self.assertEqual(response.status_code, 200)

    def test_create_post(self):
        response = self.client.post(
            reverse("consent_create"),
            {
                "child": self.child.pk,
                "consent_type": Consent.ConsentType.PERSONAL_DATA,
                "signed_on": "2024-06-01",
                "note": "",
            },
        )
        self.assertEqual(response.status_code, 302)
        self.assertTrue(Consent.objects.filter(child=self.child).exists())


class PaymentViewTests(NewViewsTestBase):
    def test_create_get(self):
        response = self.client.get(reverse("payment_create"))
        self.assertEqual(response.status_code, 200)
        self.assertIn("payment_summary_items", response.context)
        self.assertIn("payment_next_action", response.context)
        self.assertIsNone(response.context["selected_account"])
        self.assertEqual(response.context["payment_next_action"]["title"], "Выбрать счет")
        self.assertContains(response, "Контроль платежа")
        self.assertContains(response, 'id="payment-form"')

    def test_create_prefills_account(self):
        response = self.client.get(reverse("payment_create_for_account", args=[self.account.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context["selected_account"], self.account)
        self.assertEqual(response.context["payment_next_action"]["title"], "Проверить сумму")
        self.assertContains(response, self.child.full_name)

    def test_create_post(self):
        response = self.client.post(
            reverse("payment_create"),
            {
                "balance_account": self.account.pk,
                "amount": "5",
                "method": Payment.Method.CASH,
                "paid_at": timezone.localdate().isoformat(),
                "reference": "REF-1",
                "comment": "Тест",
            },
        )
        if response.status_code != 302:
            form = response.context.get("form")
            print("Form errors:", form.errors if form else "no form")
            print("Response:", response.content[:500].decode("utf-8", errors="replace"))
        self.assertEqual(response.status_code, 302)
        self.assertTrue(
            Payment.objects.filter(balance_account=self.account, amount=Decimal("5")).exists()
        )
        self.account.refresh_from_db()
        self.assertEqual(self.account.current_balance, Decimal("15"))

    def test_create_post_invalid_amount(self):
        response = self.client.post(
            reverse("payment_create"),
            {
                "balance_account": self.account.pk,
                "amount": "-5",
                "method": Payment.Method.CASH,
                "paid_at": timezone.localdate().isoformat(),
            },
        )
        self.assertEqual(response.status_code, 200)
        self.assertFalse(Payment.objects.filter(amount=Decimal("-5")).exists())


class BalancesViewTests(NewViewsTestBase):
    def test_balances_renders(self):
        response = self.client.get(reverse("balances"))
        self.assertEqual(response.status_code, 200)
        self.assertIn("balance_summary_items", response.context)
        self.assertIn("balance_next_action", response.context)
        self.assertContains(response, "Контроль балансов")
        self.assertContains(response, 'id="balance-list"')
        self.assertContains(response, reverse("payment_create_for_account", args=[self.account.pk]))
        self.assertContains(response, "balance-table")
        self.assertContains(response, 'data-label="Остаток"')
        self.assertContains(response, 'data-label="Действия"')

    def test_balance_account_forms_show_operator_control(self):
        cases = [
            (reverse("balance_account_create"), "Контроль балансового счета", "Единица учета"),
            (
                reverse("balance_account_edit", args=[self.account.pk]),
                "Контроль балансового счета",
                "Остаток и ledger",
            ),
        ]
        for url, title, control_item in cases:
            with self.subTest(url=url):
                response = self.client.get(url)
                self.assertEqual(response.status_code, 200)
                self.assertIn("object_form_control_items", response.context)
                self.assertContains(response, 'class="object-form-layout"')
                self.assertContains(response, title)
                self.assertContains(response, control_item)

    def test_balance_account_create(self):
        response = self.client.post(
            reverse("balance_account_create"),
            {
                "child": self.child.pk,
                "funding_source": self.funding.pk,
                "unit": BalanceAccount.Unit.SESSIONS,
                "service_scope": BalanceAccount.ServiceScope.SPECIFIC_SERVICE,
                "service": self.service.pk,
                "initial_amount": "3",
                "valid_from": "",
                "valid_until": "",
                "status": BalanceAccount.Status.ACTIVE,
                "notes": "",
            },
        )
        self.assertEqual(response.status_code, 302)

    def test_balance_account_edit(self):
        response = self.client.get(reverse("balance_account_edit", args=[self.account.pk]))
        self.assertEqual(response.status_code, 200)

    def test_balance_account_delete_blocks_group_participant_account(self):
        other_child = Child.objects.create(last_name="Смирнов", first_name="Илья")
        start = _local_dt(timezone.localdate() + timedelta(days=3), time(11, 0))
        appointment = Appointment.objects.create(
            child=other_child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            session_type=Appointment.SessionType.GROUP,
        )
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=self.child,
            billing_account=self.account,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )

        response = self.client.post(reverse("balance_account_delete", args=[self.account.pk]))

        self.assertRedirects(response, reverse("balances"))
        self.assertTrue(BalanceAccount.objects.filter(pk=self.account.pk).exists())

    def test_balance_account_delete_blocks_inactive_participant_account(self):
        other_child = Child.objects.create(last_name="История", first_name="Счет")
        start = _local_dt(timezone.localdate() - timedelta(days=3), time(11, 0))
        appointment = Appointment.objects.create(
            child=other_child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            session_type=Appointment.SessionType.GROUP,
            status=Appointment.Status.CANCELLED,
        )
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=self.child,
            billing_account=self.account,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=Appointment.Status.CANCELLED,
        )

        response = self.client.post(reverse("balance_account_delete", args=[self.account.pk]))

        self.assertRedirects(response, reverse("balances"))
        self.assertTrue(BalanceAccount.objects.filter(pk=self.account.pk).exists())

    def test_balance_account_delete_blocks_ledger_history(self):
        LedgerEntry.objects.create(
            account=self.account,
            entry_type=LedgerEntry.EntryType.CORRECTION,
            amount=Decimal("1"),
            reason="Историческая корректировка",
        )

        response = self.client.post(reverse("balance_account_delete", args=[self.account.pk]))

        self.assertRedirects(response, reverse("balances"))
        self.assertTrue(BalanceAccount.objects.filter(pk=self.account.pk).exists())
        self.assertTrue(LedgerEntry.objects.filter(account=self.account).exists())


class RoomViewTests(NewViewsTestBase):
    def test_room_list_renders(self):
        response = self.client.get(reverse("room_list"))
        self.assertEqual(response.status_code, 200)
        self.assertIn("room_summary_items", response.context)
        self.assertIn("room_next_action", response.context)
        self.assertEqual(
            response.context["room_next_action"]["title"],
            "Настроить групповой кабинет",
        )
        self.assertContains(response, "Следующее действие")
        self.assertContains(response, "Кабинеты и вместимость")
        self.assertContains(response, 'id="room-list"')
        self.assertContains(response, reverse("room_create"))
        self.assertContains(response, "directory-table")
        self.assertContains(response, 'data-label="Специалисты"')
        self.assertContains(response, 'data-label="Действия"')

    def test_object_form_shows_save_control_panel(self):
        response = self.client.get(reverse("room_create"))
        self.assertEqual(response.status_code, 200)
        self.assertIn("object_form_control_items", response.context)
        self.assertContains(response, 'class="object-form-layout"')
        self.assertContains(response, "Контроль кабинета")
        self.assertContains(response, "Вместимость специалистов")
        self.assertContains(response, "Разовое исключение")
        self.assertContains(response, reverse("room_list"))

    def test_room_create(self):
        response = self.client.post(
            reverse("room_create"),
            {
                "name": "Большой зал",
                "room_type": Room.RoomType.GROUP,
                "capacity": "12",
                "limit_staff_count": "on",
                "max_staff_count": "2",
                "limit_recipient_count": "on",
                "max_recipient_count": "8",
                "allow_group_sessions": "on",
                "is_active": "on",
                "color": "#ff6600",
            },
        )

        room = Room.objects.get(name="Большой зал")
        self.assertRedirects(response, reverse("room_edit", args=[room.pk]))
        self.assertEqual(room.max_staff_count, 2)
        self.assertEqual(room.max_recipient_count, 8)
        self.assertTrue(room.allow_group_sessions)

    def test_room_edit_updates_limits(self):
        response = self.client.post(
            reverse("room_edit", args=[self.room.pk]),
            {
                "name": self.room.name,
                "room_type": Room.RoomType.GROUP,
                "capacity": "10",
                "max_staff_count": "3",
                "limit_recipient_count": "on",
                "max_recipient_count": "6",
                "allow_group_sessions": "on",
                "is_active": "on",
                "color": "#00aa00",
            },
        )

        self.assertRedirects(response, reverse("room_list"))
        self.room.refresh_from_db()
        self.assertFalse(self.room.limit_staff_count)
        self.assertEqual(self.room.max_staff_count, 3)
        self.assertTrue(self.room.limit_recipient_count)
        self.assertEqual(self.room.max_recipient_count, 6)
        self.assertTrue(self.room.allow_group_sessions)


class ServiceViewTests(NewViewsTestBase):
    def test_service_list_renders(self):
        response = self.client.get(reverse("service_list"))
        self.assertEqual(response.status_code, 200)
        self.assertIn("service_summary_items", response.context)
        self.assertIn("service_next_action", response.context)
        self.assertEqual(response.context["service_next_action"]["title"], "Услуги готовы к работе")
        self.assertContains(response, self.service.name)
        self.assertContains(response, "Следующее действие")
        self.assertContains(response, "Услуги и направления")
        self.assertContains(response, 'id="service-list"')
        self.assertContains(response, reverse("service_create"))
        self.assertContains(response, "directory-table")
        self.assertContains(response, 'data-label="Длительность"')
        self.assertContains(response, 'data-label="Действия"')

    def test_service_forms_show_operator_control(self):
        cases = [
            (reverse("service_create"), "Контроль услуги", "Длительность"),
            (reverse("service_edit", args=[self.service.pk]), "Контроль услуги", "Цена"),
        ]
        for url, title, control_item in cases:
            with self.subTest(url=url):
                response = self.client.get(url)
                self.assertEqual(response.status_code, 200)
                self.assertIn("object_form_control_items", response.context)
                self.assertContains(response, 'class="object-form-layout"')
                self.assertContains(response, title)
                self.assertContains(response, control_item)

    def test_service_create(self):
        response = self.client.post(
            reverse("service_create"),
            {
                "name": "Психолог",
                "code": "PSY",
                "category": Service.Category.OTHER,
                "default_duration_minutes": 50,
                "default_price": "2200.00",
                "is_active": "on",
                "color": "#336699",
            },
        )

        service = Service.objects.get(code="PSY")
        self.assertRedirects(response, reverse("service_edit", args=[service.pk]))
        self.assertEqual(service.default_duration_minutes, 50)
        self.assertEqual(service.default_price, Decimal("2200.00"))

    def test_service_edit(self):
        response = self.client.post(
            reverse("service_edit", args=[self.service.pk]),
            {
                "name": "Логопед индивидуально",
                "code": self.service.code,
                "category": Service.Category.SPEECH,
                "default_duration_minutes": 40,
                "default_price": "1700.00",
                "is_active": "on",
                "color": "#112233",
            },
        )

        self.assertRedirects(response, reverse("service_list"))
        self.service.refresh_from_db()
        self.assertEqual(self.service.name, "Логопед индивидуально")
        self.assertEqual(self.service.default_duration_minutes, 40)

    def test_service_archive_and_restore(self):
        archive_response = self.client.post(reverse("service_archive", args=[self.service.pk]))
        self.assertRedirects(archive_response, reverse("service_list"))
        self.service.refresh_from_db()
        self.assertIsNotNone(self.service.archived_at)
        self.assertFalse(Service.objects.filter(pk=self.service.pk).exists())

        restore_response = self.client.post(reverse("service_restore", args=[self.service.pk]))
        self.assertRedirects(restore_response, reverse("service_list"))
        self.service.refresh_from_db()
        self.assertIsNone(self.service.archived_at)


class StaffMemberViewTests(NewViewsTestBase):
    def test_staff_member_list_renders(self):
        response = self.client.get(reverse("staff_member_list"))
        self.assertEqual(response.status_code, 200)
        self.assertIn("staff_summary_items", response.context)
        self.assertIn("staff_next_action", response.context)
        self.assertEqual(
            response.context["staff_next_action"]["title"],
            "Специалисты готовы к работе",
        )
        self.assertContains(response, self.staff.full_name)
        self.assertContains(response, "Следующее действие")
        self.assertContains(response, "Профили специалистов")
        self.assertContains(response, 'id="staff-member-list"')
        self.assertContains(response, reverse("staff_member_create"))
        self.assertContains(response, reverse("staff_timesheet", args=[self.staff.pk]))
        self.assertContains(response, reverse("staff_mass_reschedule", args=[self.staff.pk]))
        self.assertContains(response, "directory-table")
        self.assertContains(response, 'data-label="Пользователь"')
        self.assertContains(response, 'data-label="Действия"')

    def test_staff_member_forms_show_operator_control(self):
        cases = [
            (reverse("staff_member_create"), "Контроль специалиста", "Учетная запись"),
            (
                reverse("staff_member_edit", args=[self.staff.pk]),
                "Контроль специалиста",
                "Мобильный кабинет",
            ),
        ]
        for url, title, control_item in cases:
            with self.subTest(url=url):
                response = self.client.get(url)
                self.assertEqual(response.status_code, 200)
                self.assertIn("object_form_control_items", response.context)
                self.assertContains(response, 'class="object-form-layout"')
                self.assertContains(response, title)
                self.assertContains(response, control_item)

    def test_staff_member_create(self):
        response = self.client.post(
            reverse("staff_member_create"),
            {
                "user": "",
                "full_name": "Новый специалист",
                "specializations": "Психолог",
                "phone": "+7 900 000-20-01",
                "email": "new-staff@example.local",
                "status": StaffMember.Status.ACTIVE,
                "color": "#123456",
                "can_use_mobile": "on",
            },
        )

        staff = StaffMember.objects.get(full_name="Новый специалист")
        self.assertRedirects(response, reverse("staff_member_edit", args=[staff.pk]))
        self.assertEqual(staff.specializations, "Психолог")
        self.assertTrue(staff.can_use_mobile)

    def test_staff_member_edit_can_bind_user(self):
        user = User.objects.create_user("staff2", password="x")

        response = self.client.post(
            reverse("staff_member_edit", args=[self.staff.pk]),
            {
                "user": user.pk,
                "full_name": "Иванова Наталья",
                "specializations": "Логопед, дефектолог",
                "phone": "+7 900 000-20-02",
                "email": "staff2@example.local",
                "status": StaffMember.Status.ACTIVE,
                "color": "#654321",
                "can_use_mobile": "on",
            },
        )

        self.assertRedirects(response, reverse("staff_member_list"))
        self.staff.refresh_from_db()
        self.assertEqual(self.staff.user, user)
        self.assertEqual(self.staff.full_name, "Иванова Наталья")

    def test_staff_member_archive_and_restore(self):
        archive_response = self.client.post(reverse("staff_member_archive", args=[self.staff.pk]))
        self.assertRedirects(archive_response, reverse("staff_member_list"))
        self.staff.refresh_from_db()
        self.assertIsNotNone(self.staff.archived_at)
        self.assertFalse(StaffMember.objects.filter(pk=self.staff.pk).exists())

        restore_response = self.client.post(reverse("staff_member_restore", args=[self.staff.pk]))
        self.assertRedirects(restore_response, reverse("staff_member_list"))
        self.staff.refresh_from_db()
        self.assertIsNone(self.staff.archived_at)


class FundingSourceViewTests(NewViewsTestBase):
    def test_funding_source_list_renders(self):
        response = self.client.get(reverse("funding_source_list"))
        self.assertEqual(response.status_code, 200)
        self.assertIn("funding_summary_items", response.context)
        self.assertIn("funding_next_action", response.context)
        self.assertEqual(
            response.context["funding_next_action"]["title"],
            "Проверить сроки грантов и фондов",
        )
        self.assertContains(response, self.funding_grant.name)
        self.assertContains(response, "Следующее действие")
        self.assertContains(response, "Источники и правила средств")
        self.assertContains(response, 'id="funding-source-list"')
        self.assertContains(response, reverse("funding_source_create"))
        self.assertContains(response, f"{reverse('grant_report')}?funding={self.funding_grant.pk}")
        self.assertContains(response, "directory-table")
        self.assertContains(response, 'data-label="Передача средств"')
        self.assertContains(response, 'data-label="Действия"')

    def test_funding_source_list_prioritizes_all_archived_state(self):
        self.funding.archive()
        self.funding_grant.archive()

        response = self.client.get(reverse("funding_source_list"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.context["funding_next_action"]["title"],
            "Создать или восстановить источник",
        )
        self.assertContains(response, "Все источники в архиве")

    def test_funding_source_forms_show_operator_control(self):
        cases = [
            (reverse("funding_source_create"), "Контроль источника", "Тип источника"),
            (
                reverse("funding_source_edit", args=[self.funding_grant.pk]),
                "Контроль источника",
                "Перенос средств",
            ),
        ]
        for url, title, control_item in cases:
            with self.subTest(url=url):
                response = self.client.get(url)
                self.assertEqual(response.status_code, 200)
                self.assertIn("object_form_control_items", response.context)
                self.assertContains(response, 'class="object-form-layout"')
                self.assertContains(response, title)
                self.assertContains(response, control_item)

    def test_funding_source_create(self):
        response = self.client.post(
            reverse("funding_source_create"),
            {
                "name": "Фонд помощи",
                "source_type": FundingSource.SourceType.CHARITY_FUND,
                "starts_on": "2026-01-01",
                "ends_on": "2026-12-31",
                "transfer_policy": FundingSource.TransferPolicy.WITHIN_CHILD,
                "notes": "Годовая программа",
            },
        )

        source = FundingSource.objects.get(name="Фонд помощи")
        self.assertRedirects(response, reverse("funding_source_edit", args=[source.pk]))
        self.assertEqual(source.source_type, FundingSource.SourceType.CHARITY_FUND)
        self.assertEqual(source.transfer_policy, FundingSource.TransferPolicy.WITHIN_CHILD)

    def test_funding_source_edit(self):
        response = self.client.post(
            reverse("funding_source_edit", args=[self.funding_grant.pk]),
            {
                "name": "Грант 2026",
                "source_type": FundingSource.SourceType.GRANT,
                "starts_on": "2026-01-01",
                "ends_on": "2026-12-31",
                "transfer_policy": FundingSource.TransferPolicy.BETWEEN_CHILDREN,
                "notes": "",
            },
        )

        self.assertRedirects(response, reverse("funding_source_list"))
        self.funding_grant.refresh_from_db()
        self.assertEqual(self.funding_grant.name, "Грант 2026")
        self.assertEqual(
            self.funding_grant.transfer_policy,
            FundingSource.TransferPolicy.BETWEEN_CHILDREN,
        )

    def test_funding_source_archive_and_restore(self):
        archive_response = self.client.post(
            reverse("funding_source_archive", args=[self.funding_grant.pk])
        )
        self.assertRedirects(archive_response, reverse("funding_source_list"))
        self.funding_grant.refresh_from_db()
        self.assertIsNotNone(self.funding_grant.archived_at)
        self.assertFalse(FundingSource.objects.filter(pk=self.funding_grant.pk).exists())

        restore_response = self.client.post(
            reverse("funding_source_restore", args=[self.funding_grant.pk])
        )
        self.assertRedirects(restore_response, reverse("funding_source_list"))
        self.funding_grant.refresh_from_db()
        self.assertIsNone(self.funding_grant.archived_at)


class ExpenseDirectoryViewTests(NewViewsTestBase):
    def _financial_counts(self):
        return {
            "balances": BalanceAccount.objects.count(),
            "ledger": LedgerEntry.objects.count(),
            "payments": Payment.objects.count(),
            "expenses": CenterExpense.objects.count(),
            "donation_contracts": DonationContract.objects.count(),
            "payroll": PayrollAccrual.objects.count(),
        }

    def test_expense_directory_list_renders_categories_and_counterparties(self):
        CenterExpense.objects.create(
            category=self.expense_category,
            counterparty=self.counterparty,
            title="Расход по справочнику",
            total_amount=Decimal("500.00"),
        )
        DonationContract.objects.create(
            counterparty=self.counterparty,
            funding_source=self.funding_grant,
            contract_type=DonationContract.ContractType.PROJECT,
            number="DIR-1",
        )

        response = self.client.get(reverse("expense_directory_list"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("directory_summary_items", response.context)
        self.assertIn("directory_next_action", response.context)
        self.assertContains(response, "Справочники расходов")
        self.assertContains(response, "Категории расходов")
        self.assertContains(response, "Контрагенты")
        self.assertContains(response, self.expense_category.name)
        self.assertContains(response, self.counterparty.name)
        self.assertContains(response, "Расходов: 1")
        self.assertContains(response, "Договоров: 1")
        self.assertContains(response, reverse("expense_category_create"))
        self.assertContains(response, reverse("counterparty_create"))
        self.assertContains(response, 'id="expense-directory-categories"')
        self.assertContains(response, 'id="expense-directory-counterparties"')

    def test_expense_directory_list_filters_archived_counterparties(self):
        self.counterparty.archive()

        archived_response = self.client.get(
            reverse("expense_directory_list"),
            {"counterparty_status": "archived"},
        )
        active_response = self.client.get(
            reverse("expense_directory_list"),
            {"counterparty_status": "active"},
        )

        self.assertContains(archived_response, self.counterparty.name)
        self.assertNotContains(active_response, self.counterparty.name)

    def test_expense_category_create_edit_and_toggle_without_financial_facts(self):
        counts_before = self._financial_counts()

        create_response = self.client.post(
            reverse("expense_category_create"),
            {
                "name": "Транспорт",
                "expense_type": CenterExpenseCategory.ExpenseType.OTHER,
                "is_active": "on",
                "sort_order": "30",
                "notes": "Такси и доставка",
            },
        )

        category = CenterExpenseCategory.objects.get(name="Транспорт")
        self.assertRedirects(create_response, reverse("expense_category_edit", args=[category.pk]))
        self.assertEqual(self._financial_counts(), counts_before)

        edit_response = self.client.post(
            reverse("expense_category_edit", args=[category.pk]),
            {
                "name": "Транспортные расходы",
                "expense_type": CenterExpenseCategory.ExpenseType.SERVICES,
                "sort_order": "35",
                "notes": "",
            },
        )

        self.assertRedirects(edit_response, reverse("expense_directory_list"))
        category.refresh_from_db()
        self.assertEqual(category.name, "Транспортные расходы")
        self.assertFalse(category.is_active)
        self.assertEqual(self._financial_counts(), counts_before)

        activate_response = self.client.post(
            reverse("expense_category_activate", args=[category.pk])
        )
        self.assertRedirects(activate_response, reverse("expense_directory_list"))
        category.refresh_from_db()
        self.assertTrue(category.is_active)

        deactivate_response = self.client.post(
            reverse("expense_category_deactivate", args=[category.pk])
        )
        self.assertRedirects(deactivate_response, reverse("expense_directory_list"))
        category.refresh_from_db()
        self.assertFalse(category.is_active)
        self.assertEqual(self._financial_counts(), counts_before)

    def test_expense_category_form_rejects_case_insensitive_duplicate(self):
        response = self.client.post(
            reverse("expense_category_create"),
            {
                "name": self.expense_category.name.upper(),
                "expense_type": CenterExpenseCategory.ExpenseType.HOUSEHOLD,
                "is_active": "on",
                "sort_order": "10",
                "notes": "",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Категория с таким названием уже есть.")

    def test_counterparty_create_edit_archive_restore_without_financial_facts(self):
        counts_before = self._financial_counts()

        create_response = self.client.post(
            reverse("counterparty_create"),
            {
                "name": "Благотворительный партнер",
                "counterparty_type": Counterparty.CounterpartyType.SPONSOR,
                "inn": "2500000000",
                "kpp": "",
                "ogrn": "",
                "legal_address": "",
                "postal_address": "",
                "bank_details": "",
                "contact_person": "Петров Петр",
                "phone": "+7 900 000-10-10",
                "email": "partner@example.local",
                "notes": "",
            },
        )

        counterparty = Counterparty.objects.get(name="Благотворительный партнер")
        self.assertRedirects(create_response, reverse("counterparty_edit", args=[counterparty.pk]))
        self.assertEqual(self._financial_counts(), counts_before)

        edit_response = self.client.post(
            reverse("counterparty_edit", args=[counterparty.pk]),
            {
                "name": "Благотворительный партнер 2026",
                "counterparty_type": Counterparty.CounterpartyType.FOUNDATION,
                "inn": "2500000001",
                "kpp": "",
                "ogrn": "",
                "legal_address": "Адрес",
                "postal_address": "",
                "bank_details": "Реквизиты",
                "contact_person": "",
                "phone": "",
                "email": "",
                "notes": "Обновлено",
            },
        )

        self.assertRedirects(edit_response, reverse("expense_directory_list"))
        counterparty.refresh_from_db()
        self.assertEqual(counterparty.name, "Благотворительный партнер 2026")
        self.assertEqual(counterparty.counterparty_type, Counterparty.CounterpartyType.FOUNDATION)
        self.assertEqual(self._financial_counts(), counts_before)

        archive_response = self.client.post(reverse("counterparty_archive", args=[counterparty.pk]))
        self.assertRedirects(archive_response, reverse("expense_directory_list"))
        counterparty.refresh_from_db()
        self.assertTrue(counterparty.is_archived)
        self.assertFalse(Counterparty.objects.filter(pk=counterparty.pk).exists())

        restore_response = self.client.post(reverse("counterparty_restore", args=[counterparty.pk]))
        self.assertRedirects(restore_response, reverse("expense_directory_list"))
        counterparty.refresh_from_db()
        self.assertFalse(counterparty.is_archived)
        self.assertTrue(Counterparty.objects.filter(pk=counterparty.pk).exists())
        self.assertEqual(self._financial_counts(), counts_before)

    def test_expense_directory_rejects_specialist(self):
        self.client.logout()
        self.client.force_login(self.staff_user)

        response = self.client.get(reverse("expense_directory_list"))

        self.assertEqual(response.status_code, 302)


class CenterLegalProfileViewTests(NewViewsTestBase):
    def test_center_legal_profile_get_renders_form(self):
        response = self.client.get(reverse("center_legal_profile"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Юридический профиль центра")
        self.assertContains(response, "Реквизиты центра")
        self.assertContains(response, "Один активный профиль")
        self.assertContains(response, 'name="full_name"')

    def test_center_legal_profile_post_creates_active_profile(self):
        response = self.client.post(
            reverse("center_legal_profile"),
            {
                "full_name": "Автономная некоммерческая организация Радость моя",
                "short_name": "АНО Радость моя",
                "director_full_name": "Иванов Иван Иванович",
                "director_short_name": "И. И. Иванов",
                "director_position": "Директор",
                "authority_basis": "Устава",
                "license_number": "Л-25-000001",
                "license_date": "2026-01-15",
                "license_authority": "Министерство здравоохранения",
                "ogrn": "1234567890123",
                "inn": "2500000000",
                "kpp": "250001001",
                "legal_address": "г. Владивосток, ул. Юридическая, 1",
                "location_address": "г. Владивосток, ул. Рабочая, 2",
                "phone": "+7 900 000-00-01",
                "email": "center@example.local",
                "site": "rm.example.local",
                "bank_name": "Тест Банк",
                "bank_bik": "040000001",
                "bank_account": "40703810000000000001",
                "bank_corr_account": "30101810000000000001",
                "is_active": "on",
                "notes": "",
            },
        )

        self.assertRedirects(response, reverse("center_legal_profile"))
        profile = CenterLegalProfile.objects.get()
        self.assertTrue(profile.is_active)
        self.assertEqual(profile.short_name, "АНО Радость моя")
        self.assertEqual(CenterLegalProfile.get_active(), profile)


class CenterExpenseViewTests(NewViewsTestBase):
    def _expense_post_data(
        self,
        *,
        title: str = "Покупка материалов",
        first_source=None,
        second_source=None,
        first_amount: str = "700.00",
        second_amount: str = "300.00",
    ):
        first_source = first_source or self.funding
        second_source = second_source or self.funding_grant
        return {
            "expense_date": timezone.localdate().isoformat(),
            "category": self.expense_category.pk,
            "title": title,
            "description": "Расход для занятий",
            "counterparty": self.counterparty.pk,
            "total_amount": "1000.00",
            "notes": "",
            "funding_splits-TOTAL_FORMS": "2",
            "funding_splits-INITIAL_FORMS": "0",
            "funding_splits-MIN_NUM_FORMS": "0",
            "funding_splits-MAX_NUM_FORMS": "1000",
            "funding_splits-0-funding_source": first_source.pk,
            "funding_splits-0-amount": first_amount,
            "funding_splits-0-notes": "",
            "funding_splits-1-funding_source": second_source.pk,
            "funding_splits-1-amount": second_amount,
            "funding_splits-1-notes": "",
        }

    def test_center_expense_list_renders(self):
        expense = CenterExpense.objects.create(
            category=self.expense_category,
            counterparty=self.counterparty,
            title="Канцтовары",
            total_amount=Decimal("500.00"),
        )
        ExpenseFundingSplit.objects.create(
            expense=expense,
            funding_source=self.funding,
            amount=Decimal("500.00"),
        )

        response = self.client.get(reverse("center_expense_list"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("expense_summary_items", response.context)
        self.assertIn("expense_next_action", response.context)
        self.assertContains(response, "Расходы центра")
        self.assertContains(response, "Реестр расходов")
        self.assertContains(response, 'id="center-expense-list"')
        self.assertContains(response, reverse("center_expense_create"))
        self.assertContains(response, "Канцтовары")
        self.assertContains(response, "распределен")
        self.assertContains(response, 'data-label="Распределение"')

    def test_center_expense_report_renders(self):
        expense = CenterExpense.objects.create(
            category=self.expense_category,
            counterparty=self.counterparty,
            title="Отчетный расход",
            total_amount=Decimal("1000.00"),
        )
        ExpenseFundingSplit.objects.create(
            expense=expense,
            funding_source=self.funding,
            amount=Decimal("1000.00"),
        )

        response = self.client.get(reverse("center_expense_report"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("report", response.context)
        self.assertIn("expense_report_summary_items", response.context)
        self.assertIn("expense_report_attention_items", response.context)
        self.assertContains(response, "Отчет руководителя")
        self.assertContains(response, "Расходы центра")
        self.assertContains(response, "Категории")
        self.assertContains(response, "Источники покрытия")
        self.assertContains(response, "Отчетный расход")
        self.assertContains(response, reverse("center_expense_list"))

    def test_center_expense_report_filters_period_without_writes(self):
        inside = CenterExpense.objects.create(
            expense_date=timezone.localdate(),
            category=self.expense_category,
            counterparty=self.counterparty,
            title="В периоде",
            total_amount=Decimal("1000.00"),
        )
        outside = CenterExpense.objects.create(
            expense_date=timezone.localdate() - timedelta(days=5),
            category=self.expense_category,
            counterparty=self.counterparty,
            title="Вне периода",
            total_amount=Decimal("500.00"),
        )
        ExpenseFundingSplit.objects.create(
            expense=inside,
            funding_source=self.funding,
            amount=Decimal("1000.00"),
        )
        ExpenseFundingSplit.objects.create(
            expense=outside,
            funding_source=self.funding,
            amount=Decimal("500.00"),
        )
        ledger_count = LedgerEntry.objects.count()

        response = self.client.get(
            reverse("center_expense_report"),
            {
                "date_from": timezone.localdate().isoformat(),
                "date_to": timezone.localdate().isoformat(),
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context["report"].summary.expense_count, 1)
        self.assertContains(response, "В периоде")
        self.assertNotContains(response, "Вне периода")
        self.assertEqual(LedgerEntry.objects.count(), ledger_count)

    def test_center_expense_report_filters_funding_source_by_split_amount(self):
        mixed = CenterExpense.objects.create(
            category=self.expense_category,
            counterparty=self.counterparty,
            title="Смешанное покрытие",
            total_amount=Decimal("1000.00"),
        )
        ExpenseFundingSplit.objects.create(
            expense=mixed,
            funding_source=self.funding,
            amount=Decimal("400.00"),
        )
        ExpenseFundingSplit.objects.create(
            expense=mixed,
            funding_source=self.funding_grant,
            amount=Decimal("600.00"),
        )

        response = self.client.get(
            reverse("center_expense_report"),
            {"funding_source": self.funding.pk},
        )

        report = response.context["report"]
        self.assertEqual(response.status_code, 200)
        self.assertEqual(report.summary.total_amount, Decimal("1000.00"))
        self.assertEqual(report.summary.allocated_amount, Decimal("400.00"))
        self.assertEqual(len(report.funding_rows), 1)
        self.assertContains(response, "Смешанное покрытие")

    def test_center_expense_create_saves_draft_with_funding_splits(self):
        ledger_count = LedgerEntry.objects.count()

        response = self.client.post(
            reverse("center_expense_create"),
            self._expense_post_data(),
        )

        expense = CenterExpense.objects.get(title="Покупка материалов")
        self.assertRedirects(response, reverse("center_expense_edit", args=[expense.pk]))
        self.assertEqual(expense.status, CenterExpense.Status.DRAFT)
        self.assertEqual(expense.created_by, self.admin)
        self.assertEqual(expense.funding_splits.count(), 2)
        self.assertEqual(expense.funding_split_total, Decimal("1000.00"))
        self.assertEqual(LedgerEntry.objects.count(), ledger_count)

    def test_center_expense_create_rejects_duplicate_funding_source(self):
        response = self.client.post(
            reverse("center_expense_create"),
            self._expense_post_data(
                first_source=self.funding,
                second_source=self.funding,
            ),
        )

        self.assertEqual(response.status_code, 200)
        self.assertFalse(CenterExpense.objects.filter(title="Покупка материалов").exists())
        self.assertContains(response, "Один источник финансирования нельзя указать дважды")

    def test_center_expense_edit_updates_draft(self):
        expense = CenterExpense.objects.create(
            category=self.expense_category,
            counterparty=self.counterparty,
            title="Старое название",
            total_amount=Decimal("1000.00"),
        )

        response = self.client.post(
            reverse("center_expense_edit", args=[expense.pk]),
            self._expense_post_data(title="Новое название"),
        )

        self.assertRedirects(response, reverse("center_expense_list"))
        expense.refresh_from_db()
        self.assertEqual(expense.title, "Новое название")
        self.assertEqual(expense.funding_splits.count(), 2)

    def test_center_expense_edit_shows_existing_split_control(self):
        expense = CenterExpense.objects.create(
            category=self.expense_category,
            counterparty=self.counterparty,
            title="Распределенный расход",
            total_amount=Decimal("1000.00"),
        )
        ExpenseFundingSplit.objects.create(
            expense=expense,
            funding_source=self.funding,
            amount=Decimal("1000.00"),
        )

        response = self.client.get(reverse("center_expense_edit", args=[expense.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Сумма распределения совпадает с суммой расхода.")

    def test_center_expense_edit_blocks_non_draft(self):
        expense = CenterExpense.objects.create(
            category=self.expense_category,
            counterparty=self.counterparty,
            title="Утвержденный расход",
            total_amount=Decimal("1000.00"),
            status=CenterExpense.Status.APPROVED,
        )

        response = self.client.post(
            reverse("center_expense_edit", args=[expense.pk]),
            self._expense_post_data(title="Попытка правки"),
        )

        self.assertRedirects(response, reverse("center_expense_edit", args=[expense.pk]))
        expense.refresh_from_db()
        self.assertEqual(expense.title, "Утвержденный расход")


class EquipmentAssetViewTests(NewViewsTestBase):
    def _equipment_expense(self, title: str = "Покупка оборудования") -> CenterExpense:
        return CenterExpense.objects.create(
            category=self.equipment_expense_category,
            counterparty=self.counterparty,
            title=title,
            total_amount=Decimal("2500.00"),
        )

    def _asset_post_data(self, *, purchase_expense=None, status=EquipmentAsset.Status.ACTIVE):
        purchase_expense = purchase_expense or self._equipment_expense()
        return {
            "name": "Балансировочная платформа",
            "asset_type": EquipmentAsset.AssetType.THERAPY_EQUIPMENT,
            "inventory_number": "ASSET-001",
            "purchase_date": timezone.localdate().isoformat(),
            "purchase_expense": purchase_expense.pk,
            "total_amount": "2500.00",
            "status": status,
            "location": "Кабинет ЛФК",
            "responsible_staff": self.staff.pk,
            "notes": "Для занятий ЛФК",
        }

    def test_equipment_asset_list_renders(self):
        expense = self._equipment_expense()
        EquipmentAsset.objects.create(
            name="Балансировочная платформа",
            asset_type=EquipmentAsset.AssetType.THERAPY_EQUIPMENT,
            inventory_number="ASSET-001",
            purchase_expense=expense,
            total_amount=Decimal("2500.00"),
            location="Кабинет ЛФК",
            responsible_staff=self.staff,
        )

        response = self.client.get(reverse("equipment_asset_list"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("asset_summary_items", response.context)
        self.assertIn("asset_next_action", response.context)
        self.assertContains(response, "Оборудование")
        self.assertContains(response, "Реестр активов")
        self.assertContains(response, "Балансировочная платформа")
        self.assertContains(response, reverse("equipment_asset_create"))
        self.assertContains(response, 'id="equipment-asset-list"')

    def test_equipment_asset_create_links_equipment_expense(self):
        ledger_count = LedgerEntry.objects.count()
        expense = self._equipment_expense()

        response = self.client.post(
            reverse("equipment_asset_create"),
            self._asset_post_data(purchase_expense=expense),
        )

        asset = EquipmentAsset.objects.get(inventory_number="ASSET-001")
        self.assertRedirects(response, reverse("equipment_asset_edit", args=[asset.pk]))
        self.assertEqual(asset.purchase_expense, expense)
        self.assertEqual(asset.responsible_staff, self.staff)
        self.assertEqual(LedgerEntry.objects.count(), ledger_count)

    def test_equipment_asset_create_rejects_non_equipment_expense(self):
        household_expense = CenterExpense.objects.create(
            category=self.expense_category,
            counterparty=self.counterparty,
            title="Хозяйственный расход",
            total_amount=Decimal("1000.00"),
        )

        response = self.client.post(
            reverse("equipment_asset_create"),
            self._asset_post_data(purchase_expense=household_expense),
        )

        self.assertEqual(response.status_code, 200)
        self.assertFalse(EquipmentAsset.objects.filter(inventory_number="ASSET-001").exists())
        self.assertContains(response, "Выберите корректный вариант")

    def test_equipment_asset_edit_status_does_not_delete_expense(self):
        expense = self._equipment_expense()
        asset = EquipmentAsset.objects.create(
            name="Балансировочная платформа",
            asset_type=EquipmentAsset.AssetType.THERAPY_EQUIPMENT,
            inventory_number="ASSET-001",
            purchase_expense=expense,
            total_amount=Decimal("2500.00"),
            responsible_staff=self.staff,
        )
        ledger_count = LedgerEntry.objects.count()

        response = self.client.post(
            reverse("equipment_asset_edit", args=[asset.pk]),
            self._asset_post_data(
                purchase_expense=expense,
                status=EquipmentAsset.Status.WRITTEN_OFF,
            ),
        )

        self.assertRedirects(response, reverse("equipment_asset_list"))
        asset.refresh_from_db()
        self.assertEqual(asset.status, EquipmentAsset.Status.WRITTEN_OFF)
        self.assertTrue(CenterExpense.objects.filter(pk=expense.pk).exists())
        self.assertEqual(LedgerEntry.objects.count(), ledger_count)


class ContractRegistryViewTests(NewViewsTestBase):
    def _financial_counts(self):
        return {
            "balances": BalanceAccount.objects.count(),
            "ledger": LedgerEntry.objects.count(),
            "payments": Payment.objects.count(),
            "payroll": PayrollAccrual.objects.count(),
        }

    def _signer_link(self) -> RecipientRepresentative:
        return RecipientRepresentative.objects.get(child=self.child, representative=self.parent)

    def _service_template(self) -> ContractTemplate:
        return ContractTemplate.objects.create(
            template_type=ContractTemplate.TemplateType.RECIPIENT_SERVICE,
            title="Шаблон договора с получателем",
            version="1",
        )

    def _donation_template(self) -> ContractTemplate:
        return ContractTemplate.objects.create(
            template_type=ContractTemplate.TemplateType.SPONSOR,
            title="Шаблон пожертвования",
            version="1",
        )

    def _contract_document(self, *, child=None, title: str = "Файл договора") -> Document:
        return Document.objects.create(
            child=child or self.child,
            category=Document.Category.CONTRACT,
            title=title,
            file="documents/contract.txt",
        )

    def _docx_upload(self, filename: str, text: str) -> SimpleUploadedFile:
        from docx import Document as WordDocument

        payload = BytesIO()
        document = WordDocument()
        document.add_paragraph(text)
        document.save(payload)
        return SimpleUploadedFile(
            filename,
            payload.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def _docx_text(self, payload: bytes) -> str:
        from docx import Document as WordDocument

        document = WordDocument(BytesIO(payload))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)

    def test_contract_list_renders(self):
        template = self._service_template()
        contract = ServiceContract.objects.create(
            child=self.child,
            representative_link=self._signer_link(),
            number="S-001",
            signed_on=timezone.localdate(),
            template=template,
        )
        document = self._contract_document(title="Договор S-001")
        contract.document = document
        contract.save(update_fields=["document", "updated_at"])
        ContractLegalSnapshot.objects.create(
            contract_kind=ContractLegalSnapshot.ContractKind.SERVICE,
            service_contract=contract,
            document=document,
            contract_snapshot={"number": "S-001"},
        )

        response = self.client.get(reverse("contract_list"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("contract_summary_items", response.context)
        self.assertIn("contract_next_action", response.context)
        self.assertContains(response, "Договоры")
        self.assertContains(response, "Реестр договоров")
        self.assertContains(response, "Договоры с получателями")
        self.assertContains(response, "S-001")
        self.assertContains(response, reverse("service_contract_create"))
        self.assertContains(response, reverse("service_contract_word", args=[contract.pk]))
        self.assertContains(response, "Word")
        self.assertContains(response, "реквизиты зафиксированы:")

    def test_contract_template_create(self):
        response = self.client.post(
            reverse("contract_template_create"),
            {
                "template_type": ContractTemplate.TemplateType.RECIPIENT_SERVICE,
                "title": "Новый шаблон",
                "version": "2",
                "is_active": "on",
                "notes": "Для будущей генерации",
            },
        )

        template = ContractTemplate.objects.get(title="Новый шаблон")
        self.assertRedirects(response, reverse("contract_template_edit", args=[template.pk]))
        self.assertEqual(template.version, "2")

    def test_contract_template_create_shows_placeholder_reference(self):
        response = self.client.get(reverse("contract_template_create"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "center.full_name")
        self.assertContains(response, "representative.passport_number")
        self.assertContains(response, "service_spec.rows")
        self.assertContains(response, "certificate.number")
        self.assertContains(response, ".docx")

    def test_contract_template_rejects_legacy_doc_upload(self):
        response = self.client.post(
            reverse("contract_template_create"),
            {
                "template_type": ContractTemplate.TemplateType.RECIPIENT_SERVICE,
                "title": "Legacy template",
                "version": "1",
                "file": SimpleUploadedFile(
                    "legacy.doc",
                    b"legacy word payload",
                    content_type="application/msword",
                ),
                "is_active": "on",
                "notes": "",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertFalse(ContractTemplate.objects.filter(title="Legacy template").exists())
        self.assertContains(response, ".docx")

    def test_donation_contract_create_links_funding_source_without_financial_facts(self):
        template = self._donation_template()
        ledger_count = LedgerEntry.objects.count()
        payment_count = Payment.objects.count()

        response = self.client.post(
            reverse("donation_contract_create"),
            {
                "counterparty": self.counterparty.pk,
                "funding_source": self.funding_grant.pk,
                "contract_type": DonationContract.ContractType.PROJECT,
                "number": "D-001",
                "signed_on": timezone.localdate().isoformat(),
                "valid_from": timezone.localdate().isoformat(),
                "valid_until": "",
                "amount_limit": "100000.00",
                "status": DonationContract.Status.ACTIVE,
                "template": template.pk,
                "document": "",
                "notes": "Грантовый договор",
            },
        )

        contract = DonationContract.objects.get(number="D-001")
        self.assertRedirects(response, reverse("donation_contract_edit", args=[contract.pk]))
        self.assertEqual(contract.funding_source, self.funding_grant)
        self.assertIsNone(contract.document_id)
        self.assertEqual(LedgerEntry.objects.count(), ledger_count)
        self.assertEqual(Payment.objects.count(), payment_count)

    def test_service_contract_create_links_child_signer_template_and_document(self):
        template = self._service_template()
        document = self._contract_document()

        response = self.client.post(
            reverse("service_contract_create"),
            {
                "child": self.child.pk,
                "representative_link": self._signer_link().pk,
                "contract_type": ServiceContract.ContractType.STANDARD,
                "number": "S-002",
                "signed_on": timezone.localdate().isoformat(),
                "valid_from": timezone.localdate().isoformat(),
                "valid_until": "",
                "status": ServiceContract.Status.ACTIVE,
                "template": template.pk,
                "document": document.pk,
                "notes": "Основной договор",
            },
        )

        contract = ServiceContract.objects.get(number="S-002")
        self.assertRedirects(response, reverse("service_contract_edit", args=[contract.pk]))
        self.assertEqual(contract.child, self.child)
        self.assertEqual(contract.representative_link, self._signer_link())
        self.assertEqual(contract.document, document)

    def test_service_contract_create_rejects_signer_from_other_child(self):
        other_parent = ParentGuardian.objects.create(
            last_name="Петрова",
            first_name="Анна",
            phone="+7 900 000-00-09",
        )
        other_child = Child.objects.create(
            last_name="Петров",
            first_name="Петр",
            primary_parent=other_parent,
        )
        other_link = RecipientRepresentative.objects.get(
            child=other_child,
            representative=other_parent,
        )

        response = self.client.post(
            reverse("service_contract_create"),
            {
                "child": self.child.pk,
                "representative_link": other_link.pk,
                "contract_type": ServiceContract.ContractType.STANDARD,
                "number": "S-003",
                "signed_on": timezone.localdate().isoformat(),
                "valid_from": "",
                "valid_until": "",
                "status": ServiceContract.Status.DRAFT,
                "template": "",
                "document": "",
                "notes": "",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertFalse(ServiceContract.objects.filter(number="S-003").exists())
        self.assertContains(response, "Выберите корректный вариант")

    def test_contract_import_preview_get(self):
        response = self.client.get(reverse("contract_import_preview"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("form", response.context)
        self.assertContains(response, "Preview")
        self.assertContains(response, "Контрагенты")
        self.assertContains(response, "Расходы")

    def test_contract_import_preview_post_expenses_does_not_create_records(self):
        from django.core.files.uploadedfile import SimpleUploadedFile

        before_count = CenterExpense.objects.count()
        upload = SimpleUploadedFile(
            "expenses.csv",
            (
                "Дата расхода;Категория;Название;Сумма;Источник финансирования;Сумма источника\n"
                f"2026-07-17;{self.expense_category.name};Тестовый расход;1000;{self.funding.name};1000\n"
            ).encode(),
            content_type="text/csv",
        )

        response = self.client.post(
            reverse("contract_import_preview"),
            {
                "import_type": "expenses",
                "file": upload,
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(CenterExpense.objects.count(), before_count)
        self.assertEqual(response.context["preview"].valid_count, 1)
        self.assertContains(response, "Тестовый расход")

    def test_service_contract_pdf_download_does_not_create_document_or_financial_facts(self):
        template = self._service_template()
        contract = ServiceContract.objects.create(
            child=self.child,
            representative_link=self._signer_link(),
            number="S-004",
            signed_on=timezone.localdate(),
            template=template,
        )
        document_count = Document.objects.count()
        ledger_count = LedgerEntry.objects.count()
        payment_count = Payment.objects.count()

        response = self.client.get(reverse("service_contract_pdf", args=[contract.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"], "application/pdf")
        payload = b"".join(response.streaming_content)
        self.assertTrue(payload.startswith(b"%PDF"))
        self.assertEqual(Document.objects.count(), document_count)
        self.assertEqual(LedgerEntry.objects.count(), ledger_count)
        self.assertEqual(Payment.objects.count(), payment_count)

    def test_donation_contract_pdf_download_does_not_create_document_or_financial_facts(self):
        template = self._donation_template()
        contract = DonationContract.objects.create(
            counterparty=self.counterparty,
            funding_source=self.funding_grant,
            contract_type=DonationContract.ContractType.PROJECT,
            number="D-004",
            signed_on=timezone.localdate(),
            amount_limit=Decimal("10000.00"),
            template=template,
        )
        document_count = Document.objects.count()
        ledger_count = LedgerEntry.objects.count()
        payment_count = Payment.objects.count()

        response = self.client.get(reverse("donation_contract_pdf", args=[contract.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"], "application/pdf")
        payload = b"".join(response.streaming_content)
        self.assertTrue(payload.startswith(b"%PDF"))
        self.assertEqual(Document.objects.count(), document_count)
        self.assertEqual(LedgerEntry.objects.count(), ledger_count)
        self.assertEqual(Payment.objects.count(), payment_count)

    def test_service_contract_word_generates_document_from_template_without_financial_facts(self):
        template = ContractTemplate.objects.create(
            template_type=ContractTemplate.TemplateType.RECIPIENT_SERVICE,
            title="Word шаблон услуг",
            version="1",
            file=self._docx_upload(
                "service_template.docx",
                "Договор {{ contract.number }} для {{ child.full_name }}, подписант {{ representative.full_name }}.",
            ),
        )
        contract = ServiceContract.objects.create(
            child=self.child,
            representative_link=self._signer_link(),
            number="S-WORD",
            signed_on=timezone.localdate(),
            template=template,
        )
        document_count = Document.objects.count()
        snapshot_count = ContractLegalSnapshot.objects.count()
        counts_before = self._financial_counts()

        response = self.client.post(reverse("service_contract_word", args=[contract.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        payload = b"".join(response.streaming_content)
        self.assertTrue(payload.startswith(b"PK"))
        generated_text = self._docx_text(payload)
        self.assertIn("S-WORD", generated_text)
        self.assertIn(self.child.full_name, generated_text)
        self.assertIn(self.parent.full_name, generated_text)
        self.assertNotIn("{{ contract.number }}", generated_text)

        contract.refresh_from_db()
        self.assertIsNotNone(contract.document_id)
        self.assertEqual(Document.objects.count(), document_count + 1)
        self.assertEqual(contract.document.child, self.child)
        self.assertEqual(contract.document.category, Document.Category.CONTRACT)
        self.assertEqual(contract.document.uploaded_by, self.admin)
        self.assertTrue(contract.document.file.name.endswith(".docx"))
        snapshot = contract.document.contract_legal_snapshot
        self.assertEqual(ContractLegalSnapshot.objects.count(), snapshot_count + 1)
        self.assertEqual(snapshot.contract_kind, ContractLegalSnapshot.ContractKind.SERVICE)
        self.assertEqual(snapshot.service_contract, contract)
        self.assertIsNone(snapshot.donation_contract_id)
        self.assertEqual(snapshot.generated_by, self.admin)
        self.assertEqual(snapshot.contract_snapshot["number"], "S-WORD")
        self.assertEqual(snapshot.recipient_snapshot["full_name"], self.child.full_name)
        self.assertEqual(snapshot.representative_snapshot["full_name"], self.parent.full_name)
        self.assertEqual(snapshot.template_snapshot["title"], "Word шаблон услуг")
        self.assertEqual(self._financial_counts(), counts_before)

    def test_service_contract_word_updates_existing_document_without_duplicate(self):
        existing_document = self._contract_document(title="Старый файл договора")
        contract = ServiceContract.objects.create(
            child=self.child,
            representative_link=self._signer_link(),
            number="S-REGEN",
            signed_on=timezone.localdate(),
            document=existing_document,
        )
        document_count = Document.objects.count()
        snapshot_count = ContractLegalSnapshot.objects.count()
        counts_before = self._financial_counts()

        response = self.client.post(reverse("service_contract_word", args=[contract.pk]))

        self.assertEqual(response.status_code, 200)
        contract.refresh_from_db()
        existing_document.refresh_from_db()
        self.assertEqual(Document.objects.count(), document_count)
        self.assertEqual(contract.document, existing_document)
        self.assertIn("S-REGEN", existing_document.title)
        self.assertTrue(existing_document.file.name.endswith(".docx"))
        snapshot = existing_document.contract_legal_snapshot
        self.assertEqual(ContractLegalSnapshot.objects.count(), snapshot_count + 1)
        self.assertEqual(snapshot.service_contract, contract)
        self.assertEqual(snapshot.contract_snapshot["number"], "S-REGEN")
        self.assertEqual(self._financial_counts(), counts_before)

        contract.number = "S-REGEN-2"
        contract.save(update_fields=["number", "updated_at"])
        response = self.client.post(reverse("service_contract_word", args=[contract.pk]))

        self.assertEqual(response.status_code, 200)
        snapshot.refresh_from_db()
        self.assertEqual(ContractLegalSnapshot.objects.count(), snapshot_count + 1)
        self.assertEqual(snapshot.contract_snapshot["number"], "S-REGEN-2")

    def test_service_contract_word_rejects_document_snapshot_from_other_contract(self):
        existing_document = self._contract_document(title="Старый файл договора")
        owner_contract = ServiceContract.objects.create(
            child=self.child,
            representative_link=self._signer_link(),
            number="S-OWNER",
            signed_on=timezone.localdate(),
            document=existing_document,
        )
        self.client.post(reverse("service_contract_word", args=[owner_contract.pk]))
        existing_document.refresh_from_db()
        self.assertIn("S-OWNER", existing_document.title)

        other_contract = ServiceContract.objects.create(
            child=self.child,
            representative_link=self._signer_link(),
            number="S-OTHER",
            signed_on=timezone.localdate(),
            document=existing_document,
        )
        response = self.client.post(reverse("service_contract_word", args=[other_contract.pk]))

        self.assertRedirects(response, reverse("contract_list"))
        existing_document.refresh_from_db()
        snapshot = existing_document.contract_legal_snapshot
        self.assertIn("S-OWNER", existing_document.title)
        self.assertEqual(snapshot.service_contract, owner_contract)
        self.assertEqual(snapshot.contract_snapshot["number"], "S-OWNER")

    def test_service_contract_word_replaces_v2_placeholders_with_blank_fallback(self):
        template = ContractTemplate.objects.create(
            template_type=ContractTemplate.TemplateType.RECIPIENT_SERVICE,
            title="Service v2 placeholders",
            version="2",
            file=self._docx_upload(
                "service_template_v2.docx",
                "Phone {{ representative.phone }} email {{ representative.email }} "
                "center {{ center.full_name }} child-address {{ child.address }} "
                "spec {{ service_spec.rows }} certificate {{ certificate.number }}.",
            ),
        )
        contract = ServiceContract.objects.create(
            child=self.child,
            representative_link=self._signer_link(),
            number="S-V2",
            signed_on=timezone.localdate(),
            template=template,
        )

        response = self.client.post(reverse("service_contract_word", args=[contract.pk]))

        self.assertEqual(response.status_code, 200)
        payload = b"".join(response.streaming_content)
        generated_text = self._docx_text(payload)
        self.assertIn(self.parent.phone, generated_text)
        self.assertIn(self.parent.email, generated_text)
        self.assertIn("_______________", generated_text)
        self.assertNotIn("{{ representative.phone }}", generated_text)
        self.assertNotIn("{{ center.full_name }}", generated_text)
        self.assertNotIn("{{ service_spec.rows }}", generated_text)

    def test_service_contract_word_uses_active_center_legal_profile(self):
        profile = CenterLegalProfile.objects.create(
            full_name="Автономная некоммерческая организация Радость моя",
            short_name="АНО Радость моя",
            director_full_name="Иванов Иван Иванович",
            director_position="Директор",
            authority_basis="Устава",
            inn="2500000000",
            bank_bik="040000001",
        )
        template = ContractTemplate.objects.create(
            template_type=ContractTemplate.TemplateType.RECIPIENT_SERVICE,
            title="Center profile placeholders",
            version="1",
            file=self._docx_upload(
                "service_center_profile.docx",
                "{{ center.full_name }} {{ center.short_name }} {{ center.director_position }} "
                "{{ center.authority_basis }} {{ center.inn }} {{ center.bank_bik }}",
            ),
        )
        contract = ServiceContract.objects.create(
            child=self.child,
            representative_link=self._signer_link(),
            number="S-CENTER",
            signed_on=timezone.localdate(),
            template=template,
        )

        response = self.client.post(reverse("service_contract_word", args=[contract.pk]))

        self.assertEqual(response.status_code, 200)
        payload = b"".join(response.streaming_content)
        generated_text = self._docx_text(payload)
        self.assertIn("Автономная некоммерческая организация Радость моя", generated_text)
        self.assertIn("АНО Радость моя", generated_text)
        self.assertIn("Директор", generated_text)
        self.assertIn("Устава", generated_text)
        self.assertIn("2500000000", generated_text)
        self.assertIn("040000001", generated_text)
        self.assertNotIn("{{ center.full_name }}", generated_text)
        contract.refresh_from_db()
        snapshot = contract.document.contract_legal_snapshot
        self.assertEqual(
            snapshot.center_snapshot["full_name"],
            "Автономная некоммерческая организация Радость моя",
        )

        profile.full_name = "Новое юридическое название центра"
        profile.save(update_fields=["full_name", "updated_at"])
        snapshot.refresh_from_db()
        self.assertEqual(
            snapshot.center_snapshot["full_name"],
            "Автономная некоммерческая организация Радость моя",
        )

    def test_donation_contract_word_generates_document_without_financial_facts(self):
        template = ContractTemplate.objects.create(
            template_type=ContractTemplate.TemplateType.SPONSOR,
            title="Word шаблон пожертвования",
            version="1",
            file=self._docx_upload(
                "donation_template.docx",
                "Пожертвование {{ contract.number }}: {{ counterparty.name }} / {{ funding_source.name }} / {{ donation.amount_limit }}.",
            ),
        )
        contract = DonationContract.objects.create(
            counterparty=self.counterparty,
            funding_source=self.funding_grant,
            contract_type=DonationContract.ContractType.PROJECT,
            number="D-WORD",
            signed_on=timezone.localdate(),
            amount_limit=Decimal("12345.67"),
            template=template,
        )
        document_count = Document.objects.count()
        snapshot_count = ContractLegalSnapshot.objects.count()
        counts_before = self._financial_counts()

        response = self.client.post(reverse("donation_contract_word", args=[contract.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        payload = b"".join(response.streaming_content)
        self.assertTrue(payload.startswith(b"PK"))
        generated_text = self._docx_text(payload)
        self.assertIn("D-WORD", generated_text)
        self.assertIn(self.counterparty.name, generated_text)
        self.assertIn(self.funding_grant.name, generated_text)
        self.assertIn("12 345,67 ₽", generated_text)

        contract.refresh_from_db()
        self.assertIsNotNone(contract.document_id)
        self.assertEqual(Document.objects.count(), document_count + 1)
        self.assertEqual(contract.document.category, Document.Category.CONTRACT)
        self.assertEqual(contract.document.target_type, Document.TargetType.COUNTERPARTY)
        self.assertEqual(contract.document.counterparty, self.counterparty)
        self.assertIsNone(contract.document.child_id)
        self.assertEqual(contract.document.uploaded_by, self.admin)
        self.assertTrue(contract.document.file.name.endswith(".docx"))
        snapshot = contract.document.contract_legal_snapshot
        self.assertEqual(ContractLegalSnapshot.objects.count(), snapshot_count + 1)
        self.assertEqual(snapshot.contract_kind, ContractLegalSnapshot.ContractKind.DONATION)
        self.assertEqual(snapshot.donation_contract, contract)
        self.assertIsNone(snapshot.service_contract_id)
        self.assertEqual(snapshot.generated_by, self.admin)
        self.assertEqual(snapshot.contract_snapshot["number"], "D-WORD")
        self.assertEqual(snapshot.contract_snapshot["amount_limit"], "12345.67")
        self.assertEqual(snapshot.counterparty_snapshot["name"], self.counterparty.name)
        self.assertEqual(snapshot.funding_source_snapshot["name"], self.funding_grant.name)
        self.assertEqual(snapshot.template_snapshot["title"], "Word шаблон пожертвования")
        self.assertEqual(self._financial_counts(), counts_before)

    def test_donation_contract_word_updates_existing_document_without_duplicate(self):
        existing_document = Document.objects.create(
            target_type=Document.TargetType.COUNTERPARTY,
            counterparty=self.counterparty,
            category=Document.Category.CONTRACT,
            title="Старый договор пожертвования",
            file="documents/donation-old.txt",
        )
        contract = DonationContract.objects.create(
            counterparty=self.counterparty,
            funding_source=self.funding_grant,
            contract_type=DonationContract.ContractType.PROJECT,
            number="D-REGEN",
            signed_on=timezone.localdate(),
            document=existing_document,
        )
        document_count = Document.objects.count()
        snapshot_count = ContractLegalSnapshot.objects.count()
        counts_before = self._financial_counts()

        response = self.client.post(reverse("donation_contract_word", args=[contract.pk]))

        self.assertEqual(response.status_code, 200)
        contract.refresh_from_db()
        existing_document.refresh_from_db()
        self.assertEqual(Document.objects.count(), document_count)
        self.assertEqual(contract.document, existing_document)
        self.assertIn("D-REGEN", existing_document.title)
        self.assertEqual(existing_document.counterparty, self.counterparty)
        self.assertTrue(existing_document.file.name.endswith(".docx"))
        snapshot = existing_document.contract_legal_snapshot
        self.assertEqual(ContractLegalSnapshot.objects.count(), snapshot_count + 1)
        self.assertEqual(snapshot.donation_contract, contract)
        self.assertEqual(snapshot.contract_snapshot["number"], "D-REGEN")
        self.assertEqual(self._financial_counts(), counts_before)

        contract.number = "D-REGEN-2"
        contract.save(update_fields=["number", "updated_at"])
        response = self.client.post(reverse("donation_contract_word", args=[contract.pk]))

        self.assertEqual(response.status_code, 200)
        snapshot.refresh_from_db()
        self.assertEqual(ContractLegalSnapshot.objects.count(), snapshot_count + 1)
        self.assertEqual(snapshot.contract_snapshot["number"], "D-REGEN-2")

    def test_donation_contract_word_rejects_document_snapshot_from_other_contract(self):
        existing_document = Document.objects.create(
            target_type=Document.TargetType.COUNTERPARTY,
            counterparty=self.counterparty,
            category=Document.Category.CONTRACT,
            title="Старый договор пожертвования",
            file="documents/donation-old.txt",
        )
        owner_contract = DonationContract.objects.create(
            counterparty=self.counterparty,
            funding_source=self.funding_grant,
            contract_type=DonationContract.ContractType.PROJECT,
            number="D-OWNER",
            signed_on=timezone.localdate(),
            document=existing_document,
        )
        self.client.post(reverse("donation_contract_word", args=[owner_contract.pk]))
        existing_document.refresh_from_db()
        self.assertIn("D-OWNER", existing_document.title)

        other_contract = DonationContract.objects.create(
            counterparty=self.counterparty,
            funding_source=self.funding_grant,
            contract_type=DonationContract.ContractType.PROJECT,
            number="D-OTHER",
            signed_on=timezone.localdate(),
            document=existing_document,
        )
        response = self.client.post(reverse("donation_contract_word", args=[other_contract.pk]))

        self.assertRedirects(response, reverse("contract_list"))
        existing_document.refresh_from_db()
        snapshot = existing_document.contract_legal_snapshot
        self.assertIn("D-OWNER", existing_document.title)
        self.assertEqual(snapshot.donation_contract, owner_contract)
        self.assertEqual(snapshot.contract_snapshot["number"], "D-OWNER")

    def test_donation_contract_word_replaces_v2_placeholders_with_blank_fallback(self):
        self.counterparty.contact_person = "Legal Contact"
        self.counterparty.phone = "+7 900 000-00-88"
        self.counterparty.email = "vendor@example.local"
        self.counterparty.save(update_fields=["contact_person", "phone", "email", "updated_at"])
        template = ContractTemplate.objects.create(
            template_type=ContractTemplate.TemplateType.SPONSOR,
            title="Donation v2 placeholders",
            version="2",
            file=self._docx_upload(
                "donation_template_v2.docx",
                "Counterparty {{ counterparty.contact_person }} {{ counterparty.phone }} "
                "{{ counterparty.email }} funding {{ funding_source.type }} "
                "monthly {{ donation.monthly_amount }} bank {{ counterparty.bank_bik }}.",
            ),
        )
        contract = DonationContract.objects.create(
            counterparty=self.counterparty,
            funding_source=self.funding_grant,
            contract_type=DonationContract.ContractType.MONTHLY,
            number="D-V2",
            signed_on=timezone.localdate(),
            amount_limit=Decimal("1000.00"),
            template=template,
        )

        response = self.client.post(reverse("donation_contract_word", args=[contract.pk]))

        self.assertEqual(response.status_code, 200)
        payload = b"".join(response.streaming_content)
        generated_text = self._docx_text(payload)
        self.assertIn("Legal Contact", generated_text)
        self.assertIn("+7 900 000-00-88", generated_text)
        self.assertIn("vendor@example.local", generated_text)
        self.assertIn("_______________", generated_text)
        self.assertNotIn("{{ counterparty.contact_person }}", generated_text)
        self.assertNotIn("{{ donation.monthly_amount }}", generated_text)
        self.assertNotIn("{{ counterparty.bank_bik }}", generated_text)

    def test_service_contract_word_invalid_template_does_not_create_document(self):
        template = ContractTemplate.objects.create(
            template_type=ContractTemplate.TemplateType.RECIPIENT_SERVICE,
            title="Битый шаблон",
            version="1",
            file=SimpleUploadedFile(
                "broken.docx",
                b"not a docx",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        )
        contract = ServiceContract.objects.create(
            child=self.child,
            representative_link=self._signer_link(),
            number="S-BROKEN",
            signed_on=timezone.localdate(),
            template=template,
        )
        document_count = Document.objects.count()

        response = self.client.post(reverse("service_contract_word", args=[contract.pk]))

        self.assertRedirects(response, reverse("contract_list"))
        contract.refresh_from_db()
        self.assertIsNone(contract.document_id)
        self.assertEqual(Document.objects.count(), document_count)


class StaffCompensationRuleViewTests(NewViewsTestBase):
    def test_staff_compensation_rule_list_renders(self):
        response = self.client.get(reverse("staff_compensation_rule_list"))
        self.assertEqual(response.status_code, 200)
        self.assertIn("compensation_summary_items", response.context)
        self.assertIn("compensation_next_action", response.context)
        self.assertEqual(
            response.context["compensation_next_action"]["title"],
            "Создать первую ставку",
        )
        self.assertContains(response, "Следующее действие")
        self.assertContains(response, "Правила начисления")
        self.assertContains(response, 'id="compensation-rule-list"')
        self.assertContains(response, reverse("staff_compensation_rule_create"))
        self.assertContains(response, "directory-table")
        self.assertContains(response, 'data-label="Ставки"')

    def test_staff_compensation_rule_forms_show_operator_control(self):
        rule = StaffCompensationRule.objects.create(
            staff_member=self.staff,
            service=self.service,
            funding_source=self.funding_grant,
            amount=Decimal("500"),
        )
        cases = [
            (reverse("staff_compensation_rule_create"), "Контроль ставки", "Групповые занятия"),
            (
                reverse("staff_compensation_rule_edit", args=[rule.pk]),
                "Контроль ставки",
                "Период и payroll",
            ),
        ]
        for url, title, control_item in cases:
            with self.subTest(url=url):
                response = self.client.get(url)
                self.assertEqual(response.status_code, 200)
                self.assertIn("object_form_control_items", response.context)
                self.assertContains(response, 'class="object-form-layout"')
                self.assertContains(response, title)
                self.assertContains(response, control_item)

    def test_staff_compensation_rule_list_uses_current_period_status(self):
        today = timezone.localdate()
        StaffCompensationRule.objects.create(
            staff_member=self.staff,
            service=self.service,
            amount=Decimal("500"),
            starts_on=today - timedelta(days=1),
            ends_on=today + timedelta(days=1),
        )
        StaffCompensationRule.objects.create(
            staff_member=self.staff,
            service=self.service2,
            amount=Decimal("600"),
            starts_on=today + timedelta(days=10),
        )
        StaffCompensationRule.objects.create(
            staff_member=self.staff,
            amount=Decimal("700"),
            ends_on=today - timedelta(days=1),
        )

        response = self.client.get(reverse("staff_compensation_rule_list"))

        self.assertEqual(response.status_code, 200)
        active_summary = next(
            item for item in response.context["compensation_summary_items"] if item["label"] == "Активны"
        )
        self.assertEqual(active_summary["value"], "1")
        self.assertEqual(
            response.context["compensation_next_action"]["title"],
            "Проверить истекшие ставки",
        )
        self.assertContains(response, "активна")
        self.assertContains(response, "ожидает")
        self.assertContains(response, "истекла")

    def test_staff_compensation_rule_create(self):
        response = self.client.post(
            reverse("staff_compensation_rule_create"),
            {
                "staff_member": self.staff.pk,
                "service": self.service.pk,
                "funding_source": self.funding_grant.pk,
                "session_scope": StaffCompensationRule.SessionScope.ALL,
                "rate_type": StaffCompensationRule.RateType.PER_SESSION,
                "amount": "700.00",
                "group_pay_policy": StaffCompensationRule.GroupPayPolicy.PER_SESSION,
                "group_fixed_amount": "",
                "min_duration_minutes": "30",
                "max_duration_minutes": "45",
                "starts_on": "2026-01-01",
                "ends_on": "2026-12-31",
                "is_active": "on",
                "note": "Ставка руководителя",
            },
        )

        rule = StaffCompensationRule.objects.get(staff_member=self.staff, service=self.service)
        self.assertRedirects(response, reverse("staff_compensation_rule_edit", args=[rule.pk]))
        self.assertEqual(rule.amount, Decimal("700.00"))
        self.assertEqual(rule.funding_source, self.funding_grant)
        self.assertEqual(rule.session_scope, StaffCompensationRule.SessionScope.ALL)
        self.assertEqual(rule.group_pay_policy, StaffCompensationRule.GroupPayPolicy.PER_SESSION)
        self.assertEqual(rule.min_duration_minutes, 30)
        self.assertEqual(rule.max_duration_minutes, 45)

    def test_staff_compensation_rule_edit(self):
        rule = StaffCompensationRule.objects.create(
            staff_member=self.staff,
            service=self.service,
            amount=Decimal("500"),
        )

        response = self.client.post(
            reverse("staff_compensation_rule_edit", args=[rule.pk]),
            {
                "staff_member": self.staff.pk,
                "service": "",
                "funding_source": "",
                "session_scope": StaffCompensationRule.SessionScope.GROUP,
                "rate_type": StaffCompensationRule.RateType.HOURLY,
                "amount": "900.00",
                "group_pay_policy": StaffCompensationRule.GroupPayPolicy.FIXED_GROUP_AMOUNT,
                "group_fixed_amount": "450.00",
                "min_duration_minutes": "",
                "max_duration_minutes": "",
                "starts_on": "2026-01-01",
                "ends_on": "",
                "is_active": "on",
                "note": "Почасовая общая ставка",
            },
        )

        self.assertRedirects(response, reverse("staff_compensation_rule_list"))
        rule.refresh_from_db()
        self.assertIsNone(rule.service)
        self.assertEqual(rule.session_scope, StaffCompensationRule.SessionScope.GROUP)
        self.assertEqual(rule.rate_type, StaffCompensationRule.RateType.HOURLY)
        self.assertEqual(rule.amount, Decimal("900.00"))
        self.assertEqual(
            rule.group_pay_policy,
            StaffCompensationRule.GroupPayPolicy.FIXED_GROUP_AMOUNT,
        )
        self.assertEqual(rule.group_fixed_amount, Decimal("450.00"))
        self.assertIsNone(rule.min_duration_minutes)
        self.assertIsNone(rule.max_duration_minutes)

    def test_staff_compensation_rule_fixed_group_amount_required(self):
        response = self.client.post(
            reverse("staff_compensation_rule_create"),
            {
                "staff_member": self.staff.pk,
                "service": self.service.pk,
                "funding_source": "",
                "session_scope": StaffCompensationRule.SessionScope.GROUP,
                "rate_type": StaffCompensationRule.RateType.PER_SESSION,
                "amount": "700.00",
                "group_pay_policy": StaffCompensationRule.GroupPayPolicy.FIXED_GROUP_AMOUNT,
                "group_fixed_amount": "",
                "min_duration_minutes": "",
                "max_duration_minutes": "",
                "starts_on": "",
                "ends_on": "",
                "is_active": "on",
                "note": "",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertFormError(
            response.context["form"],
            "group_fixed_amount",
            "Укажите фиксированную сумму для группового занятия.",
        )

    def test_staff_compensation_rule_toggle(self):
        rule = StaffCompensationRule.objects.create(
            staff_member=self.staff,
            service=self.service,
            amount=Decimal("500"),
            is_active=True,
        )

        response = self.client.post(reverse("staff_compensation_rule_toggle", args=[rule.pk]))

        self.assertRedirects(response, reverse("staff_compensation_rule_list"))
        rule.refresh_from_db()
        self.assertFalse(rule.is_active)


class ScheduleViewTests(NewViewsTestBase):
    def test_schedule_renders(self):
        response = self.client.get(reverse("schedule"))
        self.assertEqual(response.status_code, 200)
        self.assertIn("day", response.context)
        self.assertContains(response, "staffDaySchedule")
        self.assertContains(response, "День по специалистам")
        self.assertContains(response, "roomFilter")
        self.assertContains(response, "statusFilter")
        self.assertContains(response, "staffDayStatusSummary")
        self.assertContains(response, "staffLaneMode")
        self.assertContains(response, "roomLaneMode")
        self.assertContains(response, "scheduleCreateLink")
        self.assertContains(response, "resetScheduleFilters")
        self.assertContains(response, "modalMoveLink")
        self.assertContains(response, "modalCancelLink")
        self.assertContains(response, "operations/searchable_select.js")


class SpecialistHomeTests(NewViewsTestBase):
    def test_specialist_home_renders_for_admin(self):
        response = self.client.get(reverse("specialist_home"))
        self.assertEqual(response.status_code, 200)

    def test_specialist_can_view_home(self):
        self.client.logout()
        self.client.force_login(self.staff_user)
        response = self.client.get(reverse("specialist_home"))
        self.assertEqual(response.status_code, 200)

    def test_specialist_home_shows_overview_and_schedule_anchors(self):
        from operations.models import StaffAvailability, TimeOffRequest

        day = timezone.localdate() + timedelta(days=1)
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(day, time(10, 0)),
            ends_at=_local_dt(day, time(10, 30)),
            billing_account=self.account,
        )
        StaffAvailability.objects.create(
            staff_member=self.staff,
            weekday=0,
            starts_at=time(9, 0),
            ends_at=time(17, 0),
            is_active=True,
        )
        TimeOffRequest.objects.create(
            staff_member=self.staff,
            request_type=TimeOffRequest.RequestType.DAY_OFF,
            starts_on=day,
            ends_on=day,
            reason="Семейные обстоятельства",
        )
        self.client.logout()
        self.client.force_login(self.staff_user)

        response = self.client.get(reverse("specialist_home"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("specialist_summary_items", response.context)
        self.assertIn("specialist_next_action", response.context)
        self.assertEqual(
            response.context["specialist_next_action"]["title"],
            "Следующее занятие в расписании",
        )
        self.assertContains(response, "Следующее действие")
        self.assertContains(response, "Следующее занятие в расписании")
        self.assertContains(response, 'id="specialist-schedule"')
        self.assertContains(response, 'href="#staff-availability"')
        self.assertContains(response, f'id="appointment-{appointment.pk}"')
        self.assertContains(response, "staff-ops-table")
        self.assertContains(response, 'data-label="Действие"')
        self.assertContains(response, 'data-label="Период"')

    def test_specialist_home_includes_legacy_appointment_without_assignment(self):
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=1), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=1), time(10, 30)),
            billing_account=self.account,
        )
        AppointmentStaffAssignment.objects.filter(
            appointment=appointment,
            staff_member=self.staff,
        ).delete()
        self.client.logout()
        self.client.force_login(self.staff_user)

        response = self.client.get(reverse("specialist_home"))

        self.assertEqual(response.status_code, 200)
        self.assertIn(appointment, response.context["appointments"])
        self.assertContains(response, self.child.full_name)

    def test_specialist_without_mobile_access_is_redirected(self):
        StaffMember.objects.filter(pk=self.staff.pk).update(can_use_mobile=False)
        self.client.logout()
        self.client.force_login(self.staff_user)

        response = self.client.get(reverse("specialist_home"))

        self.assertEqual(response.status_code, 403)

    def test_specialist_without_mobile_access_cannot_mark_appointment(self):
        StaffMember.objects.filter(pk=self.staff.pk).update(can_use_mobile=False)
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate(), time(10, 0)),
            ends_at=_local_dt(timezone.localdate(), time(10, 30)),
            billing_account=self.account,
        )
        self.client.logout()
        self.client.force_login(self.staff_user)

        response = self.client.post(
            reverse("mark_appointment", args=[appointment.pk]), {"action": "completed"}
        )

        self.assertEqual(response.status_code, 403)
        appointment.refresh_from_db()
        self.assertEqual(appointment.status, Appointment.Status.CONFIRMED)


class AppointmentConfirmationSendTests(NewViewsTestBase):
    def test_send_confirmation_to_additional_schedule_representative(self):
        representative = ParentGuardian.objects.create(
            last_name="Петров",
            first_name="Сергей",
            relationship_type=ParentGuardian.RelationshipType.FATHER,
            phone="+7 900 000-00-02",
            email="father@example.local",
        )
        link = RecipientRepresentative.objects.create(
            child=self.child,
            representative=representative,
            relationship_type=ParentGuardian.RelationshipType.FATHER,
            receives_schedule=True,
        )
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
        )

        response = self.client.post(
            reverse("appointment_send_confirmation", args=[appointment.pk]),
            {
                "target_type": f"{AppointmentConfirmation.TargetType.REPRESENTATIVE}:{link.pk}",
                "subject": "Подтвердите занятие",
                "message": "Проверьте дату и подтвердите занятие.",
            },
        )

        self.assertRedirects(response, reverse("appointment_detail", args=[appointment.pk]))
        confirmation = AppointmentConfirmation.objects.get(appointment=appointment)
        self.assertEqual(
            confirmation.target_type, AppointmentConfirmation.TargetType.REPRESENTATIVE
        )
        self.assertEqual(confirmation.representative, representative)
        self.assertEqual(confirmation.participant.child, self.child)
        self.assertEqual(confirmation.email, "father@example.local")

    def test_send_group_confirmation_to_second_participant_representative(self):
        second_parent = ParentGuardian.objects.create(
            last_name="Смирнова",
            first_name="Ольга",
            relationship_type=ParentGuardian.RelationshipType.MOTHER,
            phone="+7 900 000-00-04",
            email="second-parent@example.local",
        )
        second_child = Child.objects.create(
            last_name="Смирнов",
            first_name="Илья",
            primary_parent=second_parent,
        )
        second_link = second_child.representative_links.get(representative=second_parent)
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
            session_type=Appointment.SessionType.GROUP,
        )
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=second_child,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )

        response = self.client.post(
            reverse("appointment_send_confirmation", args=[appointment.pk]),
            {
                "target_type": f"{AppointmentConfirmation.TargetType.REPRESENTATIVE}:{second_link.pk}",
                "subject": "Подтвердите занятие",
                "message": "Проверьте дату и подтвердите занятие.",
            },
        )

        self.assertRedirects(response, reverse("appointment_detail", args=[appointment.pk]))
        confirmation = AppointmentConfirmation.objects.get(appointment=appointment)
        self.assertEqual(confirmation.representative, second_parent)
        self.assertEqual(confirmation.participant.child, second_child)
        self.assertEqual(confirmation.email, "second-parent@example.local")

    def test_default_group_confirmation_message_is_scoped_to_target_participant(self):
        second_parent = ParentGuardian.objects.create(
            last_name="Смирнова",
            first_name="Ольга",
            relationship_type=ParentGuardian.RelationshipType.MOTHER,
            phone="+7 900 000-00-04",
            email="second-parent@example.local",
        )
        second_child = Child.objects.create(
            last_name="Смирнов",
            first_name="Илья",
            primary_parent=second_parent,
        )
        second_link = second_child.representative_links.get(representative=second_parent)
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
            session_type=Appointment.SessionType.GROUP,
        )
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=second_child,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )
        form = AppointmentConfirmationSendForm(appointment=appointment)

        response = self.client.post(
            reverse("appointment_send_confirmation", args=[appointment.pk]),
            {
                "target_type": f"{AppointmentConfirmation.TargetType.REPRESENTATIVE}:{second_link.pk}",
                "subject": "Подтвердите занятие",
                "message": form.initial["message"],
            },
        )

        self.assertRedirects(response, reverse("appointment_detail", args=[appointment.pk]))
        confirmation = AppointmentConfirmation.objects.get(appointment=appointment)
        self.assertIn(second_child.full_name, confirmation.message)
        self.assertNotIn(self.child.full_name, confirmation.message)

    def test_same_representative_can_be_targeted_per_group_participant(self):
        second_child = Child.objects.create(
            last_name="Sibling",
            first_name="Second",
            primary_parent=self.parent,
        )
        second_link = second_child.representative_links.get(representative=self.parent)
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
            session_type=Appointment.SessionType.GROUP,
        )
        participant = AppointmentParticipant.objects.create(
            appointment=appointment,
            child=second_child,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )
        form = AppointmentConfirmationSendForm(appointment=appointment)
        target_key = f"{AppointmentConfirmation.TargetType.REPRESENTATIVE}:{second_link.pk}"

        self.assertIn(target_key, form.targets)

        response = self.client.post(
            reverse("appointment_send_confirmation", args=[appointment.pk]),
            {
                "target_type": target_key,
                "subject": "Confirm session",
                "message": form.initial["message"],
            },
        )

        self.assertRedirects(response, reverse("appointment_detail", args=[appointment.pk]))
        confirmation = AppointmentConfirmation.objects.get(appointment=appointment)
        self.assertEqual(confirmation.representative, self.parent)
        self.assertEqual(confirmation.participant, participant)
        self.assertIn(second_child.full_name, confirmation.message)
        self.assertNotIn(self.child.full_name, confirmation.message)

    def test_send_group_confirmation_to_assistant_specialist(self):
        assistant = StaffMember.objects.create(
            full_name="Assistant Specialist",
            email="assistant@example.local",
        )
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
            session_type=Appointment.SessionType.GROUP,
        )
        assignment = AppointmentStaffAssignment.objects.create(
            appointment=appointment,
            staff_member=assistant,
            role=AppointmentStaffAssignment.Role.ASSISTANT,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )

        response = self.client.post(
            reverse("appointment_send_confirmation", args=[appointment.pk]),
            {
                "target_type": f"{AppointmentConfirmation.TargetType.SPECIALIST}:{assignment.pk}",
                "subject": "Confirm session",
                "message": "Please confirm the session.",
            },
        )

        self.assertRedirects(response, reverse("appointment_detail", args=[appointment.pk]))
        confirmation = AppointmentConfirmation.objects.get(appointment=appointment)
        self.assertEqual(confirmation.target_type, AppointmentConfirmation.TargetType.SPECIALIST)
        self.assertEqual(confirmation.staff_assignment, assignment)
        self.assertEqual(confirmation.email, "assistant@example.local")

    def test_send_group_confirmation_to_second_recipient_tracks_participant(self):
        second_child = Child.objects.create(
            last_name="Recipient",
            first_name="Second",
            email="second-recipient@example.local",
        )
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
            session_type=Appointment.SessionType.GROUP,
        )
        participant = AppointmentParticipant.objects.create(
            appointment=appointment,
            child=second_child,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )

        response = self.client.post(
            reverse("appointment_send_confirmation", args=[appointment.pk]),
            {
                "target_type": f"{AppointmentConfirmation.TargetType.RECIPIENT}:{second_child.pk}",
                "subject": "Confirm session",
                "message": "Please confirm the session.",
            },
        )

        self.assertRedirects(response, reverse("appointment_detail", args=[appointment.pk]))
        confirmation = AppointmentConfirmation.objects.get(appointment=appointment)
        self.assertEqual(confirmation.target_type, AppointmentConfirmation.TargetType.RECIPIENT)
        self.assertEqual(confirmation.participant, participant)
        self.assertEqual(confirmation.email, "second-recipient@example.local")

        self.client.logout()
        public_response = self.client.get(
            reverse("appointment_confirmation_public", args=[confirmation.token])
        )

        self.assertContains(public_response, second_child.full_name)


class ConfirmationPublicViewTests(NewViewsTestBase):
    def test_get_renders_page(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
        )
        conf = AppointmentConfirmation.objects.create(
            appointment=appt,
            target_type=AppointmentConfirmation.TargetType.REPRESENTATIVE,
            representative=self.parent,
            email=self.parent.email,
            subject="Тест",
            message="Тест",
        )
        self.client.logout()
        response = self.client.get(reverse("appointment_confirmation_public", args=[conf.token]))
        self.assertEqual(response.status_code, 200)
        self.assertIn("confirmation_summary_items", response.context)
        self.assertIn("confirmation_control_items", response.context)
        self.assertIn("confirmation_next_action", response.context)
        self.assertContains(response, "Следующий шаг")
        self.assertContains(response, "Контроль согласования")
        self.assertContains(response, 'id="confirmation-form"')

    def test_group_confirmation_public_shows_participants_and_staff_assignments(self):
        assistant = StaffMember.objects.create(
            full_name="Assistant Public Specialist",
            email="assistant-public@example.local",
        )
        second_child = Child.objects.create(last_name="Public", first_name="Second")
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
            session_type=Appointment.SessionType.GROUP,
        )
        for child in (self.child, second_child):
            AppointmentParticipant.objects.update_or_create(
                appointment=appt,
                child=child,
                defaults={
                    "starts_at_snapshot": appt.starts_at,
                    "ends_at_snapshot": appt.ends_at,
                    "appointment_status": appt.status,
                },
            )
        for staff, role in (
            (self.staff, AppointmentStaffAssignment.Role.PRIMARY),
            (assistant, AppointmentStaffAssignment.Role.ASSISTANT),
        ):
            AppointmentStaffAssignment.objects.update_or_create(
                appointment=appt,
                staff_member=staff,
                defaults={
                    "role": role,
                    "starts_at_snapshot": appt.starts_at,
                    "ends_at_snapshot": appt.ends_at,
                    "appointment_status": appt.status,
                },
            )
        conf = AppointmentConfirmation.objects.create(
            appointment=appt,
            target_type=AppointmentConfirmation.TargetType.REPRESENTATIVE,
            representative=self.parent,
            email=self.parent.email,
            subject="Тест",
            message="Тест",
        )

        self.client.logout()
        response = self.client.get(reverse("appointment_confirmation_public", args=[conf.token]))

        self.assertContains(response, self.child.full_name)
        self.assertContains(response, second_child.full_name)
        self.assertContains(response, self.staff.full_name)
        self.assertContains(response, assistant.full_name)


class SpecialistActionsTests(NewViewsTestBase):
    def test_mark_appointment_completed(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() - timedelta(days=1), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() - timedelta(days=1), time(10, 30)),
            billing_account=self.account,
        )
        response = self.client.post(
            reverse("mark_appointment", args=[appt.pk]),
            {"action": "completed", "specialist_note": "Ок"},
        )
        self.assertEqual(response.status_code, 302)
        appt.refresh_from_db()
        self.assertEqual(appt.status, Appointment.Status.COMPLETED)
        self.assertEqual(appt.attendance_status, Appointment.AttendanceStatus.ATTENDED)

    def test_mark_appointment_backfills_legacy_snapshots(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() - timedelta(days=1), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() - timedelta(days=1), time(10, 30)),
            billing_account=self.account,
        )
        appt.participants.all().delete()
        appt.staff_assignments.all().delete()

        response = self.client.post(
            reverse("mark_appointment", args=[appt.pk]),
            {"action": "completed", "specialist_note": "Готово"},
        )

        self.assertEqual(response.status_code, 302)
        participant = AppointmentParticipant.objects.get(appointment=appt, child=self.child)
        self.assertEqual(participant.appointment_status, Appointment.Status.COMPLETED)
        self.assertEqual(participant.attendance_status, Appointment.AttendanceStatus.ATTENDED)
        self.assertEqual(participant.specialist_note, "Готово")
        self.assertIsNotNone(participant.marked_by_staff_at)
        staff_assignment = AppointmentStaffAssignment.objects.get(
            appointment=appt,
            staff_member=self.staff,
        )
        self.assertEqual(staff_assignment.appointment_status, Appointment.Status.COMPLETED)
        self.assertEqual(staff_assignment.starts_at_snapshot, appt.starts_at)
        self.assertEqual(staff_assignment.ends_at_snapshot, appt.ends_at)

    def test_mark_appointment_no_show(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() - timedelta(days=1), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() - timedelta(days=1), time(10, 30)),
            billing_account=self.account,
        )
        response = self.client.post(
            reverse("mark_appointment", args=[appt.pk]),
            {"action": "not_completed"},
        )
        self.assertEqual(response.status_code, 302)
        appt.refresh_from_db()
        self.assertEqual(appt.status, Appointment.Status.NO_SHOW)

    def test_staff_availability_create(self):
        from operations.models import StaffAvailability

        self.client.logout()
        self.client.force_login(self.staff_user)
        response = self.client.post(
            reverse("staff_availability_create"),
            {
                "weekday": 0,
                "starts_at": "09:00",
                "ends_at": "17:00",
                "note": "",
            },
        )
        self.assertEqual(response.status_code, 302)
        self.assertTrue(StaffAvailability.objects.filter(staff_member=self.staff).exists())

    def test_staff_availability_toggle(self):
        from operations.models import StaffAvailability

        avail = StaffAvailability.objects.create(
            staff_member=self.staff,
            weekday=0,
            starts_at=time(9, 0),
            ends_at=time(17, 0),
            is_active=True,
        )
        response = self.client.post(reverse("staff_availability_toggle", args=[avail.pk]))
        self.assertEqual(response.status_code, 302)
        avail.refresh_from_db()
        self.assertFalse(avail.is_active)

    def test_time_off_request_create(self):
        from operations.models import TimeOffRequest

        self.client.logout()
        self.client.force_login(self.staff_user)
        starts_on = timezone.localdate() + timedelta(days=10)
        response = self.client.post(
            reverse("time_off_request_create"),
            {
                "request_type": TimeOffRequest.RequestType.DAY_OFF,
                "starts_on": starts_on.isoformat(),
                "ends_on": starts_on.isoformat(),
                "reason": "Личные обстоятельства",
            },
        )
        self.assertEqual(response.status_code, 302)
        self.assertTrue(TimeOffRequest.objects.filter(staff_member=self.staff).exists())

    def test_time_off_decide_reject(self):
        from operations.models import TimeOffRequest

        starts_on = timezone.localdate() + timedelta(days=10)
        req = TimeOffRequest.objects.create(
            staff_member=self.staff,
            request_type=TimeOffRequest.RequestType.DAY_OFF,
            starts_on=starts_on,
            ends_on=starts_on,
            reason="X",
        )
        response = self.client.post(
            reverse("time_off_request_decide", args=[req.pk]),
            {"action": "reject", "admin_note": "Не сегодня"},
        )
        self.assertEqual(response.status_code, 302)
        req.refresh_from_db()
        self.assertEqual(req.status, TimeOffRequest.Status.REJECTED)


class DrainTasksCommandTests(NewViewsTestBase):
    def test_drain_empty_queue(self):
        from django.core.management import call_command

        call_command("drain_tasks", "--once", verbosity=0)


class SuggestedTransferSlotsTests(NewViewsTestBase):
    def test_returns_empty_when_no_active_staff(self):
        StaffMember.objects.update(status=StaffMember.Status.INACTIVE)
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
        )
        from operations.views.scheduling_helpers import suggested_transfer_slots

        slots = suggested_transfer_slots(appt, days=2, limit=5)
        self.assertEqual(slots, [])

    def test_returns_alternate_slots(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
        )
        from operations.views.scheduling_helpers import suggested_transfer_slots

        slots = suggested_transfer_slots(appt, days=3, limit=10)
        self.assertGreater(len(slots), 0)
        # All returned slots must NOT be the original (day, time) of the appointment.
        for slot in slots:
            if (
                slot["staff"] == self.staff
                and slot["time"] == "10:00"
                and slot["date"] == appt.starts_at.date()
            ):
                self.fail("Original slot should be excluded")

    def test_skips_conflicts(self):
        # Create a conflicting appointment on day+1 at 11:00 for the same staff
        day = timezone.localdate() + timedelta(days=6)
        Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(day, time(11, 0)),
            ends_at=_local_dt(day, time(11, 30)),
            billing_account=self.account,
        )
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
        )
        from operations.views.scheduling_helpers import suggested_transfer_slots

        slots = suggested_transfer_slots(appt, days=8, limit=20)
        # None of the returned slots should be at 11:00 on the conflict day for this staff
        for slot in slots:
            if slot["staff"] == self.staff and slot["date"] == day and slot["time"] == "11:00":
                self.fail("Conflicting slot should be excluded")

    def test_skips_group_participant_and_assistant_conflicts(self):
        self.room.limit_staff_count = False
        self.room.limit_recipient_count = False
        self.room.allow_group_sessions = True
        self.room.save(
            update_fields=["limit_staff_count", "limit_recipient_count", "allow_group_sessions"]
        )
        assistant = StaffMember.objects.create(
            full_name="Ассистент группы",
            specializations="Психолог",
            status=StaffMember.Status.ACTIVE,
        )
        second_parent = ParentGuardian.objects.create(
            last_name="Петрова",
            first_name="Мария",
            phone="+7 900 000-30-01",
        )
        second_child = Child.objects.create(
            last_name="Петров",
            first_name="Илья",
            primary_parent=second_parent,
        )
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
            session_type=Appointment.SessionType.GROUP,
        )
        AppointmentParticipant.objects.create(
            appointment=appt,
            child=second_child,
            starts_at_snapshot=appt.starts_at,
            ends_at_snapshot=appt.ends_at,
            appointment_status=appt.status,
        )
        AppointmentStaffAssignment.objects.create(
            appointment=appt,
            staff_member=assistant,
            role=AppointmentStaffAssignment.Role.ASSISTANT,
            starts_at_snapshot=appt.starts_at,
            ends_at_snapshot=appt.ends_at,
            appointment_status=appt.status,
        )
        conflict_day = timezone.localdate() + timedelta(days=6)
        Appointment.objects.create(
            child=second_child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(conflict_day, time(11, 0)),
            ends_at=_local_dt(conflict_day, time(11, 30)),
        )
        Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=assistant,
            room=self.room,
            starts_at=_local_dt(conflict_day, time(12, 0)),
            ends_at=_local_dt(conflict_day, time(12, 30)),
        )
        from operations.views.scheduling_helpers import suggested_transfer_slots

        slots = suggested_transfer_slots(appt, days=8, limit=40)

        for slot in slots:
            if slot["date"] == appt.starts_at.date() and slot["time"] == "10:00":
                self.fail("Original group slot should be excluded for every assigned staff member")
            if slot["date"] == conflict_day and slot["time"] in {"11:00", "12:00"}:
                self.fail("Group participant/assistant conflict should be excluded")

    def test_includes_legacy_child_when_group_snapshot_is_partial(self):
        other_staff = StaffMember.objects.create(
            full_name="Другой специалист",
            specializations="Логопед",
            status=StaffMember.Status.ACTIVE,
        )
        other_room = Room.objects.create(name="Кабинет 2")
        second_parent = ParentGuardian.objects.create(
            last_name="Петрова",
            first_name="Мария",
            phone="+7 900 000-31-01",
        )
        second_child = Child.objects.create(
            last_name="Петров",
            first_name="Илья",
            primary_parent=second_parent,
        )
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
            session_type=Appointment.SessionType.GROUP,
        )
        AppointmentParticipant.objects.create(
            appointment=appt,
            child=second_child,
            starts_at_snapshot=appt.starts_at,
            ends_at_snapshot=appt.ends_at,
            appointment_status=appt.status,
        )
        conflict_day = timezone.localdate() + timedelta(days=6)
        Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=other_staff,
            room=other_room,
            starts_at=_local_dt(conflict_day, time(11, 0)),
            ends_at=_local_dt(conflict_day, time(11, 30)),
        )
        from operations.views.scheduling_helpers import suggested_transfer_slots

        slots = suggested_transfer_slots(appt, days=8, limit=80)

        for slot in slots:
            if (
                slot["staff"] == self.staff
                and slot["date"] == conflict_day
                and slot["time"] == "11:00"
            ):
                self.fail("Legacy child conflict should be excluded for partial snapshots")

    def test_respects_limit(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
        )
        from operations.views.scheduling_helpers import suggested_transfer_slots

        slots = suggested_transfer_slots(appt, days=7, limit=3)
        self.assertLessEqual(len(slots), 3)

    def test_returns_shift_candidates_when_no_free_slots(self):
        day = timezone.localdate() + timedelta(days=5)
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(day, time(10, 0)),
            ends_at=_local_dt(day, time(10, 30)),
            billing_account=self.account,
        )
        other_parent = ParentGuardian.objects.create(
            last_name="Сидорова",
            first_name="Анна",
            phone="+7 900 000-40-01",
        )
        other_child = Child.objects.create(
            last_name="Сидоров",
            first_name="Максим",
            primary_parent=other_parent,
        )
        for minute in range(9 * 60, 18 * 60, 30):
            if minute == 10 * 60:
                continue
            hour, clock_minute = divmod(minute, 60)
            starts_at = _local_dt(day, time(hour, clock_minute))
            Appointment.objects.create(
                child=other_child,
                service=self.service,
                staff_member=self.staff,
                room=self.room,
                starts_at=starts_at,
                ends_at=starts_at + timedelta(minutes=30),
            )
        from operations.views.scheduling_helpers import (
            suggested_shift_candidates,
            suggested_transfer_slots,
        )

        slots = suggested_transfer_slots(appt, days=1, limit=3)
        candidates = suggested_shift_candidates(appt, days=1, limit=3)

        self.assertEqual(slots, [])
        self.assertGreater(len(candidates), 0)
        self.assertTrue(candidates[0]["conflicts"])
        self.assertIn("специалист уже занят в это время", candidates[0]["messages"])


class ProgramBlockWizardViewTests(NewViewsTestBase):
    def setUp(self):
        super().setUp()
        self.program = TreatmentProgram.objects.create(
            child=self.child,
            title="Программа занятий",
            status=TreatmentProgram.Status.ACTIVE,
        )
        self.block = ProgramBlock.objects.create(
            program=self.program,
            number=1,
            title="Логопедический каскад",
            service=self.service,
            staff_member=self.staff,
            planned_sessions=3,
            balance_account=self.account,
        )

    def _wizard_payload(self, day, action="preview"):
        return {
            "start_date": day.isoformat(),
            "end_date": day.isoformat(),
            "weekdays": [str(day.weekday())],
            "time_from": "10:00",
            "time_until": "12:00",
            "duration_minutes": "30",
            "requested_count": "2",
            "staff_member": str(self.staff.pk),
            "room": str(self.room.pk),
            "appointment_status": Appointment.Status.PROPOSED,
            "action": action,
        }

    def test_program_object_forms_show_operator_control(self):
        cases = [
            (
                reverse("program_create_for_child", args=[self.child.pk]),
                "Контроль программы",
                "Каскады и серии",
            ),
            (
                reverse("program_block_create", args=[self.program.pk]),
                "Контроль каскада",
                "План и счет",
            ),
        ]
        for url, title, control_item in cases:
            with self.subTest(url=url):
                response = self.client.get(url)
                self.assertEqual(response.status_code, 200)
                self.assertIn("object_form_control_items", response.context)
                self.assertContains(response, 'class="object-form-layout"')
                self.assertContains(response, title)
                self.assertContains(response, control_item)

    def test_schedule_wizard_get(self):
        response = self.client.get(reverse("program_block_schedule_wizard", args=[self.block.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertIn("form", response.context)
        self.assertEqual(response.context["block"], self.block)
        self.assertContains(response, "Период и частота")
        self.assertContains(response, "Окно поиска")
        self.assertContains(response, "Сводка каскада")
        self.assertContains(response, "осталось по плану: 3")
        self.assertContains(response, "Критичных предупреждений сейчас нет")

    def test_schedule_wizard_preview_shows_operational_summary(self):
        day = timezone.localdate() + timedelta(days=8)
        response = self.client.post(
            reverse("program_block_schedule_wizard", args=[self.block.pk]),
            self._wizard_payload(day, action="preview"),
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Предложенные окна")
        self.assertContains(response, "запрошено: 2")
        self.assertContains(response, "доступно по оплате")
        self.assertContains(response, "Предложено / на согласование")
        self.assertContains(response, "Конфликты:")
        self.assertContains(response, "program-wizard-table")
        self.assertContains(response, 'data-label="Дата и время"')
        self.assertContains(response, 'data-label="Вместимость"')
        self.assertContains(response, 'data-label="График"')

    def test_schedule_wizard_create_creates_appointments(self):
        day = timezone.localdate() + timedelta(days=8)
        response = self.client.post(
            reverse("program_block_schedule_wizard", args=[self.block.pk]),
            self._wizard_payload(day, action="create"),
        )
        self.assertEqual(response.status_code, 302)
        appointments = Appointment.objects.filter(program_block=self.block).order_by(
            "sequence_number"
        )
        self.assertEqual(appointments.count(), 2)
        self.assertEqual([appt.sequence_number for appt in appointments], [1, 2])

    def test_schedule_wizard_create_can_auto_pick_staff_and_room(self):
        day = timezone.localdate() + timedelta(days=8)
        payload = self._wizard_payload(day, action="create")
        payload["staff_member"] = ""
        payload["room"] = ""

        response = self.client.post(
            reverse("program_block_schedule_wizard", args=[self.block.pk]), payload
        )

        self.assertEqual(response.status_code, 302)
        appointment = Appointment.objects.get(program_block=self.block, sequence_number=1)
        self.assertIsNotNone(appointment.staff_member_id)
        self.assertIsNotNone(appointment.room_id)

    def test_transfer_funds_get_shows_operator_context(self):
        response = self.client.get(reverse("program_block_transfer_funds", args=[self.block.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertIn("transfer_summary_items", response.context)
        self.assertIn("transfer_control_items", response.context)
        self.assertIn("transfer_next_action", response.context)
        self.assertContains(response, "Следующий шаг")
        self.assertContains(response, "Контроль переноса средств")
        self.assertContains(response, 'id="transfer-form"')

    def test_transfer_funds_between_block_accounts(self):
        movable_source = FundingSource.objects.create(
            name="Переносимый грант",
            source_type=FundingSource.SourceType.GRANT,
            transfer_policy=FundingSource.TransferPolicy.WITHIN_CHILD,
        )
        source = BalanceAccount.objects.create(
            child=self.child,
            funding_source=movable_source,
            unit=BalanceAccount.Unit.SESSIONS,
            service_scope=BalanceAccount.ServiceScope.SPECIFIC_SERVICE,
            service=self.service2,
            initial_amount=Decimal("5"),
        )
        target = BalanceAccount.objects.create(
            child=self.child,
            funding_source=movable_source,
            unit=BalanceAccount.Unit.SESSIONS,
            service_scope=BalanceAccount.ServiceScope.SPECIFIC_SERVICE,
            service=self.service,
            initial_amount=Decimal("0"),
        )
        self.block.balance_account = target
        self.block.save(update_fields=["balance_account"])

        response = self.client.post(
            reverse("program_block_transfer_funds", args=[self.block.pk]),
            {
                "from_account": source.pk,
                "to_account": target.pk,
                "amount": "2",
                "reason": "Перенос тест",
            },
        )

        self.assertEqual(response.status_code, 302)
        source.refresh_from_db()
        target.refresh_from_db()
        self.assertEqual(source.current_balance, Decimal("3"))
        self.assertEqual(target.current_balance, Decimal("2"))


class AppointmentDetailAndMoveTests(NewViewsTestBase):
    def test_appointment_detail_renders_with_suggestions(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
        )
        response = self.client.get(reverse("appointment_detail", args=[appt.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertIn("appointment", response.context)
        self.assertIn("suggested_slots", response.context)

    def test_appointment_detail_uses_single_participant_billing_snapshot(self):
        participant_child = Child.objects.create(last_name="Карточка", first_name="Счет")
        participant_account = BalanceAccount.objects.create(
            child=participant_child,
            funding_source=self.funding,
            unit=BalanceAccount.Unit.SESSIONS,
            service_scope=BalanceAccount.ServiceScope.ANY,
            initial_amount=Decimal("5"),
        )
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=None,
        )
        appt.participants.filter(child=self.child).delete()
        participant = AppointmentParticipant.objects.create(
            appointment=appt,
            child=participant_child,
            billing_decision=Appointment.BillingDecision.CHARGE,
            billing_account=participant_account,
            attendance_status=Appointment.AttendanceStatus.MISSED,
            starts_at_snapshot=appt.starts_at,
            ends_at_snapshot=appt.ends_at,
            appointment_status=appt.status,
        )

        response = self.client.get(reverse("appointment_detail", args=[appt.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.context["appointment_billing_summary_label"],
            participant.get_billing_decision_display(),
        )
        self.assertEqual(
            response.context["appointment_attendance_summary_label"],
            participant.get_attendance_status_display(),
        )
        self.assertEqual(
            response.context["appointment_billing_account_label"],
            str(participant_account),
        )
        self.assertEqual(response.context["appointment_payment_account"], participant_account)
        self.assertContains(
            response,
            reverse("payment_create_for_account", args=[participant_account.pk]),
        )

    def test_appointment_detail_hides_payment_account_for_group(self):
        second_child = Child.objects.create(last_name="Оплата", first_name="Группа")
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
            session_type=Appointment.SessionType.GROUP,
        )
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=second_child,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )

        response = self.client.get(reverse("appointment_detail", args=[appointment.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertIsNone(response.context["appointment_payment_account"])

    def test_appointment_detail_related_appointments_follow_single_participant(self):
        participant_child = Child.objects.create(last_name="Связанные", first_name="Участник")
        day = timezone.localdate() + timedelta(days=5)
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(day, time(10, 0)),
            ends_at=_local_dt(day, time(10, 30)),
        )
        appointment.participants.filter(child=self.child).delete()
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=participant_child,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )
        legacy_related = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(day + timedelta(days=1), time(10, 0)),
            ends_at=_local_dt(day + timedelta(days=1), time(10, 30)),
        )
        participant_related = Appointment.objects.create(
            child=participant_child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(day + timedelta(days=2), time(10, 0)),
            ends_at=_local_dt(day + timedelta(days=2), time(10, 30)),
        )

        response = self.client.get(reverse("appointment_detail", args=[appointment.pk]))

        self.assertEqual(response.status_code, 200)
        related_ids = {item.pk for item in response.context["related_child_appointments"]}
        self.assertIn(participant_related.pk, related_ids)
        self.assertNotIn(legacy_related.pk, related_ids)

    def test_appointment_detail_summarizes_group_attendance(self):
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            attendance_status=Appointment.AttendanceStatus.ATTENDED,
            session_type=Appointment.SessionType.GROUP,
        )
        AppointmentParticipant.objects.filter(appointment=appointment, child=self.child).update(
            attendance_status=Appointment.AttendanceStatus.ATTENDED,
        )
        second_child = Child.objects.create(last_name="Посещение", first_name="Группа")
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=second_child,
            attendance_status=Appointment.AttendanceStatus.UNKNOWN,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )

        response = self.client.get(reverse("appointment_detail", args=[appointment.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.context["appointment_attendance_summary_label"],
            "Есть неотмеченные участники: 1 из 2",
        )

    def test_appointment_detail_shows_audit_history(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
        )
        appt.title = "История занятия"
        appt.save(update_fields=["title", "updated_at"])

        response = self.client.get(reverse("appointment_detail", args=[appt.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "История изменений")
        audit_entries = response.context["audit_entries"]
        self.assertGreaterEqual(len(audit_entries), 1)
        self.assertIn(str(appt.pk), {entry.object_pk for entry in audit_entries})

    def test_appointment_detail_updates_group_participant_program_block(self):
        second_parent = ParentGuardian.objects.create(
            last_name="Петрова",
            first_name="Мария",
            phone="+7 900 000-30-01",
        )
        second_child = Child.objects.create(
            last_name="Петров",
            first_name="Илья",
            primary_parent=second_parent,
        )
        program = TreatmentProgram.objects.create(child=second_child, title="Групповая программа")
        block = ProgramBlock.objects.create(
            program=program,
            number=1,
            title="Коммуникация",
            service=self.service,
            planned_sessions=4,
        )
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            session_type=Appointment.SessionType.GROUP,
            title="Группа коммуникации",
        )
        participant = AppointmentParticipant.objects.create(
            appointment=appt,
            child=second_child,
            starts_at_snapshot=appt.starts_at,
            ends_at_snapshot=appt.ends_at,
            appointment_status=appt.status,
        )

        detail = self.client.get(reverse("appointment_detail", args=[appt.pk]))
        self.assertContains(detail, "Каскады участников")

        response = self.client.post(
            reverse("appointment_participant_program", args=[appt.pk]),
            {
                "participant_id": participant.pk,
                "program_block": block.pk,
                "sequence_number": "",
            },
        )

        self.assertRedirects(response, reverse("appointment_detail", args=[appt.pk]))
        participant.refresh_from_db()
        self.assertEqual(participant.program_block, block)
        self.assertEqual(participant.sequence_number, 1)
        self.assertEqual(block.scheduled_count, 1)

    def test_appointment_detail_rejects_participant_program_block_for_other_child(self):
        second_parent = ParentGuardian.objects.create(
            last_name="Петрова",
            first_name="Мария",
            phone="+7 900 000-30-01",
        )
        second_child = Child.objects.create(
            last_name="Петров",
            first_name="Илья",
            primary_parent=second_parent,
        )
        other_program = TreatmentProgram.objects.create(child=self.child, title="Чужая программа")
        other_block = ProgramBlock.objects.create(
            program=other_program,
            number=1,
            title="Чужой блок",
            service=self.service,
            planned_sessions=4,
        )
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            session_type=Appointment.SessionType.GROUP,
            title="Группа коммуникации",
        )
        participant = AppointmentParticipant.objects.create(
            appointment=appt,
            child=second_child,
            starts_at_snapshot=appt.starts_at,
            ends_at_snapshot=appt.ends_at,
            appointment_status=appt.status,
        )

        response = self.client.post(
            reverse("appointment_participant_program", args=[appt.pk]),
            {
                "participant_id": participant.pk,
                "program_block": other_block.pk,
                "sequence_number": "",
            },
        )

        self.assertEqual(response.status_code, 400)
        participant.refresh_from_db()
        self.assertIsNone(participant.program_block)
        self.assertIsNone(participant.sequence_number)

    def test_appointment_create_get_with_params(self):
        response = self.client.get(
            reverse("appointment_create"),
            {
                "child_id": self.child.pk,
                "service_id": self.service.pk,
                "staff_id": self.staff.pk,
                "room_id": self.room.pk,
                "date": timezone.localdate().isoformat(),
                "time": "10:00",
            },
        )
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Формат и время")
        self.assertContains(response, "Участники")
        self.assertContains(response, "Сводка")
        self.assertContains(response, "получателей: 1 · специалистов: 1")
        self.assertContains(response, "Кабинет 1")
        self.assertContains(response, "10:00")

    def test_appointment_edit_get(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
        )
        response = self.client.get(reverse("appointment_edit", args=[appt.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Программа и оплата")
        self.assertContains(response, "редактирование")
        self.assertContains(response, "получателей: 1 · специалистов: 1")
        self.assertContains(response, "Кабинет 1")

    def test_appointment_edit_post_valid(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
        )
        new_day = timezone.localdate() + timedelta(days=6)
        response = self.client.post(
            reverse("appointment_edit", args=[appt.pk]),
            {
                "child": self.child.id,
                "service": self.service.id,
                "staff_member": self.staff.id,
                "room": self.room.id,
                "billing_account": self.account.id,
                "status": Appointment.Status.CONFIRMED,
                "date": new_day.isoformat(),
                "time": "11:00",
                "duration_minutes": "30",
            },
        )
        self.assertEqual(response.status_code, 302)
        appt.refresh_from_db()
        self.assertEqual(timezone.localtime(appt.starts_at).date(), new_day)

    def test_appointment_move_get(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
        )
        response = self.client.get(reverse("appointment_move", args=[appt.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertIn("suggested_slots", response.context)
        self.assertIn("appointment_move_summary_items", response.context)
        self.assertIn("appointment_move_next_action", response.context)
        self.assertContains(response, "Контроль переноса")
        self.assertContains(response, 'id="appointment-move-form"')
        self.assertContains(response, 'id="transfer-slots"')
        self.assertContains(response, "Предложенные окна системы")
        self.assertContains(response, reverse("appointment_reschedule_plan_create", args=[appt.pk]))

    def test_group_appointment_move_rejects_same_time_room_with_assistant(self):
        assistant = StaffMember.objects.create(
            full_name="Move Assistant",
            status=StaffMember.Status.ACTIVE,
        )
        starts_at = _local_dt(timezone.localdate() + timedelta(days=5), time(10, 0))
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=starts_at,
            ends_at=starts_at + timedelta(minutes=30),
            billing_account=self.account,
            session_type=Appointment.SessionType.GROUP,
        )
        AppointmentStaffAssignment.objects.create(
            appointment=appt,
            staff_member=assistant,
            role=AppointmentStaffAssignment.Role.ASSISTANT,
            starts_at_snapshot=appt.starts_at,
            ends_at_snapshot=appt.ends_at,
            appointment_status=appt.status,
        )

        response = self.client.post(
            reverse("appointment_move", args=[appt.pk]),
            {
                "date": timezone.localtime(appt.starts_at).date().isoformat(),
                "time": "10:00",
                "duration_minutes": "30",
                "staff_member": assistant.pk,
                "room": self.room.pk,
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Новое время совпадает с текущим занятием")
        appt.refresh_from_db()
        self.assertEqual(appt.status, Appointment.Status.CONFIRMED)
        self.assertFalse(Appointment.objects.filter(source_appointment=appt).exists())

    def test_group_appointment_headers_show_participants_and_staff(self):
        assistant = StaffMember.objects.create(
            full_name="Group Assistant",
            status=StaffMember.Status.ACTIVE,
        )
        second_child = Child.objects.create(last_name="Group", first_name="Second")
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
            session_type=Appointment.SessionType.GROUP,
        )
        AppointmentParticipant.objects.create(
            appointment=appt,
            child=second_child,
            starts_at_snapshot=appt.starts_at,
            ends_at_snapshot=appt.ends_at,
            appointment_status=appt.status,
        )
        AppointmentStaffAssignment.objects.create(
            appointment=appt,
            staff_member=assistant,
            role=AppointmentStaffAssignment.Role.ASSISTANT,
            starts_at_snapshot=appt.starts_at,
            ends_at_snapshot=appt.ends_at,
            appointment_status=appt.status,
        )

        detail = self.client.get(reverse("appointment_detail", args=[appt.pk]))
        edit = self.client.get(reverse("appointment_edit", args=[appt.pk]))
        move = self.client.get(reverse("appointment_move", args=[appt.pk]))
        cancel = self.client.get(reverse("appointment_cancel", args=[appt.pk]))

        self.assertContains(detail, second_child.full_name)
        self.assertContains(detail, assistant.full_name)
        self.assertContains(edit, second_child.full_name)
        self.assertContains(move, second_child.full_name)
        self.assertContains(move, assistant.full_name)
        self.assertContains(cancel, second_child.full_name)

    def test_appointment_detail_shows_operator_summary_and_attention(self):
        assistant = StaffMember.objects.create(
            full_name="Сводка Ассистент",
            status=StaffMember.Status.ACTIVE,
        )
        second_child = Child.objects.create(last_name="Сводка", first_name="Участник")
        starts_at = _local_dt(timezone.localdate() - timedelta(days=1), time(10, 0))
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=starts_at,
            ends_at=starts_at + timedelta(minutes=30),
            billing_account=self.account,
            session_type=Appointment.SessionType.GROUP,
            title="Группа сводки",
        )
        AppointmentParticipant.objects.create(
            appointment=appt,
            child=second_child,
            starts_at_snapshot=appt.starts_at,
            ends_at_snapshot=appt.ends_at,
            appointment_status=appt.status,
        )
        AppointmentStaffAssignment.objects.create(
            appointment=appt,
            staff_member=assistant,
            role=AppointmentStaffAssignment.Role.ASSISTANT,
            starts_at_snapshot=appt.starts_at,
            ends_at_snapshot=appt.ends_at,
            appointment_status=appt.status,
        )

        response = self.client.get(reverse("appointment_detail", args=[appt.pk]))

        self.assertContains(response, "Группа сводки")
        self.assertContains(response, "Получателей: 2 · специалистов: 2")
        self.assertContains(response, "Посещение не отмечено")
        self.assertContains(response, "Решение по списанию не принято")
        self.assertContains(response, "Нерешенных участников: 2")
        self.assertContains(response, "appointment-detail-table")
        self.assertContains(response, 'data-label="Получатель"')
        self.assertContains(response, 'data-label="График"')
        self.assertContains(response, 'data-label="Изменения"')

    def test_group_appointment_move_preserves_participants_staff_and_programs(self):
        self.room.limit_staff_count = False
        self.room.limit_recipient_count = False
        self.room.allow_group_sessions = True
        self.room.save(
            update_fields=["limit_staff_count", "limit_recipient_count", "allow_group_sessions"]
        )
        assistant = StaffMember.objects.create(
            full_name="Ассистент группы",
            specializations="Психолог",
            status=StaffMember.Status.ACTIVE,
        )
        second_parent = ParentGuardian.objects.create(
            last_name="Петрова",
            first_name="Мария",
            phone="+7 900 000-30-01",
        )
        second_child = Child.objects.create(
            last_name="Петров",
            first_name="Илья",
            primary_parent=second_parent,
        )
        program = TreatmentProgram.objects.create(child=second_child, title="Групповая программа")
        block = ProgramBlock.objects.create(
            program=program,
            number=1,
            title="Коммуникация",
            service=self.service,
            planned_sessions=4,
        )
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
            session_type=Appointment.SessionType.GROUP,
            title="Группа коммуникации",
        )
        participant = AppointmentParticipant.objects.create(
            appointment=appt,
            child=second_child,
            program_block=block,
            starts_at_snapshot=appt.starts_at,
            ends_at_snapshot=appt.ends_at,
            appointment_status=appt.status,
        )
        AppointmentStaffAssignment.objects.create(
            appointment=appt,
            staff_member=assistant,
            role=AppointmentStaffAssignment.Role.ASSISTANT,
            starts_at_snapshot=appt.starts_at,
            ends_at_snapshot=appt.ends_at,
            appointment_status=appt.status,
        )
        new_day = timezone.localdate() + timedelta(days=6)

        response = self.client.post(
            reverse("appointment_move", args=[appt.pk]),
            {
                "date": new_day.isoformat(),
                "time": "11:00",
                "duration_minutes": "30",
                "staff_member": self.staff.pk,
                "room": self.room.pk,
                "admin_note": "Перенос группы",
            },
        )

        self.assertEqual(response.status_code, 302)
        appt.refresh_from_db()
        self.assertEqual(appt.status, Appointment.Status.RESCHEDULED)
        participant.refresh_from_db()
        self.assertEqual(participant.appointment_status, Appointment.Status.RESCHEDULED)
        new_appt = Appointment.objects.get(source_appointment=appt)
        self.assertEqual(new_appt.session_type, Appointment.SessionType.GROUP)
        self.assertEqual(new_appt.participants.count(), 2)
        moved_participant = new_appt.participants.get(child=second_child)
        self.assertEqual(moved_participant.source_participant, participant)
        self.assertEqual(moved_participant.program_block, block)
        self.assertEqual(moved_participant.sequence_number, participant.sequence_number)
        self.assertEqual(moved_participant.billing_decision, Appointment.BillingDecision.UNDECIDED)
        self.assertCountEqual(
            new_appt.staff_assignments.values_list("staff_member_id", flat=True),
            [self.staff.pk, assistant.pk],
        )
        self.assertEqual(block.scheduled_count, 1)

    def test_appointment_cancel_get(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
        )
        response = self.client.get(reverse("appointment_cancel", args=[appt.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertIn("appointment_cancel_summary_items", response.context)
        self.assertIn("appointment_cancel_next_action", response.context)
        self.assertEqual(
            response.context["appointment_cancel_next_action"]["title"],
            "Сохранить отмену",
        )
        self.assertContains(response, "Контроль отмены")
        self.assertContains(response, 'id="appointment-cancel-form"')
        self.assertContains(response, "Решение по списанию принимается отдельно")

    def test_same_day_cancel_requires_billing_acknowledgement(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate(), time(10, 0)),
            ends_at=_local_dt(timezone.localdate(), time(10, 30)),
            billing_account=self.account,
        )

        response = self.client.post(
            reverse("appointment_cancel", args=[appt.pk]),
            {
                "status": Appointment.Status.CANCELLED,
                "reason": "representative_cancel",
                "admin_note": "Сообщили утром",
            },
        )

        self.assertEqual(response.status_code, 200)
        appt.refresh_from_db()
        self.assertEqual(appt.status, Appointment.Status.CONFIRMED)
        self.assertFormError(
            response.context["form"],
            "same_day_billing_ack",
            "Подтвердите отдельное решение по списанию для отмены день-в-день.",
        )

    def test_same_day_cancel_with_acknowledgement_keeps_billing_for_admin(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate(), time(10, 0)),
            ends_at=_local_dt(timezone.localdate(), time(10, 30)),
            billing_account=self.account,
        )

        response = self.client.post(
            reverse("appointment_cancel", args=[appt.pk]),
            {
                "status": Appointment.Status.CANCELLED,
                "reason": "representative_cancel",
                "same_day_billing_ack": "on",
                "admin_note": "Сообщили утром",
            },
        )

        self.assertRedirects(response, reverse("appointment_detail", args=[appt.pk]))
        appt.refresh_from_db()
        self.assertEqual(appt.status, Appointment.Status.CANCELLED)
        self.assertEqual(appt.billing_decision, Appointment.BillingDecision.UNDECIDED)
        self.assertIn("Отмена день-в-день", appt.admin_note)


class ReschedulePlanViewTests(NewViewsTestBase):
    def _appointment(self, *, days=5):
        starts_at = _local_dt(timezone.localdate() + timedelta(days=days), time(10, 0))
        return Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=starts_at,
            ends_at=starts_at + timedelta(minutes=30),
            billing_account=self.account,
        )

    def _reschedule_chain_fixture(self, *, days=5):
        first_source = self._appointment(days=days)
        second_start = _local_dt(timezone.localdate() + timedelta(days=days), time(11, 0))
        second_source = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=second_start,
            ends_at=second_start + timedelta(minutes=30),
            billing_account=self.account,
        )
        plan = AppointmentReschedulePlan.objects.create(
            status=AppointmentReschedulePlan.Status.READY,
            plan_type=AppointmentReschedulePlan.PlanType.MANUAL,
            reason="Chain fixture",
            created_by=self.admin,
        )
        first = AppointmentRescheduleStep.objects.create(
            plan=plan,
            position=1,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.VALID,
            source_appointment=second_source,
            proposed_starts_at=_local_dt(timezone.localdate() + timedelta(days=days), time(12, 0)),
            proposed_ends_at=_local_dt(timezone.localdate() + timedelta(days=days), time(12, 30)),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
        )
        second = AppointmentRescheduleStep.objects.create(
            plan=plan,
            position=2,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.VALID,
            source_appointment=first_source,
            proposed_starts_at=second_start,
            proposed_ends_at=second_start + timedelta(minutes=30),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
        )
        chain_result = plan_svc.create_chain_for_steps(
            plan,
            step_ids=[first.pk, second.pk],
            dependencies=[
                {
                    "predecessor_step_id": first.pk,
                    "successor_step_id": second.pk,
                    "reason": "first step frees second target slot",
                }
            ],
            title="Buffer chain",
            actor=self.admin,
        )
        return plan, chain_result.chain

    def _chain_with_status(self, status, title, *, days=5):
        plan, chain = self._reschedule_chain_fixture(days=days)
        chain.title = title
        chain.status = status
        chain.save(update_fields=["title", "status", "updated_at"])
        return plan, chain

    def test_appointment_detail_can_create_reschedule_plan(self):
        appointment = self._appointment()

        response = self.client.post(
            reverse("appointment_reschedule_plan_create", args=[appointment.pk])
        )

        plan = AppointmentReschedulePlan.objects.get(root_appointment=appointment)
        self.assertRedirects(response, reverse("appointment_reschedule_plan_detail", args=[plan.pk]))
        self.assertEqual(plan.status, AppointmentReschedulePlan.Status.READY)
        self.assertGreater(plan.steps.count(), 0)
        appointment.refresh_from_db()
        self.assertEqual(appointment.status, Appointment.Status.CONFIRMED)

    def test_reschedule_plan_list_shows_active_plan(self):
        appointment = self._appointment()
        plan = plan_svc.create_plan_for_appointment(
            appointment, actor=self.admin, days=2, limit=1
        )

        response = self.client.get(reverse("reschedule_plan_list"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Планы переноса")
        self.assertContains(response, "reschedule-plan-list-table")
        self.assertContains(response, appointment.child.full_name)
        self.assertContains(response, reverse("appointment_reschedule_plan_detail", args=[plan.pk]))
        self.assertContains(response, 'data-label="Согласования"')
        self.assertContains(response, 'data-label="Контроль"')
        self.assertContains(response, "Не запрошено")
        self.assertContains(response, "Можно применять: 1")

    def test_reschedule_plan_list_shows_period_metrics(self):
        appointment = self._appointment()
        current_plan = plan_svc.create_plan_for_appointment(
            appointment, actor=self.admin, days=2, limit=1
        )
        applied_plan = AppointmentReschedulePlan.objects.create(
            status=AppointmentReschedulePlan.Status.APPLIED,
            plan_type=AppointmentReschedulePlan.PlanType.MANUAL,
            reason="Старый план, примененный сейчас",
            created_by=self.admin,
            applied_by=self.admin,
            applied_at=timezone.now(),
        )
        AppointmentReschedulePlan.objects.filter(pk=applied_plan.pk).update(
            created_at=timezone.now() - timedelta(days=45)
        )
        old_cancelled_plan = AppointmentReschedulePlan.objects.create(
            status=AppointmentReschedulePlan.Status.CANCELLED,
            plan_type=AppointmentReschedulePlan.PlanType.MANUAL,
            reason="Старый отмененный план",
            created_by=self.admin,
            cancelled_by=self.admin,
            cancelled_at=timezone.now() - timedelta(days=45),
        )
        AppointmentReschedulePlan.objects.filter(pk=old_cancelled_plan.pk).update(
            created_at=timezone.now() - timedelta(days=45)
        )

        response = self.client.get(reverse("reschedule_plan_list"), {"metrics_period": "7"})

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Динамика")
        self.assertContains(response, "Период метрик")
        self.assertContains(response, "7 дней")
        metric_values = {item["label"]: item["value"] for item in response.context["metric_items"]}
        self.assertEqual(response.context["current_metrics_period"], "7")
        self.assertEqual(metric_values["Создано"], 1)
        self.assertEqual(metric_values["Шагов"], current_plan.steps.count())
        self.assertEqual(metric_values["Применено"], 1)
        self.assertEqual(metric_values["Отменено"], 0)

    def test_reschedule_plan_list_filters_waiting_confirmations(self):
        self.staff.email = "qa-plan-list-staff@example.local"
        self.staff.save(update_fields=["email", "updated_at"])
        appointment = self._appointment()
        plan = plan_svc.create_plan_for_appointment(
            appointment, actor=self.admin, days=2, limit=1
        )
        step = plan.steps.get()
        plan_svc.create_confirmations_for_step(step, actor=self.admin)

        response = self.client.get(
            reverse("reschedule_plan_list"),
            {"confirmation": AppointmentRescheduleStep.ConfirmationStatus.WAITING},
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Ждет ответов: 1")
        self.assertContains(response, reverse("appointment_reschedule_plan_detail", args=[plan.pk]))

    def test_reschedule_plan_list_filters_manual_review_focus(self):
        appointment = self._appointment()
        day = timezone.localtime(appointment.starts_at).date()
        plan = plan_svc.create_staff_absence_plan(
            self.staff,
            date_from=day,
            date_to=day,
            reason="Больничный",
            actor=self.admin,
        )

        response = self.client.get(reverse("reschedule_plan_list"), {"focus": "manual_review"})

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Ручной разбор")
        self.assertContains(response, "Ручной разбор: 1")
        self.assertContains(response, reverse("appointment_reschedule_plan_detail", args=[plan.pk]))

    def test_reschedule_plan_list_filters_stale_focus(self):
        appointment = self._appointment()
        plan = plan_svc.create_plan_for_appointment(
            appointment, actor=self.admin, days=2, limit=1
        )
        step = plan.steps.get()
        step.status = AppointmentRescheduleStep.Status.STALE
        step.save(update_fields=["status", "updated_at"])

        response = self.client.get(reverse("reschedule_plan_list"), {"focus": "stale"})

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Устарело: 1")
        self.assertContains(response, reverse("appointment_reschedule_plan_detail", args=[plan.pk]))

    def test_reschedule_plan_list_filters_ready_to_apply_focus(self):
        appointment = self._appointment()
        plan = plan_svc.create_plan_for_appointment(
            appointment, actor=self.admin, days=2, limit=1
        )

        response = self.client.get(reverse("reschedule_plan_list"), {"focus": "ready_to_apply"})

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Можно применять: 1")
        self.assertContains(response, reverse("appointment_reschedule_plan_detail", args=[plan.pk]))

    def test_reschedule_plan_list_shows_chain_focus_metrics(self):
        plan, chain = self._reschedule_chain_fixture()
        plan_svc.revalidate_chain(chain)

        response = self.client.get(reverse("reschedule_plan_list"), {"focus": "chain_ready"})

        self.assertEqual(response.status_code, 200)
        chain.refresh_from_db()
        self.assertEqual(chain.status, AppointmentRescheduleChain.Status.READY)
        summary_values = {item["label"]: item["value"] for item in response.context["summary_items"]}
        metric_values = {item["label"]: item["value"] for item in response.context["metric_items"]}
        self.assertEqual(summary_values["Цепочки"], 1)
        self.assertEqual(summary_values["Готовые цепочки"], 1)
        self.assertEqual(metric_values["Цепочек"], 1)
        self.assertContains(response, 'value="chain_ready" selected')
        self.assertContains(response, "Цепочка готова: 1")
        self.assertContains(response, "Открыть цепочку")
        self.assertContains(
            response,
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#chain-{chain.pk}',
        )

    def test_reschedule_plan_list_links_to_priority_chain_anchor(self):
        plan, ready_chain = self._chain_with_status(
            AppointmentRescheduleChain.Status.READY,
            "Ready chain",
            days=5,
        )
        stale_chain = AppointmentRescheduleChain.objects.create(
            plan=plan,
            title="Stale chain",
            status=AppointmentRescheduleChain.Status.STALE,
        )
        failed_chain = AppointmentRescheduleChain.objects.create(
            plan=plan,
            title="Failed chain",
            status=AppointmentRescheduleChain.Status.FAILED,
        )

        response = self.client.get(reverse("reschedule_plan_list"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Открыть цепочку")
        self.assertContains(
            response,
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#chain-{failed_chain.pk}',
        )
        self.assertNotContains(
            response,
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#chain-{stale_chain.pk}',
        )
        self.assertNotContains(
            response,
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#chain-{ready_chain.pk}',
        )

    def test_reschedule_plan_list_links_to_priority_step_anchor(self):
        source = self._appointment()
        plan = AppointmentReschedulePlan.objects.create(
            status=AppointmentReschedulePlan.Status.READY,
            plan_type=AppointmentReschedulePlan.PlanType.MANUAL,
            reason="Step attention fixture",
            created_by=self.admin,
        )
        ready_step = AppointmentRescheduleStep.objects.create(
            plan=plan,
            position=1,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.VALID,
            source_appointment=source,
            proposed_starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(12, 0)),
            proposed_ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(12, 30)),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
        )
        stale_step = AppointmentRescheduleStep.objects.create(
            plan=plan,
            position=2,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.STALE,
            source_appointment=source,
            proposed_starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(13, 0)),
            proposed_ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(13, 30)),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
            validation_messages=["slot changed"],
        )
        failed_step = AppointmentRescheduleStep.objects.create(
            plan=plan,
            position=3,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.FAILED,
            source_appointment=source,
            proposed_starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(14, 0)),
            proposed_ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(14, 30)),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
        )

        response = self.client.get(reverse("reschedule_plan_list"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Открыть шаг")
        self.assertContains(
            response,
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#step-{failed_step.pk}',
        )
        self.assertNotContains(
            response,
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#step-{stale_step.pk}',
        )
        self.assertNotContains(
            response,
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#step-{ready_step.pk}',
        )

    def test_reschedule_plan_list_filters_chain_attention_focuses(self):
        ready_plan, _ready_chain = self._chain_with_status(
            AppointmentRescheduleChain.Status.READY,
            "Ready chain",
            days=5,
        )
        stale_plan, _stale_chain = self._chain_with_status(
            AppointmentRescheduleChain.Status.STALE,
            "Stale chain",
            days=8,
        )
        failed_plan, _failed_chain = self._chain_with_status(
            AppointmentRescheduleChain.Status.FAILED,
            "Failed chain",
            days=11,
        )

        cases = [
            ("chain_ready", ready_plan),
            ("chain_stale", stale_plan),
            ("chain_failed", failed_plan),
        ]
        for focus, expected_plan in cases:
            with self.subTest(focus=focus):
                response = self.client.get(reverse("reschedule_plan_list"), {"focus": focus})

                self.assertEqual(response.status_code, 200)
                self.assertEqual(list(response.context["plans"]), [expected_plan])
                self.assertContains(response, f'value="{focus}" selected')

    def test_reschedule_plan_list_focus_excludes_terminal_plan_attention(self):
        chain_plan, _chain = self._chain_with_status(
            AppointmentRescheduleChain.Status.READY,
            "Terminal ready chain",
            days=5,
        )
        chain_plan.status = AppointmentReschedulePlan.Status.APPLIED
        chain_plan.save(update_fields=["status", "updated_at"])
        source = self._appointment(days=9)
        step_plan = AppointmentReschedulePlan.objects.create(
            status=AppointmentReschedulePlan.Status.CANCELLED,
            plan_type=AppointmentReschedulePlan.PlanType.MANUAL,
            reason="Terminal failed step fixture",
            created_by=self.admin,
        )
        AppointmentRescheduleStep.objects.create(
            plan=step_plan,
            position=1,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.FAILED,
            source_appointment=source,
            proposed_starts_at=_local_dt(timezone.localdate() + timedelta(days=9), time(12, 0)),
            proposed_ends_at=_local_dt(timezone.localdate() + timedelta(days=9), time(12, 30)),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
        )

        cases = [
            (AppointmentReschedulePlan.Status.APPLIED, "chain_ready"),
            (AppointmentReschedulePlan.Status.CANCELLED, "failed"),
        ]
        for status, focus in cases:
            with self.subTest(status=status, focus=focus):
                response = self.client.get(
                    reverse("reschedule_plan_list"),
                    {"status": status, "focus": focus},
                )

                self.assertEqual(response.status_code, 200)
                self.assertEqual(list(response.context["plans"]), [])
                self.assertContains(response, "Планов переноса по выбранным фильтрам нет.")

    def test_reschedule_plan_list_terminal_rows_are_read_only_history(self):
        chain_plan, chain = self._chain_with_status(
            AppointmentRescheduleChain.Status.READY,
            "Terminal history chain",
            days=5,
        )
        chain_plan.status = AppointmentReschedulePlan.Status.APPLIED
        chain_plan.save(update_fields=["status", "updated_at"])
        chain_plan.steps.update(
            confirmation_status=AppointmentRescheduleStep.ConfirmationStatus.WAITING
        )
        source = self._appointment(days=9)
        step_plan = AppointmentReschedulePlan.objects.create(
            status=AppointmentReschedulePlan.Status.APPLIED,
            plan_type=AppointmentReschedulePlan.PlanType.MANUAL,
            reason="Terminal history step fixture",
            created_by=self.admin,
        )
        failed_step = AppointmentRescheduleStep.objects.create(
            plan=step_plan,
            position=1,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.FAILED,
            confirmation_status=AppointmentRescheduleStep.ConfirmationStatus.DECLINED,
            source_appointment=source,
            proposed_starts_at=_local_dt(timezone.localdate() + timedelta(days=9), time(12, 0)),
            proposed_ends_at=_local_dt(timezone.localdate() + timedelta(days=9), time(12, 30)),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
        )

        response = self.client.get(
            reverse("reschedule_plan_list"),
            {"status": AppointmentReschedulePlan.Status.APPLIED},
        )

        self.assertEqual(response.status_code, 200)
        plans = list(response.context["plans"])
        self.assertEqual(set(plans), {chain_plan, step_plan})
        for plan in plans:
            with self.subTest(plan=plan.pk):
                self.assertTrue(plan.is_terminal)
                self.assertIsNone(plan.primary_attention_chain)
                self.assertIsNone(plan.primary_attention_step)
        self.assertContains(response, "Архив согласований", count=2)
        self.assertContains(response, "План завершен", count=2)
        self.assertContains(
            response,
            reverse("appointment_reschedule_plan_detail", args=[chain_plan.pk]),
        )
        self.assertContains(
            response,
            reverse("appointment_reschedule_plan_detail", args=[step_plan.pk]),
        )
        self.assertNotContains(response, "Открыть цепочку")
        self.assertNotContains(response, "Открыть шаг")
        self.assertNotContains(response, "Ждет ответов:")
        self.assertNotContains(response, "Есть отказ:")
        self.assertNotContains(
            response,
            f'{reverse("appointment_reschedule_plan_detail", args=[chain_plan.pk])}#chain-{chain.pk}',
        )
        self.assertNotContains(
            response,
            f'{reverse("appointment_reschedule_plan_detail", args=[step_plan.pk])}#step-{failed_step.pk}',
        )

    def test_dashboard_surfaces_ready_reschedule_chains(self):
        _plan, chain = self._reschedule_chain_fixture()
        plan_svc.revalidate_chain(chain)

        response = self.client.get(reverse("dashboard"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context["ready_chain_count"], 1)
        self.assertEqual(response.context["chain_attention_count"], 1)
        self.assertContains(response, "Применить цепочки")
        self.assertContains(response, "?focus=chain_ready")
        self.assertContains(response, reverse("reschedule_plan_list"))

    def test_dashboard_surfaces_reschedule_step_attention(self):
        source = self._appointment()
        plan = AppointmentReschedulePlan.objects.create(
            status=AppointmentReschedulePlan.Status.READY,
            plan_type=AppointmentReschedulePlan.PlanType.MANUAL,
            reason="Dashboard step fixture",
            created_by=self.admin,
        )
        AppointmentRescheduleStep.objects.create(
            plan=plan,
            position=1,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.FAILED,
            source_appointment=source,
            proposed_starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(12, 0)),
            proposed_ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(12, 30)),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
        )

        response = self.client.get(reverse("dashboard"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context["reschedule_step_count"], 1)
        self.assertEqual(response.context["failed_step_count"], 1)
        self.assertEqual(response.context["priority_total"], 1)
        self.assertContains(response, "Разобрать шаги переноса")
        self.assertContains(response, "шагов переноса в очереди")
        self.assertContains(response, f'{reverse("work_queue")}#queue-reschedule-steps')

    def test_work_queue_surfaces_ready_reschedule_chains(self):
        plan, chain = self._reschedule_chain_fixture()
        plan_svc.revalidate_chain(chain)

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertIn(chain, list(response.context["reschedule_chains"]))
        self.assertContains(response, "#queue-reschedule-chains")
        self.assertContains(response, 'id="queue-reschedule-chains"')
        self.assertContains(response, "Открыть цепочку")
        self.assertContains(
            response,
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#chain-{chain.pk}',
        )

    def test_work_queue_surfaces_priority_reschedule_steps(self):
        source = self._appointment()
        plan = AppointmentReschedulePlan.objects.create(
            status=AppointmentReschedulePlan.Status.READY,
            plan_type=AppointmentReschedulePlan.PlanType.MANUAL,
            reason="Step queue fixture",
            created_by=self.admin,
        )
        ready_step = AppointmentRescheduleStep.objects.create(
            plan=plan,
            position=1,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.VALID,
            source_appointment=source,
            proposed_starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(12, 0)),
            proposed_ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(12, 30)),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
        )
        stale_step = AppointmentRescheduleStep.objects.create(
            plan=plan,
            position=2,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.STALE,
            source_appointment=source,
            proposed_starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(13, 0)),
            proposed_ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(13, 30)),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
            validation_messages=["slot changed"],
        )
        failed_step = AppointmentRescheduleStep.objects.create(
            plan=plan,
            position=3,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.FAILED,
            source_appointment=source,
            proposed_starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(14, 0)),
            proposed_ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(14, 30)),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
        )

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        steps = list(response.context["reschedule_steps"])
        self.assertEqual(steps[:3], [failed_step, stale_step, ready_step])
        self.assertContains(response, "#queue-reschedule-steps")
        self.assertContains(response, 'id="queue-reschedule-steps"')
        self.assertContains(response, "Открыть шаг")
        self.assertContains(
            response,
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#step-{failed_step.pk}',
        )
        step_summary = next(
            item
            for item in response.context["queue_summary_items"]
            if item["href"] == "#queue-reschedule-steps"
        )
        self.assertEqual(step_summary["value"], 3)
        self.assertEqual(step_summary["tone"], "danger")

    def test_work_queue_excludes_chain_steps_from_step_attention(self):
        plan, chain = self._reschedule_chain_fixture()
        plan_svc.revalidate_chain(chain)

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertIn(chain, list(response.context["reschedule_chains"]))
        self.assertEqual(list(response.context["reschedule_steps"]), [])
        self.assertContains(
            response,
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#chain-{chain.pk}',
        )

    def test_dashboard_orders_chain_attention_by_operational_priority(self):
        self._chain_with_status(AppointmentRescheduleChain.Status.READY, "Ready chain", days=5)
        self._chain_with_status(AppointmentRescheduleChain.Status.STALE, "Stale chain", days=8)
        self._chain_with_status(AppointmentRescheduleChain.Status.FAILED, "Failed chain", days=11)

        response = self.client.get(reverse("dashboard"))

        self.assertEqual(response.status_code, 200)
        chain_items = [
            item
            for item in response.context["dashboard_focus_items"]
            if "chain_" in item["href"]
        ]
        self.assertEqual(
            [item["href"].split("focus=")[1] for item in chain_items],
            ["chain_failed", "chain_stale", "chain_ready"],
        )
        self.assertEqual(response.context["failed_chain_count"], 1)
        self.assertEqual(response.context["stale_chain_count"], 1)
        self.assertEqual(response.context["ready_chain_count"], 1)
        self.assertEqual(response.context["chain_attention_count"], 3)

    def test_work_queue_orders_chain_attention_by_operational_priority(self):
        self._chain_with_status(AppointmentRescheduleChain.Status.READY, "Ready chain", days=5)
        self._chain_with_status(AppointmentRescheduleChain.Status.STALE, "Stale chain", days=8)
        self._chain_with_status(AppointmentRescheduleChain.Status.FAILED, "Failed chain", days=11)

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        chains = list(response.context["reschedule_chains"])
        self.assertEqual(
            [chain.status for chain in chains[:3]],
            [
                AppointmentRescheduleChain.Status.FAILED,
                AppointmentRescheduleChain.Status.STALE,
                AppointmentRescheduleChain.Status.READY,
            ],
        )
        chain_summary = next(
            item
            for item in response.context["queue_summary_items"]
            if item["href"] == "#queue-reschedule-chains"
        )
        self.assertEqual(chain_summary["value"], 3)
        self.assertEqual(chain_summary["tone"], "danger")

    def test_work_queue_labels_chain_next_actions(self):
        self._chain_with_status(AppointmentRescheduleChain.Status.READY, "Ready chain", days=5)
        self._chain_with_status(AppointmentRescheduleChain.Status.STALE, "Stale chain", days=8)
        self._chain_with_status(AppointmentRescheduleChain.Status.FAILED, "Failed chain", days=11)

        response = self.client.get(reverse("work_queue"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(
            response,
            "Следующее действие: открыть план, разобрать ошибку цепочки и повторить после исправления.",
        )
        self.assertContains(
            response,
            "Следующее действие: перепроверить цепочку в плане перед применением.",
        )
        self.assertContains(
            response,
            "Следующее действие: открыть план и применить цепочку после финальной проверки.",
        )

    def test_reschedule_plan_detail_renders_steps(self):
        appointment = self._appointment()
        self.client.post(reverse("appointment_reschedule_plan_create", args=[appointment.pk]))
        plan = AppointmentReschedulePlan.objects.get(root_appointment=appointment)

        response = self.client.get(reverse("appointment_reschedule_plan_detail", args=[plan.pk]))
        step = plan.steps.order_by("position").first()

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "План переноса")
        self.assertContains(response, "reschedule-plan-table")
        self.assertContains(response, f'id="step-{step.pk}"')
        self.assertContains(response, 'data-label="Команда"')
        self.assertContains(
            response,
            reverse("appointment_move", args=[appointment.pk]),
        )
        self.assertContains(
            response,
            "План не меняет расписание, пока администратор не применит шаг.",
        )
        self.assertContains(
            response,
            "Перед применением шага система повторно проверяет получателей",
        )
        self.assertContains(response, "Перепроверить план")

    def test_reschedule_plan_detail_hides_revalidate_for_terminal_plan(self):
        appointment = self._appointment()
        plan = plan_svc.create_plan_for_appointment(
            appointment, actor=self.admin, days=2, limit=1
        )
        plan.status = AppointmentReschedulePlan.Status.APPLIED
        plan.save(update_fields=["status", "updated_at"])

        response = self.client.get(reverse("appointment_reschedule_plan_detail", args=[plan.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertFalse(response.context["can_revalidate_plan"])
        self.assertContains(response, reverse("appointment_detail", args=[appointment.pk]))
        self.assertNotContains(response, reverse("appointment_move", args=[appointment.pk]))
        self.assertContains(
            response,
            "План применен: изменения расписания уже записаны, действия закрыты.",
        )
        self.assertContains(
            response,
            "План открыт только для просмотра истории.",
        )
        self.assertContains(response, "Повторная проверка плана недоступна")
        self.assertNotContains(
            response,
            "План не меняет расписание, пока администратор не применит шаг.",
        )
        self.assertNotContains(
            response,
            "Перед применением шага система повторно проверяет получателей",
        )
        self.assertNotContains(response, 'name="action" value="revalidate"')
        self.assertNotContains(response, "Перепроверить план")

    def test_reschedule_plan_detail_rejects_terminal_plan_revalidate_post(self):
        appointment = self._appointment()
        plan = plan_svc.create_plan_for_appointment(
            appointment, actor=self.admin, days=2, limit=1
        )
        step = plan.steps.get()
        step.status = AppointmentRescheduleStep.Status.STALE
        step.validation_messages = ["Сохраненная причина"]
        step.save(update_fields=["status", "validation_messages", "updated_at"])
        plan.status = AppointmentReschedulePlan.Status.CANCELLED
        plan.validation_summary = {"locked": "cancelled"}
        plan.save(update_fields=["status", "validation_summary", "updated_at"])

        response = self.client.post(
            reverse("appointment_reschedule_plan_detail", args=[plan.pk]),
            {"action": "revalidate"},
            follow=True,
        )

        plan.refresh_from_db()
        step.refresh_from_db()
        self.assertEqual(response.status_code, 200)
        self.assertEqual(plan.status, AppointmentReschedulePlan.Status.CANCELLED)
        self.assertEqual(plan.validation_summary, {"locked": "cancelled"})
        self.assertEqual(step.status, AppointmentRescheduleStep.Status.STALE)
        self.assertEqual(step.validation_messages, ["Сохраненная причина"])
        self.assertContains(response, "нельзя перепроверять")
        self.assertContains(response, "Повторная проверка плана недоступна")
        self.assertNotContains(response, 'name="action" value="revalidate"')

    def test_reschedule_plan_detail_hides_step_and_chain_actions_for_terminal_plan(self):
        plan, chain = self._reschedule_chain_fixture()
        chain.status = AppointmentRescheduleChain.Status.READY
        chain.validation_summary = {"ready": 2}
        chain.save(update_fields=["status", "validation_summary", "updated_at"])
        plan.status = AppointmentReschedulePlan.Status.CANCELLED
        plan.save(update_fields=["status", "updated_at"])
        plan.steps.update(
            confirmation_status=AppointmentRescheduleStep.ConfirmationStatus.WAITING
        )

        response = self.client.get(reverse("appointment_reschedule_plan_detail", args=[plan.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertFalse(response.context["can_mutate_plan"])
        self.assertFalse(response.context["chains"][0].can_revalidate)
        self.assertFalse(response.context["chains"][0].can_apply)
        self.assertContains(
            response,
            "План отменен: действия закрыты, история сохранена.",
        )
        self.assertContains(
            response,
            "План открыт только для просмотра истории.",
        )
        self.assertContains(response, "Цепочка показывает сохраненный порядок шагов.")
        self.assertContains(response, "Цепочка доступна только как история")
        self.assertNotContains(response, "Перед применением система повторно проверяет")
        self.assertContains(response, "План завершен или отменен.")
        self.assertContains(response, "Архив согласований", count=2)
        self.assertContains(response, "Последний статус:", count=2)
        self.assertContains(response, "Повторная проверка плана недоступна")
        self.assertNotContains(response, 'name="action" value="revalidate"')
        self.assertNotContains(response, 'name="action" value="revalidate_chain"')
        self.assertNotContains(response, 'name="action" value="apply_chain"')
        self.assertNotContains(response, 'name="action" value="send_step_confirmations"')
        self.assertNotContains(response, 'name="action" value="apply_step"')
        self.assertNotContains(response, 'name="action" value="mark_review_conflict_resolved"')

    def test_reschedule_plan_detail_rejects_terminal_plan_step_post(self):
        appointment = self._appointment()
        plan = plan_svc.create_plan_for_appointment(
            appointment, actor=self.admin, days=2, limit=1
        )
        step = plan.steps.get()
        plan.status = AppointmentReschedulePlan.Status.CANCELLED
        plan.validation_summary = {"locked": "cancelled"}
        plan.save(update_fields=["status", "validation_summary", "updated_at"])

        response = self.client.post(
            reverse("appointment_reschedule_plan_detail", args=[plan.pk]),
            {"action": "apply_step", "step_id": step.pk},
            follow=True,
        )

        plan.refresh_from_db()
        step.refresh_from_db()
        appointment.refresh_from_db()
        self.assertEqual(response.status_code, 200)
        self.assertEqual(plan.status, AppointmentReschedulePlan.Status.CANCELLED)
        self.assertEqual(plan.validation_summary, {"locked": "cancelled"})
        self.assertEqual(step.status, AppointmentRescheduleStep.Status.VALID)
        self.assertIsNone(step.created_appointment_id)
        self.assertEqual(appointment.status, Appointment.Status.CONFIRMED)
        self.assertContains(response, "нельзя изменять")
        self.assertNotContains(response, 'name="action" value="apply_step"')

    def test_reschedule_plan_detail_renders_chain_read_only_block(self):
        first_source = self._appointment()
        second_start = _local_dt(timezone.localdate() + timedelta(days=5), time(11, 0))
        second_source = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=second_start,
            ends_at=second_start + timedelta(minutes=30),
            billing_account=self.account,
        )
        plan = AppointmentReschedulePlan.objects.create(
            status=AppointmentReschedulePlan.Status.READY,
            plan_type=AppointmentReschedulePlan.PlanType.MANUAL,
            reason="Буферная цепочка",
            created_by=self.admin,
        )
        first = AppointmentRescheduleStep.objects.create(
            plan=plan,
            position=1,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.VALID,
            source_appointment=second_source,
            proposed_starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(12, 0)),
            proposed_ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(12, 30)),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
        )
        second = AppointmentRescheduleStep.objects.create(
            plan=plan,
            position=2,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.VALID,
            source_appointment=first_source,
            proposed_starts_at=second_start,
            proposed_ends_at=second_start + timedelta(minutes=30),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
        )
        chain_result = plan_svc.create_chain_for_steps(
            plan,
            step_ids=[first.pk, second.pk],
            dependencies=[
                {
                    "predecessor_step_id": first.pk,
                    "successor_step_id": second.pk,
                    "reason": "Освобождает окно для второго шага.",
                }
            ],
            title="Буферная цепочка",
            actor=self.admin,
        )
        chain = chain_result.chain

        response = self.client.get(reverse("appointment_reschedule_plan_detail", args=[plan.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertIn("chains", response.context)
        self.assertContains(response, "Буферная цепочка")
        self.assertContains(response, "reschedule-chain-table")
        self.assertContains(response, "reschedule-chain-dependency-table")
        self.assertContains(response, f'id="chain-{chain.pk}"')
        self.assertContains(response, 'data-label="Порядок"')
        self.assertContains(response, 'data-label="Предшественник"')
        self.assertContains(response, f'href="#step-{first.pk}"')
        self.assertContains(response, f'href="#step-{second.pk}"')
        self.assertContains(response, f'<a class="small-action" href="#step-{first.pk}">', count=2)
        self.assertContains(response, f'<a class="small-action" href="#step-{second.pk}">', count=2)
        self.assertContains(response, "Атомарно все или ничего")
        self.assertNotContains(response, "Применить цепочку")

    def test_reschedule_plan_detail_explains_stale_chain_step_checks(self):
        plan, chain = self._reschedule_chain_fixture()
        first = chain.steps.order_by("chain_position").first()
        blocker_parent = ParentGuardian.objects.create(
            last_name="Blocker",
            first_name="Parent",
            phone="+7 900 000-77-03",
        )
        blocker_child = Child.objects.create(
            last_name="Blocker",
            first_name="Child",
            primary_parent=blocker_parent,
        )
        Appointment.objects.create(
            child=blocker_child,
            staff_member=first.proposed_primary_staff,
            service=self.service2,
            starts_at=first.proposed_starts_at,
            ends_at=first.proposed_ends_at,
            room=first.proposed_room,
        )
        plan_svc.revalidate_chain(chain)
        first.refresh_from_db()

        response = self.client.get(reverse("appointment_reschedule_plan_detail", args=[plan.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Проверка расписания устарела")
        self.assertContains(response, "Шаг 1")
        self.assertContains(response, first.validation_messages[0])

    def test_reschedule_plan_detail_explains_confirmation_blocked_chain(self):
        plan, chain = self._reschedule_chain_fixture()
        first = chain.steps.order_by("chain_position").first()
        AppointmentConfirmation.objects.create(
            appointment=first.source_appointment,
            reschedule_step=first,
            target_type=AppointmentConfirmation.TargetType.REPRESENTATIVE,
            email="declined-detail@example.local",
            subject="Подтверждение переноса",
            message="Подтвердите перенос",
            status=AppointmentConfirmation.Status.DECLINED,
        )
        plan_svc.revalidate_chain(chain)

        response = self.client.get(reverse("appointment_reschedule_plan_detail", args=[plan.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Согласование блокирует шаг")
        self.assertContains(response, "статус: есть отказ")

    def test_reschedule_plan_detail_explains_failed_chain_apply_error(self):
        plan, chain = self._reschedule_chain_fixture()
        chain.status = AppointmentRescheduleChain.Status.FAILED
        chain.validation_summary = {"apply_error": ["second step failed"]}
        chain.save(update_fields=["status", "validation_summary", "updated_at"])

        response = self.client.get(reverse("appointment_reschedule_plan_detail", args=[plan.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Ошибка применения цепочки")
        self.assertContains(response, "second step failed")

    def test_reschedule_plan_detail_can_revalidate_chain(self):
        plan, chain = self._reschedule_chain_fixture()

        response = self.client.post(
            reverse("appointment_reschedule_plan_detail", args=[plan.pk]),
            {"action": "revalidate_chain", "chain_id": str(chain.pk)},
            follow=True,
        )

        chain.refresh_from_db()
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.redirect_chain[-1][0],
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#chain-{chain.pk}',
        )
        self.assertEqual(chain.status, chain.Status.READY)
        self.assertEqual(chain.validation_summary["ready"], 2)
        self.assertContains(response, "Цепочка перепроверена")
        self.assertContains(response, "Цепочка показывает зависимый порядок шагов.")
        self.assertContains(response, "Перед применением система повторно проверяет")
        self.assertContains(response, 'name="action" value="revalidate_chain"')
        self.assertContains(response, 'name="action" value="apply_chain"')

    def test_reschedule_plan_detail_hides_revalidate_for_terminal_chains(self):
        cases = [
            (AppointmentRescheduleChain.Status.APPLYING, 5),
            (AppointmentRescheduleChain.Status.APPLIED, 6),
            (AppointmentRescheduleChain.Status.CANCELLED, 7),
        ]
        for status, days in cases:
            with self.subTest(status=status):
                plan, chain = self._reschedule_chain_fixture(days=days)
                chain.status = status
                chain.save(update_fields=["status", "updated_at"])

                response = self.client.get(
                    reverse("appointment_reschedule_plan_detail", args=[plan.pk])
                )

                self.assertEqual(response.status_code, 200)
                self.assertFalse(response.context["chains"][0].can_revalidate)
                self.assertContains(
                    response,
                    "Повторная проверка недоступна для текущего статуса цепочки.",
                )
                self.assertNotContains(response, 'name="action" value="revalidate_chain"')

    def test_reschedule_plan_detail_can_apply_ready_chain(self):
        plan, chain = self._reschedule_chain_fixture()
        plan_svc.revalidate_chain(chain)

        response = self.client.post(
            reverse("appointment_reschedule_plan_detail", args=[plan.pk]),
            {"action": "apply_chain", "chain_id": str(chain.pk)},
            follow=True,
        )

        chain.refresh_from_db()
        plan.refresh_from_db()
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.redirect_chain[-1][0],
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#chain-{chain.pk}',
        )
        self.assertEqual(chain.status, chain.Status.APPLIED)
        self.assertEqual(plan.status, AppointmentReschedulePlan.Status.APPLIED)
        self.assertContains(response, "Цепочка применена")
        self.assertContains(response, "Цепочка доступна только как история")
        self.assertContains(response, "Цепочка показывает сохраненный порядок шагов.")
        self.assertNotContains(response, "Перед применением система повторно проверяет")
        self.assertNotContains(response, 'name="action" value="revalidate_chain"')

    def test_staff_absence_plan_detail_renders_manual_review_actions(self):
        appointment = self._appointment()
        day = timezone.localtime(appointment.starts_at).date()
        plan = plan_svc.create_staff_absence_plan(
            self.staff,
            date_from=day,
            date_to=day,
            reason="Больничный",
            actor=self.admin,
        )

        response = self.client.get(reverse("appointment_reschedule_plan_detail", args=[plan.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Открыть занятие")
        self.assertContains(response, "Перенести вручную")
        self.assertContains(response, reverse("appointment_cancel", args=[appointment.pk]))
        self.assertContains(response, "Отметить разобранным")

    def test_reschedule_plan_detail_marks_review_conflict_after_manual_resolution(self):
        appointment = self._appointment()
        day = timezone.localtime(appointment.starts_at).date()
        plan = plan_svc.create_staff_absence_plan(
            self.staff,
            date_from=day,
            date_to=day,
            reason="Больничный",
            actor=self.admin,
        )
        step = plan.steps.get()

        blocked = self.client.post(
            reverse("appointment_reschedule_plan_detail", args=[plan.pk]),
            {"action": "mark_review_conflict_resolved", "step_id": step.pk},
            follow=True,
        )

        self.assertEqual(
            blocked.redirect_chain[-1][0],
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#step-{step.pk}',
        )
        self.assertContains(blocked, "Сначала перенесите или отмените занятие")
        step.refresh_from_db()
        self.assertEqual(step.status, AppointmentRescheduleStep.Status.PENDING)

        appointment.status = Appointment.Status.CANCELLED
        appointment.save(update_fields=["status", "updated_at"])

        response = self.client.post(
            reverse("appointment_reschedule_plan_detail", args=[plan.pk]),
            {"action": "mark_review_conflict_resolved", "step_id": step.pk},
            follow=True,
        )

        self.assertEqual(
            response.redirect_chain[-1][0],
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#step-{step.pk}',
        )
        self.assertContains(response, "Ручной конфликт отмечен разобранным.")
        step.refresh_from_db()
        plan.refresh_from_db()
        self.assertEqual(step.status, AppointmentRescheduleStep.Status.SKIPPED)
        self.assertEqual(plan.status, AppointmentReschedulePlan.Status.APPLIED)

    def test_reschedule_plan_detail_can_apply_valid_step(self):
        appointment = self._appointment()
        self.client.post(reverse("appointment_reschedule_plan_create", args=[appointment.pk]))
        plan = AppointmentReschedulePlan.objects.get(root_appointment=appointment)
        step = (
            plan.steps.filter(status=AppointmentRescheduleStep.Status.VALID)
            .order_by("position")
            .first()
        )

        response = self.client.post(
            reverse("appointment_reschedule_plan_detail", args=[plan.pk]),
            {"action": "apply_step", "step_id": step.pk},
        )

        self.assertEqual(
            response.url,
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#step-{step.pk}',
        )
        appointment.refresh_from_db()
        step.refresh_from_db()
        self.assertEqual(appointment.status, Appointment.Status.RESCHEDULED)
        self.assertEqual(step.status, AppointmentRescheduleStep.Status.APPLIED)
        self.assertIsNotNone(step.created_appointment)

    def test_reschedule_plan_detail_skips_alternative_after_apply(self):
        appointment = self._appointment()
        plan = AppointmentReschedulePlan.objects.create(
            status=AppointmentReschedulePlan.Status.READY,
            plan_type=AppointmentReschedulePlan.PlanType.SINGLE_MOVE,
            root_appointment=appointment,
            reason="Проверка альтернатив",
            created_by=self.admin,
        )
        step = AppointmentRescheduleStep.objects.create(
            plan=plan,
            position=1,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.VALID,
            source_appointment=appointment,
            proposed_starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(11, 0)),
            proposed_ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(11, 30)),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
        )
        alternative = AppointmentRescheduleStep.objects.create(
            plan=plan,
            position=2,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.VALID,
            source_appointment=appointment,
            proposed_starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(12, 0)),
            proposed_ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(12, 30)),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
        )

        response = self.client.post(
            reverse("appointment_reschedule_plan_detail", args=[plan.pk]),
            {"action": "apply_step", "step_id": step.pk},
            follow=True,
        )

        self.assertEqual(
            response.redirect_chain[-1][0],
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#step-{step.pk}',
        )
        self.assertContains(response, "Пропущен")
        step.refresh_from_db()
        alternative.refresh_from_db()
        plan.refresh_from_db()
        self.assertEqual(step.status, AppointmentRescheduleStep.Status.APPLIED)
        self.assertEqual(alternative.status, AppointmentRescheduleStep.Status.SKIPPED)
        self.assertEqual(plan.status, AppointmentReschedulePlan.Status.APPLIED)

    def test_reschedule_plan_detail_can_apply_step_with_room_override(self):
        appointment = self._appointment()
        proposed_start = _local_dt(timezone.localdate() + timedelta(days=5), time(11, 0))
        blocker_child = Child.objects.create(last_name="Room", first_name="Blocker")
        blocker_staff = StaffMember.objects.create(full_name="Room Blocker Specialist")
        Appointment.objects.create(
            child=blocker_child,
            service=self.service,
            staff_member=blocker_staff,
            room=self.room,
            starts_at=proposed_start,
            ends_at=proposed_start + timedelta(minutes=30),
        )
        plan = AppointmentReschedulePlan.objects.create(
            status=AppointmentReschedulePlan.Status.READY,
            plan_type=AppointmentReschedulePlan.PlanType.SINGLE_MOVE,
            root_appointment=appointment,
            reason="Проверка override кабинета",
            created_by=self.admin,
        )
        step = AppointmentRescheduleStep.objects.create(
            plan=plan,
            position=1,
            action_type=AppointmentRescheduleStep.ActionType.MOVE,
            status=AppointmentRescheduleStep.Status.VALID,
            source_appointment=appointment,
            proposed_starts_at=proposed_start,
            proposed_ends_at=proposed_start + timedelta(minutes=30),
            proposed_room=self.room,
            proposed_primary_staff=self.staff,
            requires_room_override=True,
            validation_messages=["кабинет превышает правила вместимости"],
        )

        detail = self.client.get(reverse("appointment_reschedule_plan_detail", args=[plan.pk]))

        self.assertContains(detail, "Нужно одноразовое разрешение кабинета")
        self.assertContains(detail, "Применить с разрешением кабинета")

        blocked = self.client.post(
            reverse("appointment_reschedule_plan_detail", args=[plan.pk]),
            {"action": "apply_step", "step_id": step.pk},
            follow=True,
        )

        self.assertEqual(
            blocked.redirect_chain[-1][0],
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#step-{step.pk}',
        )
        self.assertContains(blocked, "override кабинета")
        appointment.refresh_from_db()
        self.assertEqual(appointment.status, Appointment.Status.CONFIRMED)

        response = self.client.post(
            reverse("appointment_reschedule_plan_detail", args=[plan.pk]),
            {"action": "apply_step", "step_id": step.pk, "allow_room_override": "1"},
            follow=True,
        )

        self.assertEqual(
            response.redirect_chain[-1][0],
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#step-{step.pk}',
        )
        self.assertContains(response, "Шаг переноса применен с одноразовым разрешением кабинета.")
        step.refresh_from_db()
        self.assertEqual(step.status, AppointmentRescheduleStep.Status.APPLIED)
        override = AppointmentRoomOverride.objects.get(
            appointment=step.created_appointment
        )
        self.assertEqual(override.created_by, self.admin)

    def test_reschedule_plan_detail_can_send_step_confirmations(self):
        self.staff.email = "qa-plan-staff@example.local"
        self.staff.save(update_fields=["email", "updated_at"])
        appointment = self._appointment()
        self.client.post(reverse("appointment_reschedule_plan_create", args=[appointment.pk]))
        plan = AppointmentReschedulePlan.objects.get(root_appointment=appointment)
        step = (
            plan.steps.filter(status=AppointmentRescheduleStep.Status.VALID)
            .order_by("position")
            .first()
        )

        response = self.client.post(
            reverse("appointment_reschedule_plan_detail", args=[plan.pk]),
            {"action": "send_step_confirmations", "step_id": step.pk},
        )

        self.assertEqual(
            response.url,
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#step-{step.pk}',
        )
        confirmations = AppointmentConfirmation.objects.filter(reschedule_step=step)
        self.assertEqual(confirmations.count(), 2)
        self.assertTrue(
            confirmations.filter(
                target_type=AppointmentConfirmation.TargetType.REPRESENTATIVE,
                participant__child=self.child,
            ).exists()
        )
        self.assertTrue(
            confirmations.filter(
                target_type=AppointmentConfirmation.TargetType.SPECIALIST,
                email="qa-plan-staff@example.local",
            ).exists()
        )
        step.refresh_from_db()
        self.assertEqual(
            step.confirmation_status,
            AppointmentRescheduleStep.ConfirmationStatus.WAITING,
        )

        detail = self.client.get(reverse("appointment_reschedule_plan_detail", args=[plan.pk]))
        self.assertContains(detail, "Согласования")
        self.assertContains(detail, "Ожидает ответов")
        self.assertContains(detail, "Ждем ответы")

    def test_reschedule_plan_detail_can_apply_approved_step(self):
        self.staff.email = "qa-plan-approved-staff@example.local"
        self.staff.save(update_fields=["email", "updated_at"])
        appointment = self._appointment()
        plan = plan_svc.create_plan_for_appointment(
            appointment, actor=self.admin, days=2, limit=1
        )
        step = plan.steps.get()
        plan_svc.create_confirmations_for_step(step, actor=self.admin)
        AppointmentConfirmation.objects.filter(reschedule_step=step).update(
            status=AppointmentConfirmation.Status.CONFIRMED,
            responded_at=timezone.now(),
        )
        step = plan_svc.refresh_step_confirmation_status(step)

        detail = self.client.get(reverse("appointment_reschedule_plan_detail", args=[plan.pk]))

        self.assertContains(detail, "Согласовано")
        self.assertContains(detail, "Применить согласованный перенос")
        self.assertNotContains(detail, "Ждем ответы")

        response = self.client.post(
            reverse("appointment_reschedule_plan_detail", args=[plan.pk]),
            {"action": "apply_step", "step_id": step.pk},
            follow=True,
        )

        self.assertEqual(
            response.redirect_chain[-1][0],
            f'{reverse("appointment_reschedule_plan_detail", args=[plan.pk])}#step-{step.pk}',
        )
        self.assertContains(response, "Согласованный шаг переноса применен.")
        appointment.refresh_from_db()
        step.refresh_from_db()
        self.assertEqual(appointment.status, Appointment.Status.RESCHEDULED)
        self.assertEqual(step.status, AppointmentRescheduleStep.Status.APPLIED)
        self.assertEqual(
            step.confirmation_status,
            AppointmentRescheduleStep.ConfirmationStatus.APPROVED,
        )

    def test_reschedule_step_public_confirmation_does_not_confirm_source_appointment(self):
        appointment = self._appointment()
        appointment.status = Appointment.Status.PROPOSED
        appointment.save(update_fields=["status", "updated_at"])
        plan = plan_svc.create_plan_for_appointment(
            appointment, actor=self.admin, days=2, limit=1
        )
        step = plan.steps.get()
        confirmation = AppointmentConfirmation.objects.create(
            appointment=appointment,
            reschedule_step=step,
            target_type=AppointmentConfirmation.TargetType.REPRESENTATIVE,
            representative=self.parent,
            participant=appointment.primary_participant,
            email=self.parent.email,
            subject="Согласование переноса",
            message="Подтвердите перенос.",
            sent_by=self.admin,
        )
        self.client.logout()

        response = self.client.post(
            reverse("appointment_confirmation_public", args=[confirmation.token]),
            {"action": "confirm", "response_note": "Подходит"},
        )

        self.assertEqual(response.status_code, 200)
        appointment.refresh_from_db()
        confirmation.refresh_from_db()
        step.refresh_from_db()
        self.assertEqual(confirmation.status, AppointmentConfirmation.Status.CONFIRMED)
        self.assertEqual(appointment.status, Appointment.Status.PROPOSED)
        self.assertEqual(
            step.confirmation_status,
            AppointmentRescheduleStep.ConfirmationStatus.APPROVED,
        )


class AppointmentBillingTests(NewViewsTestBase):
    def test_appointment_billing_get_redirects_to_detail(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() - timedelta(days=1), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() - timedelta(days=1), time(10, 30)),
            billing_account=self.account,
        )
        response = self.client.get(reverse("appointment_billing", args=[appt.pk]))
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, reverse("appointment_detail", args=[appt.pk]))

    def test_appointment_billing_invalid_returns_400(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() - timedelta(days=1), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() - timedelta(days=1), time(10, 30)),
            billing_account=self.account,
        )
        response = self.client.post(
            reverse("appointment_billing", args=[appt.pk]),
            {
                "billing_decision": "",
                "billing_account": "",
                "amount": "",
                "reason": "",
            },
        )
        self.assertEqual(response.status_code, 400)

    def test_appointment_billing_charge(self):
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() - timedelta(days=1), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() - timedelta(days=1), time(10, 30)),
            billing_account=self.account,
        )
        response = self.client.post(
            reverse("appointment_billing", args=[appt.pk]),
            {
                "billing_decision": Appointment.BillingDecision.CHARGE,
                "billing_account": self.account.id,
                "amount": "-1",
                "reason": "OK",
                "next": "",
            },
        )
        self.assertEqual(response.status_code, 302)
        appt.refresh_from_db()
        self.assertEqual(appt.billing_decision, Appointment.BillingDecision.CHARGE)

    def test_appointment_billing_without_participant_id_uses_single_participant(self):
        participant_child = Child.objects.create(last_name="Списание", first_name="Участник")
        participant_account = BalanceAccount.objects.create(
            child=participant_child,
            funding_source=self.funding,
            unit=BalanceAccount.Unit.SESSIONS,
            service_scope=BalanceAccount.ServiceScope.ANY,
            initial_amount=Decimal("5"),
        )
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() - timedelta(days=1), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() - timedelta(days=1), time(10, 30)),
            billing_account=None,
        )
        appt.participants.filter(child=self.child).delete()
        participant = AppointmentParticipant.objects.create(
            appointment=appt,
            child=participant_child,
            billing_account=participant_account,
            starts_at_snapshot=appt.starts_at,
            ends_at_snapshot=appt.ends_at,
            appointment_status=appt.status,
        )

        response = self.client.post(
            reverse("appointment_billing", args=[appt.pk]),
            {
                "billing_decision": Appointment.BillingDecision.CHARGE,
                "billing_account": participant_account.id,
                "amount": "-1",
                "reason": "OK",
                "next": "",
            },
        )

        self.assertEqual(response.status_code, 302)
        appt.refresh_from_db()
        participant.refresh_from_db()
        self.assertEqual(appt.billing_decision, Appointment.BillingDecision.UNDECIDED)
        self.assertIsNone(appt.billing_account)
        self.assertEqual(participant.billing_decision, Appointment.BillingDecision.CHARGE)
        self.assertEqual(participant.billing_account, participant_account)
        self.assertTrue(
            LedgerEntry.objects.filter(
                appointment=appt,
                appointment_participant=participant,
                account=participant_account,
            ).exists()
        )

    def test_appointment_billing_group_requires_participant_id(self):
        second_child = Child.objects.create(last_name="Группа", first_name="Без выбора")
        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() - timedelta(days=1), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() - timedelta(days=1), time(10, 30)),
            billing_account=self.account,
            session_type=Appointment.SessionType.GROUP,
        )
        AppointmentParticipant.objects.create(
            appointment=appt,
            child=second_child,
            starts_at_snapshot=appt.starts_at,
            ends_at_snapshot=appt.ends_at,
            appointment_status=appt.status,
        )

        response = self.client.post(
            reverse("appointment_billing", args=[appt.pk]),
            {
                "billing_decision": Appointment.BillingDecision.CHARGE,
                "billing_account": self.account.id,
                "amount": "-1",
                "reason": "OK",
                "next": "",
            },
        )

        self.assertEqual(response.status_code, 400)
        self.assertFalse(LedgerEntry.objects.filter(appointment=appt).exists())


class RecipientEditTests(NewViewsTestBase):
    def test_recipient_import_preview_get(self):
        response = self.client.get(reverse("recipient_import_preview"))
        self.assertEqual(response.status_code, 200)
        self.assertIn("form", response.context)

    def test_recipient_import_preview_post_csv_does_not_create_records(self):
        from django.core.files.uploadedfile import SimpleUploadedFile

        before_count = Child.objects.count()
        upload = SimpleUploadedFile(
            "recipients.csv",
            (
                "Фамилия получателя;Имя получателя;Телефон представителя\n"
                "Сидоров;Илья;+7 900\n"
            ).encode(),
            content_type="text/csv",
        )

        response = self.client.post(reverse("recipient_import_preview"), {"file": upload})

        self.assertEqual(response.status_code, 200)
        self.assertEqual(Child.objects.count(), before_count)
        self.assertEqual(response.context["preview"].valid_count, 1)
        self.assertContains(response, "Сидоров")

    def test_recipient_import_preview_shows_operator_summary(self):
        from django.core.files.uploadedfile import SimpleUploadedFile

        upload = SimpleUploadedFile(
            "recipients.csv",
            (
                "Фамилия получателя;Имя получателя;Получает расписание\n"
                "Сидоров;Илья;да\n"
            ).encode(),
            content_type="text/csv",
        )

        response = self.client.post(reverse("recipient_import_preview"), {"file": upload})

        self.assertEqual(response.status_code, 200)
        self.assertIn("import_summary_items", response.context)
        self.assertIn("import_next_action", response.context)
        self.assertEqual(
            response.context["import_next_action"]["title"],
            "Файл готов к будущему импорту",
        )
        self.assertContains(response, "Следующее действие")
        self.assertContains(response, "Файл готов к будущему импорту")
        self.assertContains(response, "Запись в базу отключена")
        self.assertContains(response, 'id="import-upload"')
        self.assertContains(response, 'id="import-columns"')
        self.assertContains(response, 'id="import-rows"')
        self.assertContains(response, "import-columns-table")
        self.assertContains(response, "import-rows-table")
        self.assertContains(response, 'data-label="Поле системы"')
        self.assertContains(response, 'data-label="Получатель"')
        self.assertContains(response, 'data-label="Проверка"')

    def test_recipient_detail_shows_representative_links(self):
        response = self.client.get(reverse("recipient_detail", args=[self.child.pk]))
        self.assertEqual(response.status_code, 200)
        self.assertIn("representative_links", response.context)
        self.assertIn("recipient_detail_summary_items", response.context)
        self.assertIn("recipient_detail_next_action", response.context)
        self.assertContains(
            response,
            reverse("recipient_representative_create", args=[self.child.pk]),
        )
        self.assertContains(response, "Следующий шаг")
        self.assertContains(response, 'id="recipient-balances"')
        self.assertContains(response, "recipient-representatives-table")
        self.assertContains(response, "recipient-balance-table")
        self.assertContains(response, 'data-label="Расписание"')
        self.assertContains(response, 'data-label="Действия"')

    def test_recipient_detail_tables_have_mobile_labels(self):
        program = TreatmentProgram.objects.create(child=self.child, title="Карточная программа")
        ProgramBlock.objects.create(
            program=program,
            title="Первый блок",
            service=self.service,
            staff_member=self.staff,
            planned_sessions=4,
            balance_account=self.account,
        )
        future_start = _local_dt(timezone.localdate() + timedelta(days=3), time(11, 0))
        Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=future_start,
            ends_at=future_start + timedelta(minutes=30),
        )
        past_start = _local_dt(timezone.localdate() - timedelta(days=3), time(10, 0))
        Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=past_start,
            ends_at=past_start + timedelta(minutes=30),
            status=Appointment.Status.COMPLETED,
        )

        response = self.client.get(reverse("recipient_detail", args=[self.child.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "recipient-program-blocks-table")
        self.assertContains(response, "recipient-upcoming-table")
        self.assertContains(response, "recipient-recent-table")
        self.assertContains(response, 'data-label="Блок"')
        self.assertContains(response, 'data-label="План"')
        self.assertContains(response, 'data-label="Получатели"')
        self.assertContains(response, 'data-label="Статус"')

    def test_recipient_detail_includes_group_participant_appointment(self):
        other_child = Child.objects.create(last_name="Смирнов", first_name="Илья")
        assistant = StaffMember.objects.create(full_name="Ассистент карточки")
        start = _local_dt(timezone.localdate() + timedelta(days=3), time(11, 0))
        appointment = Appointment.objects.create(
            child=other_child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            session_type=Appointment.SessionType.GROUP,
        )
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=self.child,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )
        AppointmentStaffAssignment.objects.create(
            appointment=appointment,
            staff_member=assistant,
            role=AppointmentStaffAssignment.Role.ASSISTANT,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )

        response = self.client.get(reverse("recipient_detail", args=[self.child.pk]))

        self.assertEqual(response.status_code, 200)
        appointment_ids = {item.pk for item in response.context["upcoming_appointments"]}
        self.assertIn(appointment.pk, appointment_ids)
        self.assertContains(response, self.child.full_name)
        self.assertContains(response, other_child.full_name)
        self.assertContains(response, assistant.full_name)

    def test_recipient_list_counts_group_participant_appointments(self):
        other_child = Child.objects.create(last_name="Смирнов", first_name="Илья")
        start = _local_dt(timezone.localdate() + timedelta(days=3), time(11, 0))
        appointment = Appointment.objects.create(
            child=other_child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
            session_type=Appointment.SessionType.GROUP,
        )
        AppointmentParticipant.objects.create(
            appointment=appointment,
            child=self.child,
            starts_at_snapshot=appointment.starts_at,
            ends_at_snapshot=appointment.ends_at,
            appointment_status=appointment.status,
        )

        response = self.client.get(reverse("recipient_list"))

        self.assertEqual(response.status_code, 200)
        recipient = next(
            item for item in response.context["recipients"] if item.pk == self.child.pk
        )
        self.assertEqual(recipient.appointments_count, 1)

    def test_recipient_list_counts_legacy_appointment_without_participant(self):
        start = _local_dt(timezone.localdate() + timedelta(days=3), time(11, 0))
        appointment = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=start,
            ends_at=start + timedelta(minutes=30),
        )
        appointment.participants.all().delete()

        response = self.client.get(reverse("recipient_list"))

        self.assertEqual(response.status_code, 200)
        recipient = next(
            item for item in response.context["recipients"] if item.pk == self.child.pk
        )
        self.assertEqual(recipient.appointments_count, 1)

    def test_recipient_list_shows_operator_context(self):
        response = self.client.get(reverse("recipient_list"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("recipient_list_summary_items", response.context)
        self.assertIn("recipient_list_next_action", response.context)
        self.assertIn("recipient_list_control_items", response.context)
        self.assertContains(response, "Следующее действие")
        self.assertContains(response, "Контроль получателей")
        self.assertContains(response, 'id="recipient-list"')
        self.assertContains(response, reverse("recipient_import_preview"))
        self.assertContains(response, "directory-table")
        self.assertContains(response, 'data-label="Телефон"')
        self.assertContains(response, 'data-label="Действия"')

    def test_recipient_list_filters_before_limit(self):
        for index in range(85):
            Child.objects.create(last_name=f"Ааа{index:03d}", first_name="Получатель")
        target = Child.objects.create(last_name="ЯяФильтр", first_name="Найден")

        response = self.client.get(reverse("recipient_list"), {"q": "ЯяФильтр"})

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, target.full_name)
        self.assertEqual([item.pk for item in response.context["recipients"]], [target.pk])

    def test_recipient_edit_get(self):
        response = self.client.get(reverse("recipient_edit", args=[self.child.pk]))
        self.assertEqual(response.status_code, 200)

    def test_recipient_and_representative_forms_show_operator_control(self):
        link = RecipientRepresentative.objects.get(
            child=self.child,
            representative=self.parent,
        )
        cases = [
            (reverse("recipient_create"), "Контроль получателя", "Основной представитель"),
            (
                reverse("recipient_edit", args=[self.child.pk]),
                "Контроль получателя",
                "Дополнительные представители",
            ),
            (reverse("representative_create"), "Контроль контакта", "Email"),
            (
                reverse("representative_edit", args=[self.parent.pk]),
                "Контроль контакта",
                "Связь с получателем",
            ),
            (
                reverse("recipient_representative_create", args=[self.child.pk]),
                "Контроль представителя",
                "Расписание",
            ),
            (
                reverse("recipient_representative_edit", args=[link.pk]),
                "Контроль представителя",
                "Текущая связь",
            ),
        ]

        for url, title, control_item in cases:
            with self.subTest(url=url):
                response = self.client.get(url)
                self.assertEqual(response.status_code, 200)
                self.assertIn("object_form_control_items", response.context)
                self.assertContains(response, title)
                self.assertContains(response, control_item)

    def test_recipient_edit_post(self):
        response = self.client.post(
            reverse("recipient_edit", args=[self.child.pk]),
            {
                "last_name": "Иванов",
                "first_name": "Ваня",
                "middle_name": "Петрович",
                "birth_date": "2017-01-01",
                "status": Child.Status.ACTIVE,
                "primary_parent": self.parent.id,
                "diagnosis": "тест",
                "notes": "примечание",
            },
        )
        self.assertEqual(response.status_code, 302)
        self.child.refresh_from_db()
        self.assertEqual(self.child.middle_name, "Петрович")

    def test_representative_edit_get(self):
        response = self.client.get(reverse("representative_edit", args=[self.parent.pk]))
        self.assertEqual(response.status_code, 200)

    def test_representative_edit_post(self):
        response = self.client.post(
            reverse("representative_edit", args=[self.parent.pk]),
            {
                "last_name": "Иванова",
                "first_name": "Мария",
                "middle_name": "Сергеевна",
                "relationship_type": ParentGuardian.RelationshipType.MOTHER,
                "phone": "+7 900 000-00-99",
                "phone_alt": "",
                "email": self.parent.email,
                "notes": "",
            },
        )
        self.assertEqual(response.status_code, 302)
        self.parent.refresh_from_db()
        self.assertEqual(self.parent.middle_name, "Сергеевна")

    def test_recipient_create_with_representative_id(self):
        response = self.client.get(
            reverse("recipient_create"),
            {"representative_id": self.parent.id},
        )
        self.assertEqual(response.status_code, 200)

    def test_recipient_representative_create_additional(self):
        representative = ParentGuardian.objects.create(
            last_name="Петров",
            first_name="Сергей",
            relationship_type=ParentGuardian.RelationshipType.FATHER,
            phone="+7 900 000-00-02",
        )

        response = self.client.post(
            reverse("recipient_representative_create", args=[self.child.pk]),
            {
                "representative": representative.pk,
                "relationship_type": ParentGuardian.RelationshipType.FATHER,
                "receives_schedule": "on",
                "is_payer": "on",
                "notes": "Можно отправлять расписание",
            },
        )

        self.assertRedirects(response, reverse("recipient_detail", args=[self.child.pk]))
        link = RecipientRepresentative.objects.get(
            child=self.child,
            representative=representative,
        )
        self.assertFalse(link.is_primary)
        self.assertFalse(link.signs_contract)
        self.assertTrue(link.receives_schedule)
        self.assertTrue(link.is_payer)
        self.child.refresh_from_db()
        self.assertEqual(self.child.primary_parent, self.parent)

    def test_recipient_representative_create_primary_switches_signer(self):
        representative = ParentGuardian.objects.create(
            last_name="Сидорова",
            first_name="Анна",
            relationship_type=ParentGuardian.RelationshipType.GUARDIAN,
            phone="+7 900 000-00-03",
        )

        response = self.client.post(
            reverse("recipient_representative_create", args=[self.child.pk]),
            {
                "representative": representative.pk,
                "relationship_type": ParentGuardian.RelationshipType.GUARDIAN,
                "is_primary": "on",
                "receives_schedule": "on",
                "notes": "",
            },
        )

        self.assertRedirects(response, reverse("recipient_detail", args=[self.child.pk]))
        self.child.refresh_from_db()
        self.assertEqual(self.child.primary_parent, representative)
        new_link = RecipientRepresentative.objects.get(
            child=self.child,
            representative=representative,
        )
        old_link = RecipientRepresentative.objects.get(
            child=self.child,
            representative=self.parent,
        )
        self.assertTrue(new_link.is_primary)
        self.assertTrue(new_link.signs_contract)
        self.assertFalse(old_link.is_primary)
        self.assertFalse(old_link.signs_contract)


class SafeNextUrlTests(NewViewsTestBase):
    def test_safe_next_url_allows_same_host(self):
        from operations.views._common import safe_next_url

        same_host = "http://testserver/somewhere/"

        class FakeReq:
            POST = {"next": same_host}
            GET: dict = {}

            def get_host(self):
                return "testserver"

            def is_secure(self):
                return False

        result = safe_next_url(FakeReq(), "/fallback/")
        self.assertEqual(result, same_host)

    def test_safe_next_url_rejects_external(self):
        from operations.views._common import safe_next_url

        class FakeReq:
            POST: dict = {}
            GET = {"next": "http://evil.com/phish"}

            def get_host(self):
                return "testserver"

            def is_secure(self):
                return False

        result = safe_next_url(FakeReq(), "/fallback/")
        self.assertEqual(result, "/fallback/")

    def test_safe_next_url_falls_back_when_no_next(self):
        from operations.views._common import safe_next_url

        class FakeReq:
            POST: dict = {}
            GET: dict = {}

            def get_host(self):
                return "testserver"

            def is_secure(self):
                return False

        result = safe_next_url(FakeReq(), "/fallback/")
        self.assertEqual(result, "/fallback/")

    def test_csrf_failure_for_non_login(self):
        from operations.views._common import csrf_failure

        response = csrf_failure(self.client.get("/some/random/path/").wsgi_request, "test reason")
        self.assertEqual(response.status_code, 403)

    def test_csrf_failure_for_login_path(self):
        from operations.views._common import csrf_failure

        login_path = reverse("login")
        # Use a real client GET to login URL so MessageMiddleware is active
        response = csrf_failure(self.client.get(login_path).wsgi_request, "test")
        self.assertEqual(response.status_code, 302)
        self.assertIn(reverse("login"), response.url)


class SendConfirmationTaskTests(NewViewsTestBase):
    @override_settings(EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend")
    def test_returns_false_for_missing_confirmation(self):
        from operations.tasks import send_appointment_confirmation_email

        result = send_appointment_confirmation_email.call(999999)
        self.assertFalse(result)

    @override_settings(EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend")
    def test_returns_false_for_missing_email(self):
        from operations.models import AppointmentConfirmation
        from operations.tasks import send_appointment_confirmation_email

        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
        )
        conf = AppointmentConfirmation.objects.create(
            appointment=appt,
            target_type=AppointmentConfirmation.TargetType.REPRESENTATIVE,
            subject="Тест",
            message="Тест",
            email="",
        )
        result = send_appointment_confirmation_email.call(conf.pk)
        self.assertFalse(result)
        conf.refresh_from_db()
        self.assertEqual(conf.delivery_status, AppointmentConfirmation.DeliveryStatus.FAILED)

    @override_settings(EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend")
    def test_returns_true_on_success(self):
        from operations.models import AppointmentConfirmation
        from operations.tasks import send_appointment_confirmation_email

        appt = Appointment.objects.create(
            child=self.child,
            service=self.service,
            staff_member=self.staff,
            room=self.room,
            starts_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 0)),
            ends_at=_local_dt(timezone.localdate() + timedelta(days=5), time(10, 30)),
            billing_account=self.account,
        )
        conf = AppointmentConfirmation.objects.create(
            appointment=appt,
            target_type=AppointmentConfirmation.TargetType.REPRESENTATIVE,
            subject="Тест",
            message="Тест",
            email="test@example.local",
        )
        result = send_appointment_confirmation_email.call(conf.pk)
        self.assertTrue(result)
        conf.refresh_from_db()
        self.assertEqual(conf.delivery_status, AppointmentConfirmation.DeliveryStatus.SENT)
        self.assertIsNotNone(conf.sent_at)
